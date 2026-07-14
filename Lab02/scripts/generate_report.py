from __future__ import annotations

import json
import re
from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image as PILImage
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    Image,
    KeepTogether,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.platypus.tableofcontents import TableOfContents


ROOT = Path(__file__).resolve().parents[1]
SCREENSHOT_DIR = ROOT / "evidence" / "screenshots"
REPORT_DIR = ROOT / "report"
DOCX_PATH = REPORT_DIR / "21127645_LeMinh_Lab02_BufferOverflow.docx"
PDF_PATH = REPORT_DIR / "21127645_LeMinh_Lab02_BufferOverflow.pdf"

NAVY = "203748"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GOLD = "B4822A"
PALE_BLUE = "E8EEF5"
PALE_RED = "FDE9E7"
MUTED = "5A6570"
TABLE_DXA = 9360


CHAPTERS: list[tuple[str, list[tuple[str, list[str], list[str]]]]] = [
    (
        "Chương 1. Giới thiệu",
        [
            ("Bối cảnh và mục tiêu", [
                "Buffer Overflow xuất hiện khi chương trình ghi vượt sức chứa của vùng nhớ đích. Trong kiến trúc web, input HTTP có thể đi qua gateway Python rồi kích hoạt lỗi ở module native C; vì vậy ranh giới HTTP không loại bỏ rủi ro memory corruption.",
                "Lab quan sát luồng Browser -> HTTP POST -> Flask -> subprocess -> chương trình C, so sánh bản dùng strcpy với hai bản vá và binary hardened.",
            ], ["Nhận biết nguyên nhân ghi vượt.", "Quan sát AddressSanitizer, signal/exit code và GDB local.", "Chứng minh defense in depth bằng kiểm tra ở browser, Flask, C và compiler."]),
            ("Phạm vi và an toàn", [
                "Mọi thao tác chỉ chạy trên 127.0.0.1, Linux VM/WSL/Docker local. Input chỉ là văn bản bình thường hoặc ký tự lặp A/B; không có shellcode, ROP, reverse shell, persistence hay thay đổi return address.",
                "Lab chỉ chứng minh crash/dừng có kiểm soát và bản vá. Ngưỡng crash phụ thuộc kiến trúc, ABI, compiler, flags và stack layout nên chỉ được kết luận từ bằng chứng của lần chạy thật.",
            ], []),
        ],
    ),
    (
        "Chương 2. Cơ sở lý thuyết",
        [
            ("Bộ nhớ và chuỗi C", [
                "Buffer là vùng nhớ có kích thước hữu hạn. Stack lưu stack frame của lời gọi hàm, thường gồm biến local, metadata của frame và dữ liệu điều khiển; thứ tự cụ thể phụ thuộc ABI/compiler. Chuỗi C kết thúc bằng byte null, vì vậy name[32] chỉ chứa an toàn tối đa 31 byte dữ liệu.",
                "Stack Buffer Overflow là một dạng memory corruption khi ghi vượt biến local trên stack. Vùng liền kề có thể bị thay đổi, dẫn đến kết quả sai, abort hoặc signal; không phải mọi lần ghi vượt đều crash ngay.",
            ], []),
            ("API xử lý chuỗi", [
                "strcpy, gets và sprintf không biết sức chứa đích theo cách đủ để ngăn ghi vượt. snprintf giới hạn số byte ghi nhưng vẫn phải kiểm tra giá trị trả về để phát hiện truncate. strncpy có thể không thêm null terminator khi nguồn quá dài, nên đổi tên hàm mà không kiểm tra không phải bản vá hoàn chỉnh.",
            ], []),
            ("Công cụ và hardening", [
                "AddressSanitizer chèn instrumentation để phát hiện truy cập bộ nhớ sai trong kiểm thử; GDB dừng tiến trình, xem frame/local/backtrace. Stack Canary phát hiện giá trị bảo vệ bị đổi trước khi return. ASLR ngẫu nhiên hóa vị trí vùng nhớ; PIE cho executable tham gia ASLR; DEP/NX đánh dấu vùng dữ liệu không thực thi; RELRO bảo vệ một số cấu trúc dynamic linking; FORTIFY_SOURCE bổ sung kiểm tra cho một số lời gọi thư viện.",
                "Các cơ chế này giảm rủi ro hoặc tăng khả năng phát hiện, không thay thế việc loại bỏ strcpy không giới hạn.",
            ], []),
        ],
    ),
    (
        "Chương 3. Kiến trúc hệ thống",
        [
            ("Ba tầng", [
                "Browser UI thu input và mode. Flask áp dụng request limit, allowlist mode và gọi subprocess.run bằng danh sách argument với shell=False, timeout và working directory cố định. Native processor nhận đúng một argument, in metadata rồi xử lý chuỗi. Hệ điều hành trả exit code hoặc signal; Flask chuẩn hóa kết quả thành trace và HTTP response.",
            ], ["Browser: biểu mẫu, timeline và inspector.", "Flask: routing, validation, subprocess và trace.", "C/OS: buffer, copy, hardening, exit/signal."]),
            ("Sequence diagram", [
                "Browser --POST /submit--> Flask Gateway\nFlask --subprocess.run([binary, name], shell=False)--> C Processor\nC Processor --exit code / signal / stderr--> Operating System\nOperating System --kết quả tiến trình--> Flask Gateway\nFlask Gateway --HTTP response + Trace ID--> Browser",
            ], []),
        ],
    ),
    (
        "Chương 4. Thiết kế chương trình lỗi",
        [
            ("Mã nguồn và nguyên nhân", [
                "Hàm process_name khai báo char name[32] rồi gọi strcpy(name, user_input). strcpy tiếp tục copy cho tới byte null của nguồn nhưng không nhận kích thước name. Khi input có 32 byte dữ liệu, lần ghi null terminator đã vượt ít nhất một byte.",
                "Sơ đồ stack trong UI/report chỉ là mô hình giáo dục: vùng argument, return address, saved frame pointer, canary (nếu có) và name[32]. Vị trí thực tế phải quan sát bằng GDB trong đúng binary.",
            ], []),
        ],
    ),
    (
        "Chương 5. Thực nghiệm input bình thường",
        [
            ("Kịch bản Le Minh", [
                "Input `Le Minh` gồm 7 ký tự ASCII và 7 byte UTF-8. Kỳ vọng kỹ thuật là dữ liệu cùng null terminator nằm trong name[32]. Request, stdout, PID, exit code và thời gian chỉ được coi là kết quả thực nghiệm khi trace/log thật tồn tại.",
            ], ["Kiểm tra POST /submit và mode.", "Đối chiếu stdout/exit code trong Native Inspector.", "Đối chiếu visualizer: overflow_bytes = 0."]),
        ],
    ),
    (
        "Chương 6. Thực nghiệm input dài",
        [
            ("Các mốc", [
                "31 byte dữ liệu vừa đủ cùng một null trong buffer 32 byte. 32 byte dữ liệu cần byte null thứ 33 nên đã ghi vượt về mặt mô hình chuỗi C, nhưng không được đồng nhất với crash. Các input 33, 64 và 128 byte dùng để quan sát ASan/exit/signal trong giới hạn lab.",
                "Ba đại lượng phải tách biệt: độ dài đầu tiên ghi vượt theo capacity; độ dài đầu tiên ASan báo trong dữ liệu chạy thật; và độ dài đầu tiên tiến trình crash/dừng theo dữ liệu chạy thật.",
            ], []),
            ("Giới hạn kết luận", [
                "Báo cáo không khẳng định ngưỡng crash giống nhau trên mọi máy. Ký tự Unicode có thể chiếm nhiều byte UTF-8, nên validation và capacity phải dựa trên byte khi C copy byte.",
            ], []),
        ],
    ),
    (
        "Chương 7. Phân tích GDB",
        [
            ("Quy trình quan sát", [
                "Dùng binary vulnerable_debug với symbol -g, frame pointer và tối ưu O0. Break tại process_name, xem frame, info locals, sizeof(name), vùng byte quanh name, rồi tiếp tục để quan sát signal/backtrace. Địa chỉ chỉ xuất hiện trong terminal GDB local.",
                "Phiên overflow chỉ dùng 64 ký tự A. Không sửa register, không ghi memory, không nhảy địa chỉ tùy ý và không có nội dung khai thác thực thi mã.",
            ], []),
            ("Giới hạn bằng chứng", [
                "Nếu evidence/gdb chưa có log, báo cáo ghi trạng thái chưa chạy; script .gdb và hướng dẫn không được xem là kết quả thực nghiệm.",
            ], []),
        ],
    ),
    (
        "Chương 8. Bản vá 1 - kiểm tra độ dài",
        [
            ("Thiết kế", [
                "secure_length dùng strnlen với giới hạn hợp lý. Nếu input vượt 31 byte, hàm trả exit code riêng trước mọi thao tác copy. Input hợp lệ được copy đúng số byte bằng memcpy và đặt name[length] = '\\0'.",
                "Kiểm tra lại ở C chứng minh defense in depth: bỏ qua thuộc tính maxlength của browser hoặc gọi thẳng endpoint vẫn không thể ép native copy quá capacity.",
            ], []),
        ],
    ),
    (
        "Chương 9. Bản vá 2 - snprintf",
        [
            ("Thiết kế", [
                "secure_snprintf gọi snprintf(name, sizeof(name), \"%s\", user_input). Nếu giá trị trả về âm, xử lý lỗi format/encoding; nếu lớn hơn hoặc bằng sizeof(name), phát hiện truncate và từ chối thay vì chấp nhận tên bị cắt.",
                "Giới hạn ghi ngăn overflow trong name, còn kiểm tra return value đảm bảo chính sách dữ liệu rõ ràng. Đây là lý do chỉ thay strcpy bằng snprintf mà bỏ qua return value vẫn chưa đủ.",
            ], []),
        ],
    ),
    (
        "Chương 10. Request limit",
        [
            ("Defense in depth", [
                "Browser giới hạn thao tác thuận tiện, Flask đặt MAX_CONTENT_LENGTH và giới hạn name theo byte UTF-8, còn chương trình C tự áp capacity 31 byte cho hai bản vá. Mỗi lớp xử lý một ranh giới tin cậy khác nhau.",
                "Chỉ giới hạn frontend không đủ vì client nằm ngoài vùng tin cậy và request có thể được tạo không qua form. Firewall cũng không hiểu điều kiện copy nội bộ; request hợp lệ về mạng vẫn có thể mang input dài tới backend native.",
            ], []),
        ],
    ),
    (
        "Chương 11. Compiler hardening",
        [
            ("Cơ chế", [
                "Stack protector kiểm tra canary; PIE hỗ trợ ASLR cho executable; Full RELRO khóa relocation sau khởi tạo; NX giữ stack/data không thực thi; FORTIFY_SOURCE thêm kiểm tra khi compiler biết kích thước object. Các thuộc tính phải được xác minh bằng file/readelf/objdump trên binary thật.",
                "vulnerable_debug chỉ tắt stack protector để quan sát và vẫn không dùng executable stack. ASan là profile kiểm thử. Các binary secure bật nhóm hardening nhưng secure coding vẫn là lớp chính.",
            ], []),
        ],
    ),
    (
        "Chương 12. So sánh trước và sau vá",
        [
            ("Kết quả thiết kế", [
                "Bản vulnerable copy không giới hạn nên input trên capacity tạo memory write ngoài name. Bản length từ chối trước copy; bản snprintf giới hạn write và từ chối truncate. Exit code của hai bản vá thể hiện validation failure có kiểm soát, không phải crash.",
            ], []),
        ],
    ),
    (
        "Chương 13. Mức độ ảnh hưởng",
        [
            ("Tác động", [
                "Memory corruption có thể làm tiến trình crash, gây từ chối dịch vụ, làm sai dữ liệu hoặc rò rỉ dữ liệu trong các chương trình thực tế. Về lý thuyết, control data bị hỏng có thể đổi luồng điều khiển, nhưng lab không mô tả hay thực hiện quy trình khai thác.",
                "Mức độ thực tế phụ thuộc dữ liệu kề buffer, quyền tiến trình, khả năng lặp lại, hardening và cách dịch vụ phục hồi. Tiến trình nên chạy least privilege và được giám sát/restart có kiểm soát.",
            ], []),
        ],
    ),
    (
        "Chương 14. Phòng chống",
        [
            ("Biện pháp", [
                "Phòng chống bắt đầu bằng kiểm tra kích thước theo byte, API có giới hạn và kiểm tra return value. Request limit, compiler/OS hardening, sanitizer trong CI, fuzz testing, sandbox, least privilege, logging và monitoring tạo defense in depth.",
                "Với parser nhạy cảm, ưu tiên thư viện đã được kiểm thử hoặc ngôn ngữ memory-safe khi phù hợp; hardening không phải lý do giữ lại thao tác copy không an toàn.",
            ], ["Secure coding và validation tại ranh giới C.", "Sanitizer/fuzzing trong pipeline kiểm thử.", "Canary, PIE/ASLR, RELRO, NX, FORTIFY ở build/OS.", "Giảm quyền, sandbox, log và giám sát ở runtime."]),
        ],
    ),
    (
        "Chương 15. Trả lời câu hỏi báo cáo",
        [
            ("1. Buffer Overflow khác Injection như thế nào?", [
                "Buffer Overflow là vi phạm an toàn bộ nhớ do ghi vượt object; Injection xảy ra khi dữ liệu bị interpreter hiểu thành cú pháp/lệnh. Một lỗi làm hỏng memory layout, lỗi kia làm thay đổi ngữ nghĩa của ngôn ngữ đích như SQL hoặc shell.",
            ], []),
            ("2. Vì sao backend native bị kích hoạt qua HTTP?", [
                "HTTP chỉ là kênh vận chuyển. Flask trích trường name rồi truyền nó thành argument cho chương trình C; nếu C copy không kiểm tra, dữ liệu web đi qua các tầng và chạm trực tiếp buffer local.",
            ], []),
            ("3. Vì sao firewall không đủ?", [
                "Firewall kiểm soát kết nối/luồng mạng, không biết capacity của name[32] hay semantics của strcpy. Request tới đúng local service vẫn có thể chứa dữ liệu dài hợp lệ ở tầng HTTP.",
            ], []),
            ("4. Bản vá hoạt động ra sao?", [
                "Bản vá length đo byte và từ chối trên 31 trước copy. Bản vá snprintf giới hạn số byte ghi rồi kiểm tra return value để từ chối truncate. Cả hai giữ kiểm tra tại C thay vì chỉ tin browser/Flask.",
            ], []),
            ("5. Ít nhất ba cơ chế hardening", [
                "Stack Canary phát hiện giá trị bảo vệ bị đổi; PIE kết hợp ASLR làm địa chỉ executable thay đổi; DEP/NX ngăn thực thi tại vùng data. RELRO bảo vệ relocation và FORTIFY_SOURCE tăng kiểm tra thư viện khi compiler có đủ thông tin.",
            ], []),
            ("Câu hỏi bổ sung", [
                "Buffer name là biến local trong stack frame của process_name. Khi input vượt buffer, các byte ngoài object bị ghi và chương trình có undefined behavior. strcpy/gets/sprintf nguy hiểm vì không gắn thao tác với sức chứa đích; memory corruption nghiêm trọng vì có thể ảnh hưởng dữ liệu và control flow, không chỉ một điều kiện logic.",
            ], []),
        ],
    ),
    (
        "Chương 16. Kiểm thử",
        [
            ("Phạm vi kiểm thử", [
                "Pytest bao phủ route/healthcheck, input bình thường, ASan overflow, hai bản vá, request limit và các ràng buộc an toàn. Native tests chạy binary thật; report/screenshot tests kiểm tra artifact và manifest mà không OCR.",
                "Kết quả chỉ được trích từ evidence/logs/pytest.txt của lần chạy thật. Nếu file vắng hoặc không kết thúc thành công, báo cáo không tuyên bố toàn bộ test pass.",
            ], []),
        ],
    ),
    (
        "Chương 17. Kết luận",
        [
            ("Kết quả và bài học", [
                "Thiết kế lab cho thấy một input HTTP có thể tới buffer C, cách null terminator tạo ghi vượt tại ranh giới 32 byte, và cách sanitizer/debugger hỗ trợ quan sát. Hai bản vá loại bỏ ghi vượt theo hai chính sách rõ ràng, còn hardening tạo lớp giảm thiểu bổ sung.",
                "Giới hạn chính là signal/ngưỡng crash và stack layout phụ thuộc môi trường; ảnh GDB phải chụp thủ công. Hướng phát triển an toàn là mở rộng fuzz test có giới hạn, chạy sanitizer trong CI và tiếp tục giảm phạm vi code native.",
            ], []),
        ],
    ),
]


SHOTS = [
    (1, "01_home_overview.png", "http://127.0.0.1:5002/", "Mở trang chủ.", "Mục tiêu, kiến trúc và các mode.", "Tổng quan kiến trúc và chế độ thực hành."),
    (5, "02_normal_input_before_submit.png", "http://127.0.0.1:5002/vulnerable", "Nhập Le Minh, chụp trước khi gửi.", "Input, mode và giới hạn.", "Input bình thường trước khi gửi."),
    (5, "03_normal_http_request.png", "POST /submit", "Gửi Le Minh, mở Request Inspector.", "POST /submit, Content-Type, length và mode.", "HTTP request của input bình thường."),
    (5, "04_normal_native_process.png", "http://127.0.0.1:5002/vulnerable", "Mở Native Process Inspector.", "Exit code/stdout thực tế; dự kiến exit 0 cho Le Minh.", "Kết quả tiến trình native với input bình thường."),
    (5, "05_normal_memory_visualizer.png", "http://127.0.0.1:5002/vulnerable", "Mở Memory Visualizer.", "7 byte dữ liệu, null và overflow 0.", "Mô hình bộ nhớ không overflow."),
    (6, "06_overflow_32_input.png", "http://127.0.0.1:5002/vulnerable", "Nhập A lặp 32.", "32 byte dữ liệu và giải thích null terminator.", "Input 32 byte tại ranh giới chuỗi C."),
    (6, "07_overflow_32_memory_boundary.png", "POST /submit", "Gửi 32 byte, mở visualizer.", "Ít nhất 1 byte ngoài biên; không khẳng định crash.", "Ranh giới ghi vượt ở input 32 byte."),
    (6, "08_overflow_64_request.png", "POST /submit", "Gửi A lặp 64 ở vulnerable_asan.", "Request và input length 64 byte.", "Request 64 byte tới profile ASan."),
    (6, "09_overflow_64_strcpy_step.png", "http://127.0.0.1:5002/vulnerable", "Chọn bước strcpy.", "strcpy không nhận kích thước đích.", "Bước copy không giới hạn của strcpy."),
    (6, "10_overflow_64_memory_visualizer.png", "http://127.0.0.1:5002/vulnerable", "Mở visualizer của trace 64 byte.", "Buffer 32 byte và vùng overflow.", "Mô hình vùng ghi vượt với input 64 byte."),
    (6, "11_asan_detected.png", "http://127.0.0.1:5002/vulnerable", "Mở ASan Inspector.", "stack-buffer-overflow từ log thật.", "AddressSanitizer phát hiện lỗi."),
    (6, "12_asan_stack_trace.png", "http://127.0.0.1:5002/vulnerable", "Mở chi tiết stack trace.", "File, process_name và dòng strcpy.", "Stack trace ASan định vị nguồn lỗi."),
    (6, "13_native_crash_result.png", "POST /submit", "Mở Native Inspector sau trace 64.", "Exit code/signal thực tế.", "Trạng thái dừng của tiến trình native."),
    (6, "14_final_vulnerable_verdict.png", "http://127.0.0.1:5002/vulnerable", "Mở Final Security Verdict.", "Overflow, detection/crash thực tế và nguyên nhân.", "Kết luận bảo mật của bản vulnerable."),
    (6, "15_length_test_table.png", "python scripts/test_lengths.py", "Chạy test độ dài và chụp bảng thật.", "Các length, HTTP, exit/signal, ASan/crash.", "Kết quả kiểm thử theo độ dài."),
    (7, "16_gdb_breakpoint.png", "gdb -q -x gdb/inspect_normal.gdb", "Chụp khi dừng tại process_name.", "Breakpoint, file và dòng.", "GDB dừng tại process_name."),
    (7, "17_gdb_local_buffer.png", "p sizeof(name); x/64bx &name", "Quan sát local buffer.", "sizeof(name)=32 hoặc vùng byte local.", "Biến local name[32] trong GDB."),
    (7, "18_gdb_overflow_stop.png", "gdb -q -x gdb/inspect_overflow.gdb", "Chạy A lặp 64, chụp khi dừng.", "Signal/trạng thái và backtrace thật.", "Phiên overflow được quan sát bằng GDB."),
    (8, "19_secure_length_reject.png", "http://127.0.0.1:5002/secure/length", "Gửi A lặp 64.", "Từ chối trước copy và exit code thực tế.", "Bản vá length từ chối input dài."),
    (8, "20_secure_length_timeline.png", "http://127.0.0.1:5002/secure/length", "Chọn bước strnlen/compare/reject.", "Đo byte, so sánh 31 và không ghi buffer.", "Timeline kiểm tra độ dài trước copy."),
    (9, "21_secure_snprintf_reject.png", "http://127.0.0.1:5002/secure/snprintf", "Gửi A lặp 64, xem return value.", "Phát hiện truncate và từ chối.", "Bản vá snprintf từ chối truncate."),
    (9, "22_code_comparison.png", "http://127.0.0.1:5002/comparison", "Mở ba cột so sánh.", "strcpy, length check và snprintf return value.", "So sánh mã vulnerable và hai bản vá."),
    (11, "23_hardening_comparison.png", "http://127.0.0.1:5002/hardening", "Mở Hardening Inspector sau make all.", "Canary, PIE, RELRO, NX, FORTIFY từ công cụ thật.", "So sánh thuộc tính hardening."),
    (11, "24_stack_canary_explanation.png", "http://127.0.0.1:5002/hardening", "Mở giải thích canary.", "Vị trí khái niệm và kiểm tra trước return.", "Cơ chế Stack Canary."),
    (11, "25_asan_vs_hardening.png", "http://127.0.0.1:5002/hardening", "Mở bảng ASan vs hardening.", "Công cụ test khác lớp production.", "Phân biệt ASan và hardening."),
    (16, "26_presentation_mode.png", "http://127.0.0.1:5002/vulnerable", "Bật Presentation Mode cho trace có sẵn.", "Một bước chữ lớn và điều hướng.", "Presentation Mode của timeline."),
    (16, "27_pytest_passed.png", "sh scripts/run_tests.sh", "Chụp tổng kết pytest thật.", "Kết quả hiện tại; chỉ pass nếu exit 0.", "Kết quả pytest trong WSL."),
    (17, "28_report_files.png", "python scripts/generate_report.py", "Liệt kê report sau khi tạo.", "DOCX và PDF khác 0 byte.", "Hai artifact báo cáo được sinh."),
]


def safe_text(value: object) -> str:
    return str(value if value is not None else "-").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def read_json(path: Path) -> object | None:
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None


def observed_summary() -> dict[str, str]:
    summary: dict[str, str] = {}
    normal = read_json(ROOT / "evidence" / "traces" / "normal_vulnerable.json")
    if isinstance(normal, dict):
        summary["normal"] = (
            f"Trace thật: exit_code={normal.get('exit_code')}, signal={normal.get('signal')}, "
            f"ASan={normal.get('asan_detected')}, overflow_bytes={normal.get('overflow_bytes')}."
        )
    else:
        summary["normal"] = "Chưa có evidence/traces/normal_vulnerable.json; chưa kết luận kết quả chạy input bình thường."

    length_data = read_json(ROOT / "evidence" / "logs" / "length_test.json")
    rows = length_data if isinstance(length_data, list) else length_data.get("results", []) if isinstance(length_data, dict) else []
    if rows:
        first_asan = next((row.get("length") for row in rows if row.get("asan_detected")), None)
        first_crash = next((row.get("length") for row in rows if row.get("crash_detected")), None)
        summary["lengths"] = (
            "Theo length_test.json thật: first ASan="
            f"{first_asan if first_asan is not None else 'chưa quan sát'}, first crash="
            f"{first_crash if first_crash is not None else 'chưa quan sát'}."
        )
    else:
        summary["lengths"] = "Chưa có evidence/logs/length_test.json; ngưỡng ASan và crash chưa được quan sát."

    pytest_log = ROOT / "evidence" / "logs" / "pytest.txt"
    if pytest_log.exists():
        lines = [line.strip() for line in pytest_log.read_text(encoding="utf-8", errors="replace").splitlines() if line.strip()]
        summary["pytest"] = "Dòng cuối log pytest thật: " + (lines[-1] if lines else "log rỗng")
    else:
        summary["pytest"] = "Chưa có evidence/logs/pytest.txt; chưa tuyên bố toàn bộ test pass."

    gdb_logs = [ROOT / "evidence" / "gdb" / name for name in ("normal_session.txt", "overflow_session.txt", "hardened_session.txt")]
    present = [path.name for path in gdb_logs if path.exists() and path.stat().st_size]
    summary["gdb"] = f"Log GDB thật hiện có: {', '.join(present)}." if present else "Chưa có log GDB thật; chỉ có script/hướng dẫn để chạy local."
    return summary


def set_run_font(run, name: str = "Calibri", size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_field(paragraph) -> None:
    run = paragraph.add_run("Trang ")
    set_run_font(run, size=9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MUTED)
    r_pr.append(color)
    r.append(r_pr)
    text = OxmlElement("w:t")
    text.text = "1"
    r.append(text)
    field.append(r)
    paragraph._p.append(field)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        existing = tbl_pr.find(qn(tag))
        if existing is not None:
            tbl_pr.remove(existing)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    tbl_pr.append(indent)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_docx_table(doc: Document, headers: list[str], rows: Iterable[Iterable[object]], widths: list[int], caption: str) -> None:
    caption_p = doc.add_paragraph(style="Caption")
    caption_p.paragraph_format.space_before = Pt(6)
    caption_p.paragraph_format.space_after = Pt(4)
    caption_p.paragraph_format.keep_with_next = True
    run = caption_p.add_run(caption)
    set_run_font(run, size=9.5, color=DARK_BLUE, bold=True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for cell, header in zip(table.rows[0].cells, headers):
        shade_cell(cell, PALE_BLUE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(header)
        set_run_font(run, size=9, bold=True, color=NAVY)
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            set_run_font(run, size=8.5)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def configure_docx(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9.5)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("LAB 02  |  BUFFER OVERFLOW TRONG ỨNG DỤNG WEB LOCAL")
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(footer)


def add_cover(doc: Document) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(90)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    set_run_font(p.add_run("BÁO CÁO THỰC HÀNH AN TOÀN ỨNG DỤNG WEB"), size=10.5, color=GOLD, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_run_font(p.add_run("LAB 2"), size=30, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    set_run_font(p.add_run("BUFFER OVERFLOW TRONG\nỨNG DỤNG WEB LOCAL"), size=19, color=DARK_BLUE, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(64)
    set_run_font(p.add_run("Từ HTTP POST tới stack frame, AddressSanitizer và bản vá"), size=11, color=MUTED, italic=True)
    for label, value in (("MSSV", "21127645"), ("Họ tên", "Lê Minh"), ("Ngày lập báo cáo", date.today().strftime("%d/%m/%Y"))):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        set_run_font(p.add_run(f"{label}: "), size=11, color=NAVY, bold=True)
        set_run_font(p.add_run(value), size=11, color=NAVY)
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    doc.add_heading("Mục lục", level=1)
    note = doc.add_paragraph("Danh mục chương tĩnh bên dưới luôn đọc được; số trang thật được đặt bằng trường PAGE ở chân trang.")
    note.runs[0].italic = True
    for title, _ in CHAPTERS:
        paragraph = doc.add_paragraph(title)
        paragraph.paragraph_format.left_indent = Inches(0.2)
    paragraph = doc.add_paragraph("Phụ lục. Bằng chứng và câu lệnh")
    paragraph.paragraph_format.left_indent = Inches(0.2)
    doc.add_page_break()


def picture_size(path: Path, max_w: float = 6.25, max_h: float = 6.7) -> tuple[float, float]:
    with PILImage.open(path) as image:
        width, height = image.size
    scale = min(max_w / width, max_h / height)
    return width * scale, height * scale


def add_docx_shot(doc: Document, number: int, shot: tuple[int, str, str, str, str, str], missing: list[str]) -> None:
    _, filename, location, action, required, caption = shot
    path = SCREENSHOT_DIR / filename
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(8)
    heading.paragraph_format.space_after = Pt(4)
    heading.paragraph_format.keep_with_next = True
    set_run_font(heading.add_run(filename), size=10.5, color=NAVY, bold=True)
    if path.exists() and path.stat().st_size:
        width, height = picture_size(path)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        p.add_run().add_picture(str(path), width=Inches(width), height=Inches(height))
    else:
        missing.append(filename)
        table = doc.add_table(rows=4, cols=1)
        table.style = "Table Grid"
        values = (
            f"ẢNH THỦ CÔNG CẦN BỔ SUNG: {filename}",
            f"URL hoặc lệnh: {location}",
            f"Thao tác: {action}",
            f"Nội dung bắt buộc: {required}",
        )
        for index, (row, value) in enumerate(zip(table.rows, values)):
            if index == 0:
                shade_cell(row.cells[0], PALE_RED)
            run = row.cells[0].paragraphs[0].add_run(value)
            set_run_font(run, size=9, color=NAVY, bold=index == 0)
        set_table_geometry(table, [TABLE_DXA])
    p = doc.add_paragraph(style="Caption")
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.add_run(f"Hình {number}. {caption}")


def add_docx_special(doc: Document, chapter_number: int, observed: dict[str, str], table_counter: list[int]) -> None:
    if chapter_number == 3:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        run = p.add_run("Browser -> HTTP POST -> Flask -> subprocess.run -> C processor -> OS -> HTTP response")
        set_run_font(run, name="Consolas", size=9.5, color=NAVY, bold=True)
        cap = doc.add_paragraph("Sơ đồ 1. Sequence diagram khái niệm của hệ thống.", style="Caption")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif chapter_number == 4:
        code = "int process_name(const char *user_input) {\n    char name[32];\n    strcpy(name, user_input);\n    return 0;\n}"
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.right_indent = Inches(0.25)
        shade = OxmlElement("w:shd")
        shade.set(qn("w:fill"), "F2F4F7")
        p._p.get_or_add_pPr().append(shade)
        set_run_font(p.add_run(code), name="Consolas", size=9.5, color=NAVY)
        doc.add_paragraph("Đoạn mã 1. Mã nguồn tối thiểu chứa thao tác strcpy có chủ đích.", style="Caption")
    elif chapter_number == 5:
        doc.add_paragraph(observed["normal"])
    elif chapter_number == 6:
        doc.add_paragraph(observed["lengths"])
        table_counter[0] += 1
        add_docx_table(
            doc,
            ["Mốc", "Ý nghĩa", "Cách xác nhận"],
            [
                ("31 byte", "Capacity an toàn", "Mô hình name[32] + null"),
                ("32 byte", "Ít nhất 1 byte ghi vượt", "Mô hình copy tới null"),
                ("ASan đầu tiên", "Phát hiện instrumentation", "length_test.json thật"),
                ("Crash đầu tiên", "Dừng tiến trình", "exit/signal thật"),
            ],
            [1500, 3500, 4360],
            f"Bảng {table_counter[0]}. Phân biệt capacity, ASan và crash.",
        )
    elif chapter_number == 7:
        doc.add_paragraph(observed["gdb"])
    elif chapter_number == 11:
        table_counter[0] += 1
        add_docx_table(
            doc,
            ["Profile", "Mục đích", "Trạng thái hardening"],
            [
                ("vulnerable_debug", "GDB", "Phải xác minh bằng file/readelf"),
                ("vulnerable_asan", "Phát hiện lỗi", "ASan, không đại diện production"),
                ("secure_length", "Bản vá length", "Phải xác minh binary thật"),
                ("secure_snprintf", "Bản vá snprintf", "Phải xác minh binary thật"),
                ("secure_hardened", "So sánh hardening", "Phải xác minh binary thật"),
            ],
            [2200, 2800, 4360],
            f"Bảng {table_counter[0]}. Profile build và nguyên tắc xác minh.",
        )
    elif chapter_number == 12:
        table_counter[0] += 1
        add_docx_table(
            doc,
            ["Tiêu chí", "Vulnerable", "Secure length", "Secure snprintf"],
            [
                ("Hàm copy", "strcpy", "memcpy sau validation", "snprintf"),
                ("Giới hạn", "Không", "Tối đa 31 byte", "sizeof(name)"),
                ("Input dài", "Ghi vượt", "Từ chối trước copy", "Từ chối truncate"),
                ("Hardening", "Không thay bản vá", "Lớp bổ sung", "Lớp bổ sung"),
            ],
            [1700, 2300, 2680, 2680],
            f"Bảng {table_counter[0]}. So sánh trước và sau vá.",
        )
    elif chapter_number == 16:
        doc.add_paragraph(observed["pytest"])


def build_docx() -> list[str]:
    doc = Document()
    configure_docx(doc)
    add_cover(doc)
    add_toc(doc)
    observed = observed_summary()
    missing: list[str] = []
    figure = 0
    table_counter = [0]
    for chapter_number, (title, sections) in enumerate(CHAPTERS, 1):
        doc.add_heading(title, level=1)
        for section_title, paragraphs, bullets in sections:
            doc.add_heading(section_title, level=2)
            for text in paragraphs:
                for part in text.split("\n"):
                    doc.add_paragraph(part)
            for bullet in bullets:
                doc.add_paragraph(bullet, style="List Bullet")
        add_docx_special(doc, chapter_number, observed, table_counter)
        for shot in (item for item in SHOTS if item[0] == chapter_number):
            figure += 1
            add_docx_shot(doc, figure, shot, missing)

    doc.add_page_break()
    doc.add_heading("Phụ lục. Bằng chứng và câu lệnh", level=1)
    doc.add_heading("A. Mã nguồn và compiler flags", level=2)
    doc.add_paragraph("Mã nguồn: native/*.c, native/processor_common.h. Flags và target là nguồn sự thật trong Makefile; không suy diễn trạng thái hardening chỉ từ tên profile.")
    doc.add_heading("B. Request, response và trace", level=2)
    doc.add_paragraph("Request/response/trace thật nằm trong evidence/requests, evidence/responses và evidence/traces. File vắng được coi là chưa thu thập, không được thay bằng log mẫu giả.")
    doc.add_heading("C. ASan và GDB", level=2)
    doc.add_paragraph("ASan log nằm trong evidence/asan; GDB log thủ công nằm trong evidence/gdb. Địa chỉ chỉ quan sát local và không dùng để xây dựng payload.")
    doc.add_heading("D. Câu lệnh chạy", level=2)
    for command in (
        "make all",
        "python app.py",
        "python scripts/test_lengths.py",
        "gdb -q -x gdb/inspect_normal.gdb",
        "gdb -q -x gdb/inspect_overflow.gdb",
        "pytest",
        "python scripts/check_screenshots.py",
        "python scripts/generate_report.py",
    ):
        p = doc.add_paragraph(style="List Bullet")
        set_run_font(p.add_run(command), name="Consolas", size=9.5)
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)
    return missing


def register_pdf_fonts() -> tuple[str, str, str]:
    candidates = [
        (
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf"),
        ),
        (
            Path("C:/Windows/Fonts/arial.ttf"),
            Path("C:/Windows/Fonts/arialbd.ttf"),
            Path("C:/Windows/Fonts/ariali.ttf"),
        ),
        (
            Path("/mnt/c/Windows/Fonts/arial.ttf"),
            Path("/mnt/c/Windows/Fonts/arialbd.ttf"),
            Path("/mnt/c/Windows/Fonts/ariali.ttf"),
        ),
    ]
    for regular, bold, italic in candidates:
        if regular.exists() and bold.exists() and italic.exists():
            pdfmetrics.registerFont(TTFont("LabSans", str(regular)))
            pdfmetrics.registerFont(TTFont("LabSans-Bold", str(bold)))
            pdfmetrics.registerFont(TTFont("LabSans-Italic", str(italic)))
            pdfmetrics.registerFontFamily("LabSans", normal="LabSans", bold="LabSans-Bold", italic="LabSans-Italic", boldItalic="LabSans-Bold")
            return "LabSans", "LabSans-Bold", "LabSans-Italic"
    raise RuntimeError("Không tìm thấy font Unicode DejaVu Sans hoặc Arial để tạo PDF tiếng Việt.")


class ReportTemplate(BaseDocTemplate):
    def afterFlowable(self, flowable) -> None:
        if isinstance(flowable, Paragraph):
            style = flowable.style.name
            if style in ("PDF Heading 1", "PDF Heading 2"):
                level = 0 if style.endswith("1") else 1
                text = flowable.getPlainText()
                self.notify("TOCEntry", (level, text, self.page))


def pdf_styles(font: str, bold: str, italic: str) -> dict[str, ParagraphStyle]:
    sample = getSampleStyleSheet()
    return {
        "body": ParagraphStyle("PDF Body", parent=sample["BodyText"], fontName=font, fontSize=10.5, leading=14, spaceAfter=6, textColor=colors.HexColor("#1E2933")),
        "h1": ParagraphStyle("PDF Heading 1", parent=sample["Heading1"], fontName=bold, fontSize=16, leading=20, textColor=colors.HexColor(f"#{BLUE}"), spaceBefore=18, spaceAfter=10, keepWithNext=True),
        "h2": ParagraphStyle("PDF Heading 2", parent=sample["Heading2"], fontName=bold, fontSize=13, leading=16, textColor=colors.HexColor(f"#{BLUE}"), spaceBefore=14, spaceAfter=7, keepWithNext=True),
        "bullet": ParagraphStyle("PDF Bullet", parent=sample["BodyText"], fontName=font, fontSize=10.5, leading=14, leftIndent=27, firstLineIndent=-13.5, bulletIndent=13.5, spaceAfter=4),
        "caption": ParagraphStyle("PDF Caption", parent=sample["BodyText"], fontName=bold, fontSize=9, leading=12, alignment=TA_CENTER, textColor=colors.HexColor(f"#{DARK_BLUE}"), spaceBefore=3, spaceAfter=8),
        "small": ParagraphStyle("PDF Small", parent=sample["BodyText"], fontName=font, fontSize=8.5, leading=11, textColor=colors.HexColor(f"#{MUTED}")),
        "code": ParagraphStyle("PDF Code", parent=sample["Code"], fontName=font, fontSize=8.5, leading=11, leftIndent=16, rightIndent=16, backColor=colors.HexColor("#F2F4F7"), borderPadding=8, spaceAfter=8),
        "cover_kicker": ParagraphStyle("Cover Kicker", parent=sample["BodyText"], fontName=bold, fontSize=10.5, leading=13, alignment=TA_CENTER, textColor=colors.HexColor(f"#{GOLD}"), spaceAfter=16),
        "cover_title": ParagraphStyle("Cover Title", parent=sample["Title"], fontName=bold, fontSize=29, leading=34, alignment=TA_CENTER, textColor=colors.HexColor(f"#{NAVY}"), spaceAfter=8),
        "cover_sub": ParagraphStyle("Cover Subtitle", parent=sample["BodyText"], fontName=bold, fontSize=18, leading=23, alignment=TA_CENTER, textColor=colors.HexColor(f"#{DARK_BLUE}"), spaceAfter=24),
        "cover_meta": ParagraphStyle("Cover Meta", parent=sample["BodyText"], fontName=font, fontSize=11, leading=16, alignment=TA_CENTER, textColor=colors.HexColor(f"#{NAVY}")),
    }


def on_pdf_page(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont("LabSans-Bold", 8)
    canvas.setFillColor(colors.HexColor(f"#{MUTED}"))
    canvas.drawCentredString(LETTER[0] / 2, LETTER[1] - 0.48 * inch, "LAB 02  |  BUFFER OVERFLOW TRONG ỨNG DỤNG WEB LOCAL")
    canvas.setFont("LabSans", 8.5)
    canvas.drawCentredString(LETTER[0] / 2, 0.45 * inch, f"Trang {doc.page}")
    canvas.restoreState()


def pdf_table(data: list[list[str]], widths: list[float], font: str, bold: str) -> Table:
    table = Table(data, colWidths=widths, repeatRows=1, hAlign="LEFT")
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#AAB4BE")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{PALE_BLUE}")),
        ("FONTNAME", (0, 0), (-1, 0), bold),
        ("FONTNAME", (0, 1), (-1, -1), font),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("LEADING", (0, 0), (-1, -1), 10.5),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    return table


def add_pdf_shot(story: list, number: int, shot: tuple[int, str, str, str, str, str], styles: dict, font: str, bold: str) -> None:
    _, filename, location, action, required, caption = shot
    path = SCREENSHOT_DIR / filename
    components = [Paragraph(filename, styles["h2"])]
    if path.exists() and path.stat().st_size:
        width, height = picture_size(path, 6.1 * inch, 5.6 * inch)
        components.append(Image(str(path), width=width, height=height, hAlign="CENTER"))
    else:
        data = [
            [Paragraph(f"<b>ẢNH THỦ CÔNG CẦN BỔ SUNG: {safe_text(filename)}</b>", styles["small"])],
            [Paragraph(f"URL hoặc lệnh: {safe_text(location)}", styles["small"])],
            [Paragraph(f"Thao tác: {safe_text(action)}", styles["small"])],
            [Paragraph(f"Nội dung bắt buộc: {safe_text(required)}", styles["small"])],
        ]
        table = Table(data, colWidths=[6.1 * inch], hAlign="LEFT")
        table.setStyle(TableStyle([
            ("BOX", (0, 0), (-1, -1), 0.8, colors.HexColor("#C94A4A")),
            ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#BCC5CE")),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{PALE_RED}")),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ("TOPPADDING", (0, 0), (-1, -1), 7),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ]))
        components.append(table)
    components.append(Paragraph(f"Hình {number}. {safe_text(caption)}", styles["caption"]))
    story.append(KeepTogether(components))


def add_pdf_special(story: list, chapter_number: int, observed: dict[str, str], styles: dict, font: str, bold: str, table_counter: list[int]) -> None:
    if chapter_number == 3:
        story.extend([
            Paragraph("Browser -> HTTP POST -> Flask -> subprocess.run -> C processor -> OS -> HTTP response", styles["code"]),
            Paragraph("Sơ đồ 1. Sequence diagram khái niệm của hệ thống.", styles["caption"]),
        ])
    elif chapter_number == 4:
        story.extend([
            Paragraph("int process_name(const char *user_input) {<br/>    char name[32];<br/>    strcpy(name, user_input);<br/>    return 0;<br/>}", styles["code"]),
            Paragraph("Đoạn mã 1. Mã nguồn tối thiểu chứa thao tác strcpy có chủ đích.", styles["caption"]),
        ])
    elif chapter_number == 5:
        story.append(Paragraph(safe_text(observed["normal"]), styles["body"]))
    elif chapter_number == 6:
        story.append(Paragraph(safe_text(observed["lengths"]), styles["body"]))
        table_counter[0] += 1
        story.extend([
            Paragraph(f"Bảng {table_counter[0]}. Phân biệt capacity, ASan và crash.", styles["caption"]),
            pdf_table([
                ["Mốc", "Ý nghĩa", "Cách xác nhận"],
                ["31 byte", "Capacity an toàn", "Mô hình name[32] + null"],
                ["32 byte", "Ít nhất 1 byte ghi vượt", "Mô hình copy tới null"],
                ["ASan đầu tiên", "Phát hiện instrumentation", "length_test.json thật"],
                ["Crash đầu tiên", "Dừng tiến trình", "exit/signal thật"],
            ], [1.05 * inch, 2.3 * inch, 2.85 * inch], font, bold),
        ])
    elif chapter_number == 7:
        story.append(Paragraph(safe_text(observed["gdb"]), styles["body"]))
    elif chapter_number == 11:
        table_counter[0] += 1
        story.extend([
            Paragraph(f"Bảng {table_counter[0]}. Profile build và nguyên tắc xác minh.", styles["caption"]),
            pdf_table([
                ["Profile", "Mục đích", "Trạng thái hardening"],
                ["vulnerable_debug", "GDB", "Xác minh bằng file/readelf"],
                ["vulnerable_asan", "Phát hiện lỗi", "Không đại diện production"],
                ["secure_length", "Bản vá length", "Xác minh binary thật"],
                ["secure_snprintf", "Bản vá snprintf", "Xác minh binary thật"],
                ["secure_hardened", "So sánh hardening", "Xác minh binary thật"],
            ], [1.55 * inch, 1.8 * inch, 2.85 * inch], font, bold),
        ])
    elif chapter_number == 12:
        table_counter[0] += 1
        story.extend([
            Paragraph(f"Bảng {table_counter[0]}. So sánh trước và sau vá.", styles["caption"]),
            pdf_table([
                ["Tiêu chí", "Vulnerable", "Secure length", "Secure snprintf"],
                ["Hàm copy", "strcpy", "memcpy sau validation", "snprintf"],
                ["Giới hạn", "Không", "Tối đa 31 byte", "sizeof(name)"],
                ["Input dài", "Ghi vượt", "Từ chối trước copy", "Từ chối truncate"],
                ["Hardening", "Không thay bản vá", "Lớp bổ sung", "Lớp bổ sung"],
            ], [1.1 * inch, 1.45 * inch, 1.85 * inch, 1.8 * inch], font, bold),
        ])
    elif chapter_number == 16:
        story.append(Paragraph(safe_text(observed["pytest"]), styles["body"]))


def build_pdf() -> None:
    font, bold, italic = register_pdf_fonts()
    styles = pdf_styles(font, bold, italic)
    frame = Frame(inch, 0.72 * inch, LETTER[0] - 2 * inch, LETTER[1] - 1.45 * inch, id="normal")
    template = ReportTemplate(str(PDF_PATH), pagesize=LETTER, leftMargin=inch, rightMargin=inch, topMargin=0.75 * inch, bottomMargin=0.72 * inch, title="Lab 2 - Buffer Overflow trong ứng dụng web local", author="Lê Minh - 21127645")
    template.addPageTemplates(PageTemplate(id="report", frames=frame, onPage=on_pdf_page))
    story: list = [
        Spacer(1, 1.25 * inch),
        Paragraph("BÁO CÁO THỰC HÀNH AN TOÀN ỨNG DỤNG WEB", styles["cover_kicker"]),
        Paragraph("LAB 2", styles["cover_title"]),
        Paragraph("BUFFER OVERFLOW TRONG<br/>ỨNG DỤNG WEB LOCAL", styles["cover_sub"]),
        Spacer(1, 0.25 * inch),
        Paragraph("Từ HTTP POST tới stack frame, AddressSanitizer và bản vá", styles["cover_meta"]),
        Spacer(1, 0.65 * inch),
        Paragraph("<b>MSSV:</b> 21127645", styles["cover_meta"]),
        Paragraph("<b>Họ tên:</b> Lê Minh", styles["cover_meta"]),
        Paragraph(f"<b>Ngày lập báo cáo:</b> {date.today().strftime('%d/%m/%Y')}", styles["cover_meta"]),
        PageBreak(),
        Paragraph("Mục lục", styles["h1"]),
    ]
    toc = TableOfContents()
    toc.levelStyles = [
        ParagraphStyle("TOC 1", fontName=bold, fontSize=10, leading=14, leftIndent=0, firstLineIndent=0, textColor=colors.HexColor(f"#{NAVY}"), spaceBefore=3),
        ParagraphStyle("TOC 2", fontName=font, fontSize=9, leading=12, leftIndent=18, firstLineIndent=0, textColor=colors.HexColor(f"#{MUTED}")),
    ]
    story.extend([toc, PageBreak()])
    observed = observed_summary()
    figure = 0
    table_counter = [0]
    for chapter_number, (title, sections) in enumerate(CHAPTERS, 1):
        story.append(Paragraph(safe_text(title), styles["h1"]))
        for section_title, paragraphs, bullets in sections:
            story.append(Paragraph(safe_text(section_title), styles["h2"]))
            for text in paragraphs:
                for part in text.split("\n"):
                    story.append(Paragraph(safe_text(part), styles["body"]))
            for bullet in bullets:
                story.append(Paragraph(safe_text(bullet), styles["bullet"], bulletText="•"))
        add_pdf_special(story, chapter_number, observed, styles, font, bold, table_counter)
        for shot in (item for item in SHOTS if item[0] == chapter_number):
            figure += 1
            add_pdf_shot(story, figure, shot, styles, font, bold)
    story.extend([
        PageBreak(),
        Paragraph("Phụ lục. Bằng chứng và câu lệnh", styles["h1"]),
        Paragraph("A. Mã nguồn và compiler flags", styles["h2"]),
        Paragraph("Mã nguồn nằm trong native/*.c và native/processor_common.h. Makefile là nguồn sự thật về flags; tên profile không thay cho kiểm tra binary.", styles["body"]),
        Paragraph("B. Request, response và trace", styles["h2"]),
        Paragraph("Bằng chứng thật nằm trong evidence/requests, evidence/responses và evidence/traces. File vắng nghĩa là chưa thu thập.", styles["body"]),
        Paragraph("C. ASan và GDB", styles["h2"]),
        Paragraph("ASan log nằm trong evidence/asan; GDB log thủ công nằm trong evidence/gdb. Không có log giả hoặc nội dung khai thác.", styles["body"]),
        Paragraph("D. Câu lệnh chạy", styles["h2"]),
    ])
    for command in ("make all", "python app.py", "python scripts/test_lengths.py", "gdb -q -x gdb/inspect_normal.gdb", "gdb -q -x gdb/inspect_overflow.gdb", "pytest", "python scripts/check_screenshots.py", "python scripts/generate_report.py"):
        story.append(Paragraph(safe_text(command), styles["bullet"], bulletText="•"))
    template.multiBuild(story)


def main() -> int:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    SCREENSHOT_DIR.mkdir(parents=True, exist_ok=True)
    missing = build_docx()
    build_pdf()
    print(f"Đã tạo DOCX: {DOCX_PATH}")
    print(f"Đã tạo PDF:  {PDF_PATH}")
    print(f"Ảnh còn thiếu ({len(missing)}): {', '.join(missing) if missing else 'không'}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
