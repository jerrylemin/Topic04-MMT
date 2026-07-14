import os
import subprocess
import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]


def run(*args):
    return subprocess.run([sys.executable, *args], cwd=ROOT, text=True, capture_output=True,
                          encoding="utf-8", env={**os.environ, "PYTHONIOENCODING": "utf-8"})


def test_guide_and_manifest_are_synchronized():
    from screenshot_manifest import SCREENSHOTS
    guide = (ROOT / "HUONG_DAN_CHUP_ANH.md").read_text(encoding="utf-8")
    assert len(SCREENSHOTS) == 9
    assert all(item["filename"] in guide for item in SCREENSHOTS)
    assert guide.count("**Tên file:**") == 9
    assert guide.count("**Caption dùng trong báo cáo:**") == 9


def test_checker_lists_required_in_order():
    result = run("scripts/check_screenshots.py", "--list-required")
    assert result.returncode == 0
    assert result.stdout.index("01_normal_input.png") < result.stdout.index("09_tests_reports.png")


def test_report_has_nine_sections_and_detailed_placeholders():
    result = run("scripts/generate_report.py")
    assert result.returncode == 0, result.stderr
    doc = Document(ROOT / "report/21127645_LeMinh_Lab02_BufferOverflow.docx")
    text = "\n".join(p.text for p in doc.paragraphs)
    cells = "\n".join(c.text for t in doc.tables for row in t.rows for c in row.cells)
    assert "1. Mục tiêu và môi trường thực hành" in text and "9. Kết luận" in text
    assert cells.count("Chèn ảnh tại vị trí này.") == 9
    assert "ẢNH 01/09" in cells and "01_normal_input.png" in cells
    assert (ROOT / "report/21127645_LeMinh_Lab02_BufferOverflow.pdf").stat().st_size > 10_000
