import subprocess
import zipfile
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SCRIPT = ROOT / "scripts/generate_report.py"
DOCX = ROOT / "report/21127645_LeMinh_Lab04_CSRF.docx"
PDF = ROOT / "report/21127645_LeMinh_Lab04_CSRF.pdf"


def test_generator_uses_manifest_and_real_test_logs():
    source = SCRIPT.read_text(encoding="utf-8")
    for required in ("screenshot_manifest", "pytest.txt", "coverage.txt", "Chèn ảnh tại vị trí này.", "image_size"):
        assert required in source


def test_generated_docx_and_pdf_are_real_complete_artifacts():
    assert DOCX.stat().st_size > 20_000 and PDF.stat().st_size > 10_000
    with zipfile.ZipFile(DOCX) as archive:
        assert archive.testzip() is None
        document = archive.read("word/document.xml")
    for text in (b"21127645", "Mục tiêu và môi trường thực hành".encode("utf-8"), b"validate_csrf_token"):
        assert text in document
    assert PDF.read_bytes()[:5] == b"%PDF-"


def test_report_has_nine_sections_and_seven_positioned_placeholders():
    doc = Document(DOCX)
    text = "\n".join(p.text for p in doc.paragraphs)
    for heading in ("1. Mục tiêu và môi trường thực hành", "2. Kịch bản và các bước thực hiện",
                    "3. Nguyên nhân kỹ thuật", "4. Kết quả và bằng chứng", "5. Mức độ ảnh hưởng",
                    "6. Bản vá và cách phòng chống", "7. Trả lời các câu hỏi báo cáo",
                    "8. Kết quả kiểm thử", "9. Kết luận"):
        assert heading in text
    placeholders = [cell.text for table in doc.tables for row in table.rows for cell in row.cells if "Chèn ảnh tại vị trí này." in cell.text]
    assert len(placeholders) == 7
    assert all(f"ẢNH {index:02d}/07" in placeholders[index - 1] for index in range(1, 8))
