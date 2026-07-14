"""Build the Lab04 DOCX/PDF report from verified source and evidence files."""

import ast
import json
import re
import shutil
import subprocess
import sys
from html import escape
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "report"
DOCX = REPORT / "21127645_LeMinh_Lab04_CSRF.docx"
PDF = REPORT / "21127645_LeMinh_Lab04_CSRF.pdf"
FLOW_FILES = [
    ("Login", "login_victim.json"),
    ("Vulnerable email", "vulnerable_email_change.json"),
    ("Secure missing token", "secure_email_missing_token.json"),
    ("Secure invalid token", "secure_email_invalid_token.json"),
    ("Secure origin denied", "secure_email_origin_denied.json"),
    ("Secure success", "secure_email_success.json"),
    ("Logout denied", "logout_csrf_denied.json"),
    ("Logout success", "logout_success.json"),
    ("Reset denied", "reset_csrf_denied.json"),
    ("Reset success", "reset_success.json"),
]
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x55, 0x55, 0x55)


def _required() -> dict:
    paths = {
        "audit": ROOT / "evidence/audit/audit_logs.json",
        "state": ROOT / "evidence/state/state_transitions.json",
        "pytest": ROOT / "evidence/logs/pytest.txt",
        "coverage": ROOT / "evidence/logs/coverage.txt",
        "smoke": ROOT / "evidence/logs/runtime_smoke_test.txt",
    }
    paths.update({f"trace:{name}": ROOT / "evidence/traces" / filename for name, filename in FLOW_FILES})
    missing = [str(path.relative_to(ROOT)) for path in paths.values() if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing real evidence: " + ", ".join(missing))
    return paths


def _read_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


def _font(run, size=10.5, bold=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = color


def _configure(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name, normal.font.size = "Calibri", Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name, style.font.size, style.font.color.rgb = "Calibri", Pt(size), color
        style.paragraph_format.space_before, style.paragraph_format.space_after = Pt(before), Pt(after)
    doc.core_properties.title = "Lab04 - Cross-Site Request Forgery (CSRF)"
    doc.core_properties.subject = "Báo cáo thực hành bảo mật ứng dụng web"
    doc.core_properties.author = "21127645 - Lê Minh"
    doc.core_properties.keywords = "CSRF, Flask, Synchronizer Token, Origin, Referer, SQLite"
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _font(header.add_run("LAB04 · CSRF · 21127645"), 8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(footer.add_run("Trang "), 8.5, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def _shade(cell, fill="F2F4F7"):
    props = cell._tc.get_or_add_tcPr()
    shading = props.find(qn("w:shd")) or OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    if shading.getparent() is None:
        props.append(shading)


def _table(doc, headers, rows, widths, font_size=8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    table_width = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w = tbl_w if tbl_w is not None else OxmlElement("w:tblW")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(table_width * 1440)))
    if tbl_w.getparent() is None:
        tbl_pr.append(tbl_w)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    for index, (text, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[index]
        cell.width = Inches(width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _shade(cell)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        _font(paragraph.add_run(str(text)), font_size, bold=True)
    for values in rows:
        cells = table.add_row().cells
        for index, (value, width) in enumerate(zip(values, widths)):
            cells[index].width = Inches(width)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            _font(paragraph.add_run(str(value)), font_size)
    return table


def _title(doc, kicker, title, subtitle=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(115)
    p.paragraph_format.space_after = Pt(14)
    _font(p.add_run(kicker.upper()), 11, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    _font(p.add_run(title), 28, bold=True, color=DARK_BLUE)
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(70)
        _font(p.add_run(subtitle), 14, color=MUTED)


def _page(doc, title, paragraphs=(), bullets=(), table=None):
    doc.add_heading(title, level=1)
    for text in paragraphs:
        doc.add_paragraph(text)
    for text in bullets:
        doc.add_paragraph(text, style="List Bullet")
    if table:
        _table(doc, *table)
    doc.add_page_break()


def _source_function(filename, function):
    path = ROOT / filename
    source = path.read_text(encoding="utf-8")
    node = next(item for item in ast.walk(ast.parse(source)) if isinstance(item, ast.FunctionDef) and item.name == function)
    return node.lineno, node.end_lineno, "\n".join(source.splitlines()[node.lineno - 1:node.end_lineno])


def _compact(value, limit=55):
    text = json.dumps(value, ensure_ascii=False, separators=(",", ":")) if isinstance(value, (dict, list)) else str(value)
    text = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", text)
    return text if len(text) <= limit else text[:limit - 1] + "…"


def build_docx() -> Path:
    paths = _required()
    traces = [(name, _read_json(paths[f"trace:{name}"])) for name, _filename in FLOW_FILES]
    audits = _read_json(paths["audit"])
    states = _read_json(paths["state"])
    pytest_log = paths["pytest"].read_text(encoding="utf-8", errors="replace")
    coverage_log = paths["coverage"].read_text(encoding="utf-8", errors="replace")
    smoke_log = paths["smoke"].read_text(encoding="utf-8", errors="replace")

    doc = Document()
    _configure(doc)
    _title(doc, "Báo cáo thực hành bảo mật ứng dụng web", "LAB04 · CROSS-SITE REQUEST FORGERY", "Synchronizer Token · Origin/Referer · Audit & Trace")
    for line in ("Sinh viên: Lê Minh", "MSSV: 21127645", "Môi trường: Flask + SQLite, chỉ chạy loopback", "Ngày hoàn thiện: 15/07/2026"):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; _font(p.add_run(line), 11, bold="MSSV" in line)
    doc.add_page_break()

    toc = [(i + 1, title) for i, title in enumerate([
        "Thông tin, mục tiêu và phạm vi", "Kiến trúc, database, session và cookie", "Nền tảng CSRF",
        "Luồng vulnerable và secure", "Token, Origin và Referer", "SOP, SameSite và CORS",
        "POST-only, logout và reset", "Audit, trace và inspectors", "Security Controls và Code Comparison",
        "Kết quả kiểm thử, coverage và smoke test", "Ảnh hưởng, phòng chống, giới hạn và kết luận",
        "Câu hỏi báo cáo và phụ lục evidence", "Phụ lục 10 bảng trace chi tiết",
    ])]
    doc.add_heading("Mục lục", level=1)
    _table(doc, ["STT", "Nội dung"], toc, [0.6, 5.9], 9)
    doc.add_heading("Danh mục bảng", level=2)
    doc.add_paragraph("Bảng kiến trúc và schema; Origin Matrix; audit event; state transition; security control; code comparison; kết quả test; 10 bảng trace theo từng flow.")
    doc.add_page_break()

    _page(doc, "1. Thông tin, mục tiêu và phạm vi", [
        "Lab04 là ứng dụng học tập local gồm Victim Application tại 127.0.0.1:5004 và Cross-Origin Form Demo tại 127.0.0.1:9004 hoặc localhost:9004. Mục tiêu là chỉ ra vì sao request giả mạo có thể mang theo session cookie, sau đó so sánh luồng thiếu token với luồng được bảo vệ bằng Synchronizer Token Pattern.",
        "Phạm vi an toàn được cố định trong mã nguồn: không nhận host, URL, port hay route tùy ý; không kết nối Internet; không dùng browser automation; không dùng fetch, XMLHttpRequest, iframe hoặc document.cookie; không auto-submit. Form minh họa chỉ gửi khi người dùng chủ động bấm và xác nhận.",
        "Báo cáo này được tạo lại từ trace JSON, request/response TXT, audit log, state transition, pytest, coverage, runtime smoke test và source code thật. Không có ảnh giả, placeholder ảnh hoặc chương ảnh trống.",
    ], ["Tài khoản demo victim / Victim123! chỉ dùng cho dữ liệu local.", "Mật khẩu lưu dưới dạng hash Werkzeug; token và cookie luôn được che trong evidence.", "Hai ứng dụng bind loopback và không bật CORS wildcard."])

    _page(doc, "2. Kiến trúc, database, session và cookie", [
        "Victim Application xử lý login, session, validation, SQLite mutation, audit và trace. Demo Page chỉ render ba form email cố định: vulnerable, secure thiếu token và secure token giả. Hai ứng dụng không chia sẻ secret, session hay database.",
        "SQLite có users, demo_transfers, audit_logs, state_history và trace_records. Mọi SQL mutation dùng placeholder và transaction. State Inspector liên kết state_history với trace_id, vì vậy email trước/sau và quyết định update được kiểm tra từ dữ liệu database thật.",
        "Flask session cookie được cấu hình HttpOnly, SameSite=Lax, path=/ và Secure=False cho HTTP loopback. HttpOnly hạn chế JavaScript đọc cookie nhưng không ngăn browser tự gắn cookie vào request; SameSite là lớp bổ sung và không thay thế token.",
        "Data flow: Browser → Flask Router → Session Authentication → Origin/Referer → CSRF → Input Validation → SQLite → Audit/Trace → HTTP Response.",
    ], table=(["Thành phần", "Nguồn dữ liệu", "Vai trò"], [["Victim", "victim_app.py", "Validation và mutation"], ["Demo Page", "attacker_app.py", "Form local cố định"], ["SQLite", "lab04.sqlite3", "State, audit, trace"]], [1.2, 2.0, 3.3]))

    _page(doc, "3. Nền tảng CSRF", [
        "CSRF là việc lợi dụng browser của người dùng đã đăng nhập để gửi request thay đổi trạng thái mà ứng dụng nhầm là chủ ý của người dùng. Người tạo form không cần biết mật khẩu: browser quản lý cookie và tự đính kèm cookie khi chính sách cho phép.",
        "Điều kiện điển hình gồm session còn hiệu lực, route thay đổi trạng thái, dữ liệu request có thể đoán hoặc cố định, và server không yêu cầu bằng chứng không thể giả mạo như token gắn với session. SOP có thể ngăn script đọc response nhưng không ngăn HTML form gửi request, nên không đọc được response không đồng nghĩa state không đổi.",
        "CSRF khác XSS: CSRF ép browser gửi hành động bằng credential sẵn có, còn XSS chạy script trong trusted origin. XSS nghiêm trọng hơn đối với cơ chế CSRF vì script cùng origin có thể đọc token trong DOM và gửi request hợp lệ. Do đó CSP, escaping và phòng XSS vẫn là lớp bảo vệ thiết yếu.",
    ], ["Không dùng GET cho state change vì link, prefetch, cache và crawler có thể kích hoạt.", "Token phải kiểm tra ở server vì client không phải trust boundary.", "Token không đặt trong URL để tránh history, log và Referer leak."])

    _page(doc, "4. Luồng vulnerable và secure", [
        "POST /vulnerable/change-email yêu cầu session và validate email nhưng cố ý không kiểm tra Origin/Referer hoặc CSRF token. Khi form local khác origin gửi email demo_changed@lab.local và cookie hiện diện, SQLite được cập nhật, audit vulnerable_email_changed được ghi và trace mô tả nguyên nhân session-cookie-only.",
        "POST /secure/change-email kiểm tra theo thứ tự: session, Origin hoặc Referer exact match, token hiện diện, hmac.compare_digest, email hợp lệ, rồi mới UPDATE. Sau thành công token được rotate, audit secure_email_changed/csrf_token_valid/csrf_token_rotated được ghi và response gắn với trace.",
        "Token thiếu hoặc sai trả 403 và không đổi email. Origin sai được từ chối trước token. Các denial vẫn tạo audit và trace để cho thấy database update bị skip. Đây là khác biệt quan sát được giữa hai flow với cùng loại dữ liệu đầu vào.",
    ], table=(["Flow", "Origin", "Token", "Kết quả"], [["Vulnerable", "Không kiểm tra", "Không yêu cầu", "Email có thể đổi"], ["Secure", "Exact match", "Session-bound", "Chỉ đổi khi mọi check đạt"]], [1.2, 1.7, 1.5, 2.1]))

    _page(doc, "5. Token, Origin và Referer", [
        "Token được sinh bằng secrets.token_urlsafe(32), lưu cùng session, khác giữa hai session, tạo mới sau login và rotate sau secure state change. validate_csrf_token trả cấu trúc present, valid, status và reason; so sánh dùng hmac.compare_digest. Inspector chỉ nhận token đã mask ở server.",
        "Origin parser dùng urllib.parse.urlsplit và chuẩn hóa scheme, hostname, effective port. Allowlist chỉ gồm http://127.0.0.1:5004 và http://localhost:5004. Không dùng substring, startswith hoặc endswith. Nếu Origin có mặt, Origin luôn thắng; Referer chỉ là fallback khi Origin vắng; cả hai thiếu thì từ chối.",
        "Origin và Referer là lớp bổ sung vì header có thể bị thiếu bởi chính sách trình duyệt hoặc proxy. Token vẫn là lớp chính vì nó chứng minh request biết giá trị ngẫu nhiên gắn với session.",
    ], table=(["Nguồn", "Parsed", "Exact match", "Decision"], [["127.0.0.1:5004", "http / 127.0.0.1 / 5004", "Có", "Allowed"], ["127.0.0.1:9004", "http / 127.0.0.1 / 9004", "Không", "Denied"], ["Missing", "-", "Không", "Denied"]], [1.5, 2.4, 1.1, 1.5]))

    _page(doc, "6. SOP, SameSite và CORS", [
        "127.0.0.1:9004 → 127.0.0.1:5004 là cross-origin do khác port nhưng same-site do cùng scheme và host. localhost:9004 → 127.0.0.1:5004 vừa cross-origin vừa cross-site. Báo cáo phân biệt Expected theo policy, Observed từ request thật và Not observed khi chưa dùng browser DevTools.",
        "Same-Origin Policy kiểm soát script đọc dữ liệu cross-origin; nó không cấm form HTML gửi request. Demo chỉ giải thích lý thuyết, không dùng iframe/contentDocument và không tuyên bố tự động chứng minh SOP.",
        "CORS điều khiển quyền script đọc response/API; form POST chuẩn không cần CORS. Vì vậy tắt CORS hoặc không bật wildcard không phải bản vá CSRF chính. SameSite phụ thuộc browser, scheme, host, method và navigation context; token server-side vẫn là lớp chính.",
    ], table=(["Nguồn → Victim", "Origin", "Site", "Nhãn"], [["127.0.0.1:9004", "Khác", "Cùng", "Expected; browser chưa tự động xác minh"], ["localhost:9004", "Khác", "Khác", "Expected; browser chưa tự động xác minh"]], [1.8, 1.0, 1.0, 2.7]))

    _page(doc, "7. POST-only, logout và reset", [
        "Mọi secure state change dùng POST, session hợp lệ, Origin/Referer exact validation, CSRF token, input validation, audit và trace trước mutation. GET /secure/change-email chỉ render form và không đổi database. Route state-changing GET cũ đã bị loại bỏ.",
        "POST /logout từ chối token thiếu/sai hoặc Origin không hợp lệ với HTTP 403 và giữ session. Thành công ghi logout_success trước khi session.clear. POST /reset-lab cũng từ chối trước mutation; thành công reset user/balance, giữ audit/trace/state evidence và sau đó logout.",
        "Logout cần CSRF protection vì forced logout gây mất phiên và phá workflow. Reset cần bảo vệ mạnh hơn vì thay đổi toàn bộ trạng thái lab. HttpOnly không giúp trong hai trường hợp này vì browser vẫn có thể gửi cookie.",
    ], table=(["Route", "Denied event", "Success event", "State rule"], [["/logout", "logout_csrf_denied", "logout_success", "Giữ session khi denied"], ["/reset-lab", "lab_reset_csrf_denied", "lab_reset", "Giữ DB khi denied"]], [1.2, 1.8, 1.5, 2.0]))

    _page(doc, "8. Audit, trace và inspectors", [
        "Audit log lưu timestamp, user_id, username, action, route, mode, Origin, Referer, CSRF status, cookie presence, decision, reason, state before/after và trace_id. Password, full cookie, full token và secret key không được ghi. Trang audit hỗ trợ filter theo action, decision, mode, username và trace ID.",
        "Mỗi trace có request metadata thật và 16 bước: Browser UI, HTTP Request, Flask Router, Session Authentication, Origin Validation, CSRF Validation, Input Validation, SQLite Read/Write, Audit Logging, HTTP Response và Final Result. Mỗi step hiển thị timestamp, layer, title, technique, input, output, code reference, security meaning và status.",
        "Request Inspector đọc URL, query, content metadata và form đã redaction; Cookie Inspector đọc app.config và request.cookies; Token Inspector mask ở server; Origin Inspector dùng validation result; State Inspector đọc state_history; Presentation Mode chỉ trình bày trace đã có và không gửi request.",
    ], table=(["Evidence", "Số bản ghi", "Nguồn"], [["Audit", len(audits), "SQLite audit_logs"], ["State", len(states), "SQLite state_history"], ["Trace flow", len(traces), "SQLite trace_records / JSON export"]], [1.4, 1.1, 4.0]))

    snippets = []
    for filename, function in (("victim_app.py", "vulnerable_change_email"), ("victim_app.py", "secure_change_email"), ("csrf_service.py", "validate_csrf_token"), ("origin_service.py", "validate_origin_or_referer"), ("victim_app.py", "logout"), ("victim_app.py", "reset_lab")):
        start, end, code = _source_function(filename, function)
        snippets.append([filename, function, start, end, _compact(code, 110)])
    _page(doc, "9. Security Controls và Code Comparison", [
        "Security Control Panel phản ánh runtime config và code: session authentication, token, rotation, Origin, Referer fallback, SameSite, HttpOnly, Secure cookie, POST-only, input validation, parameterized SQL, audit, CSP, CORS policy và request size limit. Mỗi control nêu nguồn, route, rủi ro giảm được và giới hạn.",
        "Code Comparison dùng AST đọc source đang chạy để lấy tên file, function, line start/end và snippet. Bảng dưới đây được tạo lúc chạy generator; không phải pseudocode ghi cứng. Điều này bảo đảm nội dung báo cáo khớp với implementation cuối.",
    ], table=(["File", "Function", "Start", "End", "Source excerpt"], snippets, [1.1, 1.35, 0.55, 0.55, 2.95], 6.5))

    _page(doc, "10. Kết quả kiểm thử, coverage và runtime smoke", [
        "Pytest log và coverage log dưới đây được đọc từ evidence cuối. Generator không tự tuyên bố test đạt khi log chưa tồn tại. Bộ test bao phủ route, authentication, session, CSRF, rotation, Origin/Referer, vulnerable/secure flow, logout/reset, SQLite, audit, trace, inspectors, controls, code extraction, headers, CORS, presentation, evidence, report, runtime restriction và cleanup.",
        "Runtime smoke test khởi động hai ứng dụng loopback và chạy healthcheck, login, vulnerable email, reset hợp lệ, secure missing/invalid/origin-denied/success, logout missing và logout success. Script không kiểm tra SameSite/SOP của browser.",
        "Pytest summary: " + _compact(pytest_log.replace("\n", " "), 420),
        "Coverage summary: " + _compact(coverage_log.replace("\n", " "), 420),
        "Smoke summary: " + _compact(smoke_log.replace("\n", " "), 420),
    ])

    _page(doc, "11. Ảnh hưởng, phòng chống, giới hạn và kết luận", [
        "CSRF có thể thay đổi dữ liệu bằng quyền của nạn nhân, làm sai lệch tài khoản và gây mất tính toàn vẹn. Mức ảnh hưởng phụ thuộc route: đổi email có thể hỗ trợ chiếm quyền ở hệ thống thật; logout gây gián đoạn; reset phá trạng thái. Lab chỉ dùng dữ liệu giả lập và không cung cấp công cụ mục tiêu tùy ý.",
        "Biện pháp chính là token ngẫu nhiên gắn session, validate ở server trước mutation và rotate sau thành công. Các lớp bổ sung gồm exact Origin/Referer, SameSite phù hợp, POST-only, input validation, re-authentication cho password/transfer, CSP, audit và trace. XSS phải được phòng riêng vì có thể làm suy yếu token.",
        "Giới hạn: requests/test client không mô phỏng quyết định SameSite hay khả năng đọc response của browser; phần này chỉ được gắn nhãn Expected hoặc lý thuyết. Cookie Secure tắt do lab dùng HTTP loopback. Không có ảnh chụp thủ công theo phạm vi nhiệm vụ.",
        "Kết luận: vulnerable flow chứng minh thiếu token là nguyên nhân cốt lõi; secure flow bảo vệ state bằng token + Origin/Referer + validation trước SQL. Evidence, tests, coverage và smoke test tạo chuỗi chứng minh có thể tái lập.",
    ])

    _page(doc, "12. Câu hỏi báo cáo và phụ lục evidence", [
        "Browser tự gửi cookie vì cookie jar và matching policy thuộc browser. CSRF không cần biết mật khẩu vì session đã xác thực. Phía tạo form cross-origin thường không đọc được response do SOP, nhưng request vẫn có thể đổi state.",
        "Token phải kiểm tra server-side; URL không phù hợp vì leak qua history/log/Referer. SameSite không thay token do còn same-site cross-origin và khác biệt context. HttpOnly chỉ chặn JavaScript đọc cookie, không chặn cookie được gửi.",
        "SameSite phân loại site; SOP kiểm soát script đọc theo origin. CORS không phải bản vá chính vì form submit không cần CORS. Origin/Referer là lớp bổ sung vì có thể thiếu. Logout/reset cần token vì đều thay đổi trạng thái. XSS có thể đọc token cùng origin, nên phòng XSS là bắt buộc.",
        "Evidence appendix gồm 10 trace JSON, 7 request TXT, 7 response TXT, audit_logs.json, state_transitions.json, pytest.txt, coverage.txt, runtime_smoke_test.txt, security_review.txt và submission_cleanup.txt. Tất cả nằm dưới evidence/ và không chứa secret đầy đủ.",
    ])

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = Inches(11), Inches(8.5)
    section.top_margin = section.bottom_margin = Inches(0.55)
    section.left_margin = section.right_margin = Inches(0.5)
    section.header_distance = section.footer_distance = Inches(0.3)
    for flow_index, (name, trace) in enumerate(traces):
        doc.add_heading(f"Phụ lục Trace {flow_index + 1}: {name}", level=1)
        doc.add_paragraph(f"Trace ID: {trace['trace_id']} · HTTP {trace['http_status']} · Decision: {trace['final_result']} · Origin: {trace['origin_decision']} · CSRF: {trace['csrf_token_status']}")
        rows = []
        for step in trace["steps"]:
            ref = step.get("code_reference") or {}
            reference = f"{ref.get('file','-')}:{ref.get('function','-')}:{ref.get('line','-')}" if isinstance(ref, dict) else str(ref)
            rows.append([
                step["step_number"], _compact(step["layer"], 20), _compact(step["title"], 28),
                _compact(step.get("technique", ""), 34), _compact(step.get("input_data", {}), 42),
                _compact(step.get("output_data", {}), 42), _compact(reference, 35), step.get("status", ""),
                _compact(step.get("security_meaning", ""), 48),
            ])
        _table(doc, ["STT", "Layer", "Action", "Technique", "Input", "Output", "Code reference", "Status", "Security meaning"], rows,
               [0.35, 0.85, 1.1, 1.35, 1.25, 1.25, 1.15, 0.65, 1.5], 5.5)
        if flow_index < len(traces) - 1:
            doc.add_page_break()

    REPORT.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX)
    if not DOCX.exists() or DOCX.stat().st_size < 20_000:
        raise RuntimeError("DOCX generation failed or output is unexpectedly small.")
    return DOCX


def build_pdf() -> Path:
    if PDF.exists():
        PDF.unlink()
    soffice = shutil.which("soffice") or r"C:\Program Files\LibreOffice\program\soffice.exe"
    if not Path(soffice).exists():
        return _build_pdf_fallback()
    result = subprocess.run([soffice, "--headless", "--convert-to", "pdf", "--outdir", str(REPORT), str(DOCX)], capture_output=True, text=True, timeout=180)
    if result.returncode or not PDF.exists() or PDF.stat().st_size < 10_000:
        raise RuntimeError(f"PDF conversion failed: {result.stderr or result.stdout}")
    return PDF


def _build_pdf_fallback() -> Path:
    """Create a readable, evidence-backed PDF when LibreOffice is unavailable."""
    try:
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.units import inch
        from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer
    except ImportError as exc:
        raise RuntimeError(
            "PDF generation needs LibreOffice or the already-supported ReportLab fallback."
        ) from exc

    source = Document(DOCX)
    styles = getSampleStyleSheet()
    font_candidates = [
        Path(r"C:\Windows\Fonts\arial.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ]
    font_path = next((path for path in font_candidates if path.exists()), None)
    if font_path:
        pdfmetrics.registerFont(TTFont("ReportUnicode", str(font_path)))
        for style in styles.byName.values():
            style.fontName = "ReportUnicode"
    styles.add(ParagraphStyle(name="ReportTitle", parent=styles["Title"], alignment=TA_CENTER))
    story = []

    # Read both paragraphs and tables in document order.  Table rows become
    # textual evidence lines so the fallback never substitutes placeholders.
    for child in source.element.body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        text_content = " ".join(
            text.strip() for text in child.itertext() if text and text.strip()
        )
        if not text_content:
            continue
        style = styles["BodyText"]
        if tag == "p":
            paragraph_style = child.xpath("./w:pPr/w:pStyle/@w:val")
            if paragraph_style and str(paragraph_style[0]).startswith("Heading"):
                level = str(paragraph_style[0]).replace("Heading", "") or "1"
                style = styles.get(f"Heading{min(int(level), 3)}", styles["Heading1"])
        else:
            style = styles["Code"]
        for offset in range(0, len(text_content), 1800):
            story.append(Paragraph(escape(text_content[offset : offset + 1800]), style))
            story.append(Spacer(1, 0.08 * inch))
        if tag == "sectPr":
            story.append(PageBreak())

    SimpleDocTemplate(
        str(PDF),
        pagesize=letter,
        rightMargin=0.55 * inch,
        leftMargin=0.55 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
        title="Lab04 - CSRF",
        author="21127645 - Lê Minh",
    ).build(story)
    if not PDF.exists() or PDF.stat().st_size < 10_000:
        raise RuntimeError("ReportLab fallback did not produce a complete PDF.")
    return PDF


def main() -> int:
    build_docx()
    build_pdf()
    print(f"DOCX: {DOCX}")
    print(f"PDF: {PDF}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
