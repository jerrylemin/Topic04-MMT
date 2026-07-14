"""Sinh báo cáo Lab01 ngắn gọn; không tạo ảnh hay kết quả kiểm thử giả."""
from pathlib import Path
import sys

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
from screenshot_manifest import SCREENSHOTS

from PIL import Image as PILImage, UnidentifiedImageError
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

SHOT_DIR = ROOT / "evidence" / "screenshots"
REPORT_DIR = ROOT / "report"
DOCX_PATH = REPORT_DIR / "21127645_LeMinh_Lab01_XSS.docx"
PDF_PATH = REPORT_DIR / "21127645_LeMinh_Lab01_XSS.pdf"

SECTIONS = [
    ("1. Mục tiêu và môi trường thực hành", [
        "Mục tiêu là phân biệt Reflected, Stored và DOM-based XSS; theo dõi dữ liệu không tin cậy tới sink; kiểm chứng escaping, sanitization, textContent, CSP và cookie flags.",
        "Môi trường: Flask/SQLite và Chrome hoặc Firefox tại 127.0.0.1:5000. Payload chỉ hiển thị alert kiểm thử trong lab local; không kết nối website thật.",
    ], []),
    ("2. Kịch bản và các bước thực hiện", [
        "Reflected: gửi cùng q vào bản vulnerable và secure. Stored: reset database, POST bình luận, reload rồi mở bản secure. DOM: đặt cùng fragment vào hai trang và quan sát DOM. Cuối cùng đối chiếu code, CSP/cookie và kiểm thử.",
    ], [0, 2, 3, 5]),
    ("3. Nguyên nhân kỹ thuật", [
        "Reflected XSS: request.args['q'] được bọc Markup và chèn vào ngữ cảnh HTML nên ký tự đặc biệt không được encode.",
        "Stored XSS: comments.body được lưu trong SQLite rồi đánh dấu an toàn khi render. Lưu bằng SQL tham số chỉ chống SQL Injection, không chống XSS.",
        "DOM-based XSS: JavaScript đọc location.hash và gán vào innerHTML; lỗi xảy ra tại trình duyệt vì fragment không đi trong HTTP request.",
        "Vị trí dữ liệu Reflected là vùng kết quả HTML; bản vulnerable không escape <, >, dấu nháy. Người mở URL payload có thể chạy script trong origin của ứng dụng. Stored nguy hiểm hơn vì tồn tại và có thể tác động mọi người xem; cookie thiếu HttpOnly có thể bị script đọc.",
    ], []),
    ("4. Kết quả và bằng chứng", [
        "Các ảnh dưới đây đặt cạnh đúng kịch bản. Khi PNG hợp lệ tồn tại, generator thay khung hướng dẫn bằng ảnh thật và giữ caption.",
    ], [1, 4, 6, 7, 8]),
    ("5. Mức độ ảnh hưởng", [
        "Reflected cần nạn nhân mở URL; DOM cần fragment/nguồn client bị kiểm soát; Stored có phạm vi lớn nhất vì payload tồn tại và chạy với nhiều lượt xem. Hậu quả gồm thay đổi DOM, hành động dưới phiên người dùng và rò rỉ dữ liệu mà script truy cập được. HttpOnly giảm rủi ro đọc cookie nhưng không loại bỏ XSS.",
    ], []),
    ("6. Bản vá và cách phòng chống", [
        "Vulnerable: Markup(q), Markup(row['body']) và output.innerHTML = value. Secure: để Jinja autoescape q; sanitize nội dung HTML bằng allowlist rồi vẫn encode đúng ngữ cảnh; dùng output.textContent = value.",
        "Biện pháp bổ sung: output encoding theo ngữ cảnh HTML/attribute/URL/JavaScript; sanitization bằng thư viện tin cậy; validate input ở server; không tin URL/form/cookie/localStorage; CSP chặt; cookie HttpOnly, Secure trên HTTPS và SameSite. CSP là defense in depth, không thay việc sửa code.",
    ], []),
    ("7. Trả lời các câu hỏi báo cáo trong BaiTapTopic04.docx", [
        "So sánh: Reflected lấy payload từ request và phản hồi ngay; Stored lấy từ kho lưu trữ và lặp lại cho người xem; DOM-based do JavaScript client chuyển source như location.hash tới sink DOM, có thể không qua server.",
        "Validate input chưa đủ vì dữ liệu hợp lệ ở một ngữ cảnh vẫn nguy hiểm ở ngữ cảnh khác, có nhiều cách biểu diễn/encode và dữ liệu có thể đến từ nguồn ngoài form. Cần output encoding vì nó làm ký tự dữ liệu không còn được parser hiểu là cú pháp thực thi.",
        "CSP không thay sửa code; nó chỉ hạn chế nguồn/thực thi khi lỗi còn tồn tại. Vá từng lỗi: Reflected dùng autoescape/encoding; Stored sanitize allowlist và escape khi render; DOM thay innerHTML bằng textContent hoặc API DOM an toàn.",
    ], []),
    ("8. Kết quả kiểm thử", [], [9]),
    ("9. Kết luận", [
        "Ba lỗi cùng bắt đầu từ dữ liệu không tin cậy nhưng khác nơi lưu và nơi thực thi. Bản vá hiệu quả phải loại sink nguy hiểm/encode đúng ngữ cảnh; CSP và cookie flags chỉ tăng chiều sâu phòng thủ.",
    ], []),
]


def valid_png(path: Path) -> bool:
    try:
        if not path.is_file() or path.stat().st_size == 0:
            return False
        with PILImage.open(path) as image:
            return image.format == "PNG" and image.width >= 1 and image.height >= 1
    except (OSError, UnidentifiedImageError):
        return False


def pytest_summary() -> str:
    path = ROOT / "evidence" / "logs" / "pytest.txt"
    if not path.exists() or not path.stat().st_size:
        return "Chưa có evidence/logs/pytest.txt; báo cáo không tuyên bố kiểm thử đạt."
    lines = [line.strip() for line in path.read_text(encoding="utf-8", errors="replace").splitlines() if line.strip()]
    return "Log pytest thật (dòng cuối): " + (lines[-1] if lines else "log rỗng")


def shade(cell, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def configure_docx(doc: Document) -> None:
    section = doc.sections[0]
    section.page_height, section.page_width = Cm(29.7), Cm(21)
    section.top_margin = section.bottom_margin = Cm(1.7)
    section.left_margin = section.right_margin = Cm(1.8)
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Caption"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.styles["Normal"].paragraph_format.space_after = Pt(5)
    doc.styles["Heading 1"].font.size = Pt(15)
    doc.styles["Heading 1"].font.color.rgb = RGBColor(23, 74, 132)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(110)
    r = p.add_run("BÁO CÁO THỰC HÀNH LAB 01")
    r.bold = True; r.font.name = "Arial"; r.font.size = Pt(25); r.font.color.rgb = RGBColor(23, 74, 132)
    for text in ("CROSS-SITE SCRIPTING", "Reflected XSS - Stored XSS - DOM-based XSS", "MSSV: 21127645", "Họ tên: Lê Minh"):
        p = doc.add_paragraph(text); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def picture_inches(path: Path, max_w: float = 6.25, max_h: float = 4.3) -> tuple[float, float]:
    with PILImage.open(path) as image:
        scale = min(max_w / image.width, max_h / image.height)
        return image.width * scale, image.height * scale


def placeholder_lines(number: int, shot: dict) -> list[str]:
    action = f"Trạng thái: {shot['initial']} Thao tác: nhập {shot['data']}; {shot['button']}"
    return [
        f"ẢNH {number:02d}/{len(SCREENSHOTS):02d}", f"Tên file bắt buộc: {shot['filename']}",
        f"Tiêu đề ảnh: {shot['title']}", "Chèn ảnh tại vị trí này.", f"URL hoặc lệnh: {shot['location']}",
        action, f"Panel/DevTools: {shot['panel']}", f"Nội dung bắt buộc phải thấy: {shot['required']}",
        f"Kết quả mong đợi: {shot['expected']}", f"Caption: {shot['caption']}",
    ]


def add_docx_shot(doc: Document, index: int, missing: list[str]) -> None:
    shot = SCREENSHOTS[index]
    number = index + 1
    path = SHOT_DIR / shot["filename"]
    if valid_png(path):
        w, h = picture_inches(path)
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_with_next = True
        p.add_run().add_picture(str(path), width=Inches(w), height=Inches(h))
    else:
        missing.append(shot["filename"])
        table = doc.add_table(rows=1, cols=1); table.style = "Table Grid"
        table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        cell = table.cell(0, 0); shade(cell, "FFF4D6"); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for i, line in enumerate(placeholder_lines(number, shot)):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line); run.font.name = "Arial"; run.font.size = Pt(9); run.bold = i in (0, 1, 2, 3)
    cap = doc.add_paragraph(f"Hình {number}. {shot['caption']}", style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_docx() -> list[str]:
    doc = Document(); configure_docx(doc); add_cover(doc)
    missing: list[str] = []
    for title, paragraphs, shot_indexes in SECTIONS:
        doc.add_heading(title, 1)
        if title.startswith("8."):
            doc.add_paragraph(pytest_summary())
        for paragraph in paragraphs:
            doc.add_paragraph(paragraph)
        for index in shot_indexes:
            add_docx_shot(doc, index, missing)
    REPORT_DIR.mkdir(exist_ok=True); doc.save(DOCX_PATH)
    return missing


def pdf_font() -> tuple[str, str]:
    normal = Path("C:/Windows/Fonts/arial.ttf"); bold = Path("C:/Windows/Fonts/arialbd.ttf")
    if normal.exists() and bold.exists():
        pdfmetrics.registerFont(TTFont("ArialVN", str(normal))); pdfmetrics.registerFont(TTFont("ArialVN-B", str(bold)))
        return "ArialVN", "ArialVN-B"
    return "Helvetica", "Helvetica-Bold"


def add_pdf_shot(story: list, index: int, body, bold_font: str, missing: list[str]) -> None:
    shot = SCREENSHOTS[index]; number = index + 1; path = SHOT_DIR / shot["filename"]
    story.append(PageBreak())
    block = []
    if valid_png(path):
        with PILImage.open(path) as image:
            scale = min((17 * cm) / image.width, (11 * cm) / image.height)
            block.append(Image(str(path), width=image.width * scale, height=image.height * scale))
    else:
        if shot["filename"] not in missing: missing.append(shot["filename"])
        content = "<br/><br/>".join(line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;") for line in placeholder_lines(number, shot))
        block.append(Table([[Paragraph(content, body)]], colWidths=[17 * cm], splitInRow=0, style=TableStyle([
            ("BOX", (0, 0), (-1, -1), 1, colors.HexColor("#D49A18")), ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#FFF4D6")),
            ("LEFTPADDING", (0, 0), (-1, -1), 8), ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ("TOPPADDING", (0, 0), (-1, -1), 7), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ])))
    block.extend([Paragraph(f"Hình {number}. {shot['caption']}", body), Spacer(1, 8)])
    story.extend(block)


def build_pdf(missing: list[str]) -> None:
    font, bold = pdf_font(); samples = getSampleStyleSheet()
    body = ParagraphStyle("BodyVN", parent=samples["BodyText"], fontName=font, fontSize=9.5, leading=13, spaceAfter=5)
    h1 = ParagraphStyle("H1VN", parent=samples["Heading1"], fontName=bold, fontSize=15, leading=18, textColor=colors.HexColor("#174A84"), spaceBefore=9, spaceAfter=6)
    title = ParagraphStyle("TitleVN", parent=samples["Title"], fontName=bold, fontSize=24, textColor=colors.HexColor("#174A84"), alignment=1)
    story = [Spacer(1, 6 * cm), Paragraph("BÁO CÁO THỰC HÀNH LAB 01", title), Paragraph("CROSS-SITE SCRIPTING", title),
             Paragraph("MSSV: 21127645 - Họ tên: Lê Minh", body), PageBreak()]
    for section_title, paragraphs, shot_indexes in SECTIONS:
        story.append(Paragraph(section_title, h1))
        if section_title.startswith("8."):
            story.append(Paragraph(pytest_summary(), body))
        for paragraph in paragraphs:
            story.append(Paragraph(paragraph, body))
        for index in shot_indexes:
            add_pdf_shot(story, index, body, bold, missing)
    SimpleDocTemplate(str(PDF_PATH), pagesize=A4, leftMargin=2 * cm, rightMargin=2 * cm, topMargin=1.7 * cm, bottomMargin=1.7 * cm).build(story)


def main() -> int:
    SHOT_DIR.mkdir(parents=True, exist_ok=True); REPORT_DIR.mkdir(exist_ok=True)
    missing = build_docx(); build_pdf(missing)
    missing = [item["filename"] for item in SCREENSHOTS if item["filename"] in missing]
    print(f"Đã tạo DOCX: {DOCX_PATH}")
    print(f"Đã tạo PDF:  {PDF_PATH}")
    print(f"Ảnh còn thiếu ({len(missing)}): {', '.join(missing) if missing else 'không'}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
