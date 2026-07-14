import html
import json
import os
import re
import subprocess
import sys
from pathlib import Path

ROOT=Path(__file__).parents[1]
PAYLOAD='<img src=x onerror="alert(\'XSS\')">'

def trace(response):
    match=re.search(r'<pre class="trace-json" hidden>(.*?)</pre>',response.text,re.S)
    assert match
    return json.loads(html.unescape(match.group(1)))

def test_routes(client):
    for route in ["/","/vulnerable/search","/secure/search","/vulnerable/post/1/comments","/secure/post/1/comments","/vulnerable/dom-search","/secure/dom-search","/profile","/security-headers","/test-results"]:
        assert client.get(route).status_code==200

def test_reflected_trace_and_modes(client):
    vulnerable=client.get("/vulnerable/search",query_string={"q":PAYLOAD})
    secure=client.get("/secure/search",query_string={"q":PAYLOAD})
    assert PAYLOAD in vulnerable.text and "&lt;img" in secure.text
    vt,st=trace(vulnerable),trace(secure)
    assert vt["input_source"]=='request.args["q"]'
    assert {"Source","Tainted","Sink","XSS"}<={n["risk"] for n in vt["source_to_sink"]}
    assert vt["final_result"]["verdict"].startswith("Payload đã thực thi")
    assert st["final_result"]["verdict"]=="Payload không thực thi"
    assert st["response_summary"]["before_escape"]==PAYLOAD
    assert "&lt;img" in st["response_summary"]["after_escape"]

def test_cookie_is_masked_in_request_inspector(client):
    client.set_cookie("demo","sensitive-value")
    summary=trace(client.get("/secure/search?q=x"))["request_summary"]
    assert "sensitive-value" not in summary["cookie"] and "***" in summary["cookie"]

def test_stored_trace_database_and_sanitization(client):
    vulnerable=client.post("/vulnerable/post/1/comments",data={"author":"A","body":PAYLOAD})
    assert vulnerable.status_code==200 and PAYLOAD in vulnerable.text
    vt=trace(vulnerable); techniques=" ".join(s["title"] for s in vt["steps"])
    assert "SQLite INSERT" in techniques and "SQLite SELECT" in techniques
    assert vt["database_inspector"]["latest"]["body"]==PAYLOAD
    secure=client.get("/secure/post/1/comments"); st=trace(secure)
    assert "onerror" not in secure.text.split("Dữ liệu gốc trong DB",1)[0]
    assert st["security_controls"][1]["name"]=="Bleach sanitization"
    assert st["final_result"]["sanitized"] is True

def test_validation_errors(client):
    response=client.post("/vulnerable/post/1/comments",data={"author":"","body":""})
    assert "không được để trống" in response.text and response.status_code==200

def test_dom_trace_and_files(client):
    vulnerable=trace(client.get("/vulnerable/dom-search")); secure=trace(client.get("/secure/dom-search"))
    assert vulnerable["request_summary"]["path"]=="/vulnerable/dom-search"
    assert "fragment" in vulnerable["steps"][1]["description"].lower()
    assert vulnerable["source_to_sink"][2]["name"]=="innerHTML"
    assert secure["source_to_sink"][2]["name"]=="textContent"
    v=(ROOT/"static/js/dom_vulnerable.js").read_text(encoding="utf-8")
    s=(ROOT/"static/js/dom_secure.js").read_text(encoding="utf-8")
    assert "innerHTML" in v and "textContent" in s and "eval(" not in s and "document.write" not in s

def test_csp_cookie_and_header_trace(client):
    response=client.get("/security-headers"); policy=response.headers["Content-Security-Policy"]
    for directive in ["default-src 'self'","script-src 'self'","object-src 'none'","frame-ancestors 'none'"]: assert directive in policy
    assert "unsafe-inline" not in policy
    assert any("CSP" in step["title"] for step in trace(response)["steps"])
    cookie=client.get("/profile").headers["Set-Cookie"]
    assert "HttpOnly" in cookie and "SameSite=Lax" in cookie

def test_trace_ui_features_exist(client):
    text=client.get("/secure/search?q=demo").text
    for value in ["Presentation Mode","Auto Play","Xuất JSON","Xóa timeline","data-action=\"next\"","Request Inspector","Response Inspector","Source → Sink"]: assert value in text
    js=(ROOT/"static/js/trace-ui.js").read_text(encoding="utf-8")
    assert "application/json" in js and "replaceChildren" in js and "navigator.clipboard" in js

def test_no_external_request_or_secret():
    files=list(ROOT.glob("templates/**/*.html"))+list(ROOT.glob("static/js/*.js"))+[ROOT/"trace_service.py"]
    text="\n".join(path.read_text(encoding="utf-8") for path in files)
    assert not re.search(r"https?://(?!127\.0\.0\.1|localhost)",text)
    assert "sensitive-value" not in text

def test_no_browser_automation_dependency():
    sources=[p for p in ROOT.rglob("*") if p.is_file() and ".venv" not in p.parts and "tests" not in p.parts and p.suffix in {".py",".txt",".md"}]
    hits=[str(p) for p in sources if re.search(r"\b(playwright|selenium)\b",p.read_text(encoding="utf-8",errors="ignore"),re.I)]
    assert not hits
    assert "playwright" not in (ROOT/"requirements.txt").read_text(encoding="utf-8").lower()

def test_manual_guide_has_all_required_names():
    guide=(ROOT/"HUONG_DAN_CHUP_ANH.md").read_text(encoding="utf-8")
    from screenshot_manifest import SCREENSHOTS
    assert len(SCREENSHOTS)==10
    assert all(item["filename"] in guide for item in SCREENSHOTS)
    assert guide.count("**Tên file:**") == 10
    assert guide.count("**Caption dùng trong báo cáo:**") == 10

def test_screenshot_checker_lists_missing():
    result=subprocess.run([sys.executable,"scripts/check_screenshots.py"],cwd=ROOT,text=True,capture_output=True,encoding="utf-8",env={**os.environ,"PYTHONIOENCODING":"utf-8"})
    assert result.returncode==1 and "Đúng tên: 0/10" in result.stdout and "01_reflected_vulnerable.png" in result.stdout

def test_screenshot_checker_lists_required_in_order():
    result=subprocess.run([sys.executable,"scripts/check_screenshots.py","--list-required"],cwd=ROOT,text=True,capture_output=True,encoding="utf-8",env={**os.environ,"PYTHONIOENCODING":"utf-8"})
    assert result.returncode==0
    assert result.stdout.index("01_reflected_vulnerable.png") < result.stdout.index("10_tests_reports.png")

def test_reports_generate_with_placeholders():
    result=subprocess.run([sys.executable,"scripts/generate_report.py"],cwd=ROOT,text=True,capture_output=True,encoding="utf-8",env={**os.environ,"PYTHONIOENCODING":"utf-8"})
    assert result.returncode==0 and "Ảnh còn thiếu (10)" in result.stdout
    assert (ROOT/"report/21127645_LeMinh_Lab01_XSS.docx").stat().st_size>10_000
    assert (ROOT/"report/21127645_LeMinh_Lab01_XSS.pdf").stat().st_size>10_000
    from docx import Document
    doc=Document(ROOT/"report/21127645_LeMinh_Lab01_XSS.docx")
    text="\n".join(p.text for p in doc.paragraphs)
    cells="\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    assert "1. Mục tiêu và môi trường thực hành" in text
    assert "9. Kết luận" in text
    assert "ẢNH 01/10" in cells and "01_reflected_vulnerable.png" in cells
    assert cells.count("Chèn ảnh tại vị trí này.") == 10
