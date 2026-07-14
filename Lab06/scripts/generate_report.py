#!/usr/bin/env python3
"""Generate the evidence-backed Lab06 DOCX and PDF reports.

The generator deliberately fails closed. It will not create final artifacts unless
all required evidence families exist, pytest reports at least 80 passing tests with
no failures/errors, every required core module has at least 90% coverage, the
runtime smoke log contains an explicit success marker, and fifteen named trace
flows can be resolved from real trace JSON.

The report is text-only by design. No screenshots, generated illustrations, blank
image frames, or image placeholders are added. Diagrams are represented as text.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import sys
import tempfile
import textwrap
import zipfile
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable, Mapping, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    KeepTogether,
    LongTable,
    PageBreak,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
EVIDENCE_ROOT = ROOT / "evidence"
REPORT_ROOT = ROOT / "report"
DOCX_NAME = "21127645_LeMinh_Lab06_CookiePoisoning.docx"
PDF_NAME = "21127645_LeMinh_Lab06_CookiePoisoning.pdf"

MIN_TESTS = 80
MIN_CORE_COVERAGE = 90
MIN_PDF_PAGES = 18
MAX_PDF_PAGES = 25
MAX_INPUT_FILE_BYTES = 10 * 1024 * 1024
DESIGN_PRESET = "compact_reference_guide"
COVER_PATTERN = "editorial_cover"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_DXA = 45
CELL_MARGIN_BOTTOM_DXA = 45
CELL_MARGIN_START_DXA = 80
CELL_MARGIN_END_DXA = 80

INK = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "596675"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
GOLD = "A56F00"
WHITE = "FFFFFF"

CORE_MODULES = (
    "database.py",
    "auth_service.py",
    "cookie_service.py",
    "base64_cookie_service.py",
    "signed_cookie_service.py",
    "encrypted_cookie_service.py",
    "server_session_service.py",
    "authorization_service.py",
    "audit_service.py",
    "trace_service.py",
)

REQUIRED_SOURCE_FILES = (
    "app.py",
    "config.py",
    "database.py",
    "auth_service.py",
    "cookie_service.py",
    "base64_cookie_service.py",
    "signed_cookie_service.py",
    "encrypted_cookie_service.py",
    "server_session_service.py",
    "authorization_service.py",
    "audit_service.py",
    "trace_models.py",
    "trace_service.py",
    "schema.sql",
    "README.md",
    "evidence/logs/requirements_review.txt",
)

REQUIRED_EVIDENCE_FAMILIES = (
    "traces",
    "requests",
    "responses",
    "cookies",
    "sessions",
    "audit",
)


class ReportInputError(RuntimeError):
    """Raised when truthful report generation is not possible."""


class ReportBuildError(RuntimeError):
    """Raised when an output fails structural verification."""


@dataclass(frozen=True)
class TestSummary:
    passed: int
    failed: int
    errors: int
    skipped: int
    source: Path


@dataclass(frozen=True)
class CoverageSummary:
    modules: Mapping[str, int]
    total: int | None
    source: Path


@dataclass(frozen=True)
class SmokeSummary:
    status: str
    marker: str
    source: Path


@dataclass(frozen=True)
class EvidenceFile:
    family: str
    path: Path
    size: int
    sha256: str
    text: str
    parsed: Any | None

    @property
    def relative_path(self) -> str:
        return self.path.relative_to(ROOT).as_posix()


@dataclass(frozen=True)
class FlowSpec:
    flow_id: str
    title: str
    aliases: tuple[tuple[str, ...], ...]


@dataclass(frozen=True)
class FlowTrace:
    spec: FlowSpec
    source: EvidenceFile
    record: Mapping[str, Any]


@dataclass(frozen=True)
class TableSpec:
    title: str
    headers: tuple[str, ...]
    rows: tuple[tuple[str, ...], ...]
    widths_dxa: tuple[int, ...]
    kind: str = "standard"


@dataclass
class Chapter:
    number: int
    title: str
    paragraphs: list[str] = field(default_factory=list)
    code_blocks: list[str] = field(default_factory=list)
    tables: list[TableSpec] = field(default_factory=list)
    qa_answers: list[tuple[str, str]] = field(default_factory=list)


@dataclass(frozen=True)
class ReportInputs:
    generated_at: str
    tests: TestSummary
    coverage: CoverageSummary
    smoke: SmokeSummary
    evidence: Mapping[str, tuple[EvidenceFile, ...]]
    flows: tuple[FlowTrace, ...]
    sources: Mapping[str, str]
    source_hashes: Mapping[str, str]


FLOW_SPECS = (
    FlowSpec("plain_login", "Plain cookie login", (("plain_login_student",), ("plain_cookie_login",))),
    FlowSpec("plain_admin_denied", "Plain admin denied", (("plain_admin_denied",),)),
    FlowSpec("plain_cookie_modified", "Plain cookie modified", (("plain_admin_cookie_modified",), ("plain_cookie_modified",))),
    FlowSpec("base64_original", "Base64 original", (("base64_admin_denied",), ("base64_original",))),
    FlowSpec("base64_modified", "Base64 modified", (("base64_admin_cookie_modified",), ("base64_modified",))),
    FlowSpec("signed_cookie_valid", "Signed cookie valid", (("signed_cookie_valid",),)),
    FlowSpec("signed_cookie_invalid", "Signed cookie invalid", (("signed_cookie_tampered",), ("signed_cookie_invalid",))),
    FlowSpec("encrypted_cookie_valid", "Encrypted cookie valid", (("encrypted_cookie_valid",),)),
    FlowSpec("encrypted_cookie_tampered", "Encrypted cookie tampered", (("encrypted_cookie_tampered",),)),
    FlowSpec("server_session_student_login", "Server-side student login", (("server_session_student_login",),)),
    FlowSpec("student_admin_denied", "Student admin denied", (("server_session_student_admin_denied",), ("student_admin_denied",))),
    FlowSpec("admin_allowed", "Admin allowed", (("server_session_admin_allowed",), ("admin_allowed",))),
    FlowSpec("session_rotation", "Session rotation", (("server_session_rotated",), ("session_rotation",))),
    FlowSpec("logout_invalidation", "Logout invalidation", (("logout_invalidation",), ("server_session_logout",), ("logout", "revoked"))),
    FlowSpec("old_session_rejection", "Old session rejection", (("old_session_rejected",), ("old_session_rejection",))),
)


CHAPTER_TITLES = (
    "Giới thiệu",
    "Mục tiêu và phạm vi an toàn",
    "Cookie là gì",
    "Cấu trúc và thuộc tính cookie",
    "Cookie là dữ liệu không đáng tin cậy",
    "Cookie Poisoning là gì",
    "Kiến trúc Lab06",
    "Database schema và tài khoản demo",
    "Plain Cookie vulnerable",
    "Phân tích thay đổi role trong plain cookie",
    "Base64 Cookie vulnerable",
    "Base64 không phải mã hóa",
    "Signed Cookie",
    "HMAC và kiểm tra toàn vẹn",
    "Encrypted Cookie",
    "Authenticated Encryption",
    "Server-side Session",
    "Session ID ngẫu nhiên",
    "Session rotation",
    "Logout và server-side invalidation",
    "Server-side authorization",
    "Cookie flags",
    "Cookie Poisoning và Session Hijacking",
    "Cookie Poisoning và Session Fixation",
    "Cookie Stealing và XSS",
    "So sánh năm mô hình",
    "Code Comparison",
    "Audit logging",
    "Kết quả kiểm thử",
    "Coverage",
    "Runtime smoke test",
    "Mức độ ảnh hưởng",
    "Biện pháp phòng chống",
    "Bài học rút ra",
    "Kết luận",
)


BASE_CHAPTER_PARAGRAPHS: Mapping[int, tuple[str, ...]] = {
    1: (
        "Lab06 minh họa Cookie Poisoning trong một ứng dụng Flask chỉ chạy local. Trọng tâm là ranh giới tin cậy: browser giữ cookie nhưng server chịu trách nhiệm xác thực, quản lý phiên và phân quyền.",
        "Báo cáo được tạo từ source, trace, request, response, cookie, session, audit và log kiểm thử thực. Các số liệu vận hành trong tài liệu không được suy diễn từ cấu hình hay ghi cứng.",
    ),
    2: (
        "Mục tiêu là quan sát dữ liệu cookie, tái hiện lỗi role phía client trong hai mode cố ý dễ bị tấn công, rồi đối chiếu với signing, authenticated encryption và server-side session.",
        "Phạm vi bị giới hạn tại http://127.0.0.1:5006. Lab không truy cập website thật, không kết nối Internet khi runtime, không có proxy, cookie editor tổng quát, XSS, hijacking, replay UI hay công cụ tự động hóa trình duyệt.",
    ),
    3: (
        "Cookie là cặp tên-giá trị do server đặt trong response và browser gửi lại theo phạm vi Domain/Path cùng các thuộc tính liên quan. Cookie là cơ chế vận chuyển state, không tự tạo ra niềm tin.",
        "Dữ liệu nhận từ Cookie header phải được kiểm tra giống mọi input khác. Một tên cookie quen thuộc hoặc flag bảo mật không chứng minh nội dung đúng.",
    ),
    4: (
        "Name và Value xác định dữ liệu; Domain và Path giới hạn nơi gửi; HttpOnly hạn chế JavaScript đọc cookie; Secure giới hạn truyền qua HTTPS; SameSite giảm một số request cross-site.",
        "Các flag hỗ trợ giảm bề mặt tấn công nhưng không ký, mã hóa hoặc xác minh role trong payload. Giá trị hiển thị trong báo cáo phải lấy từ response/config evidence thực.",
    ),
    5: (
        "Người dùng kiểm soát browser, storage và request gửi đi. Vì vậy role, is_admin, balance hay permission nằm trong cookie đều có thể bị thay đổi nếu không có cơ chế integrity và policy server-side.",
        "Server phải xác minh danh tính, trạng thái phiên và quyền từ nguồn có thẩm quyền trên mỗi request nhạy cảm.",
    ),
    6: (
        "Cookie Poisoning là sửa nội dung cookie trước khi gửi lại để làm thay đổi logic ứng dụng. Trong lab, giá trị role=user được đổi thủ công thành role=admin tại đúng origin local.",
        "Root cause không phải DevTools mà là quyết định phân quyền dựa trên dữ liệu client-controlled. Tác động minh họa là leo thang đặc quyền trong trang admin giả lập.",
    ),
    7: (
        "Kiến trúc tách route Flask, service cookie/session, authorization, SQLite, audit và trace. Plain/Base64 cố ý tin dữ liệu client; Signed xác minh integrity trước khi dùng payload; Session lấy role từ database.",
    ),
    8: (
        "SQLite chứa users, server_sessions, audit_logs, cookie_events và session_events. Mọi truy vấn phải có tham số; database chỉ giữ hash Session ID, không giữ token thô.",
        "Hai tài khoản demo dùng password hash PBKDF2-SHA256 600000 vòng. Password demo chỉ dùng ở form login và không xuất hiện trong evidence.",
    ),
    9: (
        "Plain mode đặt lab06_username và lab06_role. Route admin cố ý so sánh trực tiếp role trong request.cookies, không kiểm tra integrity và không truy vấn database role.",
    ),
    10: (
        "Trạng thái ban đầu role=user dẫn tới deny. Khi người học sửa thủ công thành admin, request sau được allow vì server không phân biệt được giá trị ban đầu và giá trị đã bị sửa.",
        "Bản vá là không dùng role client làm nguồn quyền; thay bằng session opaque và database-backed authorization.",
    ),
    11: (
        "Base64 mode lưu JSON username/role trong lab06_profile_b64. Server URL-safe Base64 decode, parse JSON rồi cố ý dùng role mà không có chữ ký.",
    ),
    12: (
        "Base64 biến bytes thành ký tự thuận tiện vận chuyển. Bất kỳ ai có chuỗi đều có thể decode, sửa JSON và encode lại; do đó confidentiality và integrity đều không tồn tại.",
    ),
    13: (
        "Signed cookie gắn chữ ký HMAC thông qua itsdangerous. Server phải xác minh chữ ký trước khi chấp nhận payload; token thiếu, sai format hoặc bị sửa phải bị từ chối.",
        "Payload signed vẫn có thể đọc được nếu format lộ nội dung. Signing không phải encryption và không thay thế kiểm tra role mới nhất từ database.",
    ),
    14: (
        "HMAC tính mã xác thực từ payload và secret server-side. Sửa payload hoặc signature làm kết quả tính lại không khớp, giúp phát hiện mất integrity và authenticity.",
        "Secret không được đưa vào client, log, trace hay báo cáo. Revocation của signed payload khó hơn server-side session nếu không có state kiểm soát bổ sung.",
    ),
    15: (
        "Encrypted demo dùng Fernet cho payload read-only gồm user_id, display_name, preference và issued_at. Payload không có role, password hoặc Session ID thật.",
    ),
    16: (
        "Authenticated encryption đồng thời che plaintext và xác thực token. Sửa một byte hoặc dùng sai key làm decrypt thất bại trước khi plaintext được sử dụng.",
        "Encryption bảo vệ dữ liệu, nhưng policy vẫn phải quyết định quyền từ nguồn server-side có thể cập nhật và revoke.",
    ),
    17: (
        "Server-side session đặt một Session ID ngẫu nhiên trong lab06_session. Server hash token, tra record active/unexpired, lấy user và role từ database rồi mới authorize.",
    ),
    18: (
        "Session ID được sinh bằng secrets.token_urlsafe(32) để có entropy cao và không mang ý nghĩa nghiệp vụ. Cookie không chứa role hoặc user_id.",
        "Database lưu SHA-256 của token để giảm tác động nếu bảng session bị lộ; so khớp được thực hiện bằng cách hash token request.",
    ),
    19: (
        "Rotation sau login tạo token mới và thu hồi token trước đó. Cơ chế này cắt liên kết với định danh phiên cũ và hỗ trợ giảm rủi ro fixation.",
    ),
    20: (
        "Logout phải đổi state server thành inactive/revoked và đồng thời expire cookie browser. Chỉ xóa cookie client sẽ để token cũ còn hợp lệ ở server.",
    ),
    21: (
        "Authorization Inspector phải thể hiện subject, action, policy, nguồn role, database role, decision và reason. Student bị deny; admin chỉ được allow sau database lookup.",
        "Thay đổi role trong database phải có hiệu lực ở request kế tiếp mà không cần phát hành cookie role mới.",
    ),
    22: (
        "HttpOnly giảm khả năng script đọc cookie; Secure yêu cầu HTTPS; SameSite giảm một số CSRF. Với HTTP local, Secure phải phản ánh cấu hình thực thay vì được ghi đạt giả.",
    ),
    23: (
        "Cookie Poisoning thay đổi nội dung để đổi logic. Session Hijacking chiếm một phiên hợp lệ, thường bằng cách lấy Session ID. Hai hành vi có mục tiêu và biện pháp kiểm soát khác nhau.",
    ),
    24: (
        "Session Fixation ép nạn nhân dùng Session ID đã biết rồi chờ đăng nhập. Cookie Poisoning sửa dữ liệu cookie để lừa logic. Rotation sau login là kiểm soát quan trọng chống fixation.",
    ),
    25: (
        "Cookie Stealing qua XSS lấy cookie hiện có; HttpOnly có thể giảm khả năng script đọc token nhưng không sửa lỗi XSS và không ngăn người dùng sửa cookie của chính họ bằng DevTools.",
    ),
    26: (
        "Ma trận dưới đây so sánh nơi lưu state, khả năng đọc/sửa, integrity, confidentiality, revocation và nguồn authorization của năm mô hình.",
    ),
    27: (
        "Code Comparison được trích trực tiếp từ marker trong source hiện hành. Báo cáo không sử dụng snippet được viết lại độc lập với implementation.",
    ),
    28: (
        "Audit ghi actor, action, route, mode, trạng thái cookie, role submitted/database, decision, reason và trace ID. Log chỉ chứa giá trị đã mask hoặc fingerprint.",
    ),
    29: (
        "Kết quả chương này được parse từ evidence/logs/pytest.txt. Báo cáo chỉ được tạo khi có ít nhất 80 test pass và không có failure/error.",
    ),
    30: (
        "Coverage được parse theo từng module lõi từ evidence/logs/coverage.txt. Mọi module bắt buộc phải đạt tối thiểu 90%; không dùng total để che module thiếu.",
    ),
    31: (
        "Runtime smoke phải có marker thành công rõ ràng sau khi app thật được khởi động, health check và các flow local cố định đã chạy. Log mơ hồ hoặc chứa failure làm generator dừng.",
    ),
    32: (
        "Plain/Base64 có thể dẫn tới broken access control và leo thang đặc quyền. Signed/encrypted giảm một số rủi ro payload nhưng vẫn còn rủi ro key management, replay và state stale nếu dùng sai.",
    ),
    33: (
        "Biện pháp chính: không lưu quyền động ở client; dùng opaque session; hash token server-side; rotate; revoke khi logout; kiểm tra role từ database; ký/mã hóa đúng mục đích; giới hạn log và áp dụng cookie flags.",
    ),
    34: (
        "Bài học trung tâm là phân biệt vận chuyển state với nguồn thẩm quyền. Encoding, signing, encryption, flags và session giải quyết các thuộc tính bảo mật khác nhau; không một kỹ thuật đơn lẻ thay thế authorization.",
    ),
    35: (
        "Lab06 chứng minh lỗi phát sinh khi server trao niềm tin cho role phía client và cách chuyển niềm tin về server-side state. Kết luận vận hành chỉ dựa trên evidence và log thực đã kiểm tra.",
    ),
}


QA_ANSWERS = (
    ("Vì sao cookie là dữ liệu không đáng tin cậy?", "Cookie nằm trong browser do người dùng kiểm soát và có thể bị sửa trước khi gửi. Server phải coi nó là input không tin cậy."),
    ("Cookie Poisoning khác Session Hijacking như thế nào?", "Poisoning sửa nội dung để đổi logic; hijacking chiếm một phiên hợp lệ, thường bằng Session ID bị lộ."),
    ("Base64 có phải mã hóa không?", "Không. Base64 là encoding có thể đảo ngược công khai, không có key, confidentiality hoặc integrity."),
    ("Signed cookie giải quyết vấn đề gì?", "Nó phát hiện sửa đổi và xác thực nguồn tạo payload khi secret được bảo vệ."),
    ("Vì sao server-side authorization vẫn bắt buộc?", "Quyền có thể thay đổi hoặc bị thu hồi; server phải kiểm tra policy và nguồn role có thẩm quyền trên mỗi request."),
    ("Signed cookie khác encrypted cookie thế nào?", "Signing bảo vệ integrity/authenticity nhưng thường không che nội dung; encryption che nội dung, và Fernet còn xác thực integrity."),
    ("Vì sao mã hóa không thay thế kiểm tra quyền?", "Decrypt thành công chỉ chứng minh token hợp lệ theo key, không chứng minh subject được phép thực hiện action hiện tại."),
    ("Vì sao không nên lưu role, balance hoặc is_admin trong cookie?", "Đó là state nhạy cảm và động; nếu client sửa được hoặc dữ liệu stale, quyết định nghiệp vụ bị sai."),
    ("HttpOnly giải quyết rủi ro gì?", "Nó hạn chế JavaScript đọc cookie và giảm một phần tác động đánh cắp token qua XSS."),
    ("Vì sao HttpOnly không ngăn sửa cookie bằng DevTools?", "HttpOnly là ràng buộc với script, không tước quyền quản lý storage của người dùng/browser."),
    ("Secure flag giải quyết rủi ro gì?", "Nó yêu cầu browser chỉ gửi cookie qua HTTPS, giảm lộ token trên kênh HTTP."),
    ("SameSite giải quyết rủi ro gì?", "Nó hạn chế gửi cookie trong một số ngữ cảnh cross-site và hỗ trợ giảm CSRF."),
    ("Vì sao cookie flags không kiểm tra tính toàn vẹn nội dung?", "Flags điều khiển cách browser lưu/gửi cookie; chúng không tính MAC hoặc xác minh payload."),
    ("Session ID ngẫu nhiên hoạt động ra sao?", "Browser giữ token entropy cao; server dùng token để tìm state phiên, còn dữ liệu user/role nằm server-side."),
    ("Vì sao chỉ lưu hash Session ID tốt hơn token thô?", "Nếu database bị đọc, hash không thể dùng trực tiếp làm cookie để truy cập phiên."),
    ("Vì sao phải rotate session sau login?", "Rotation vô hiệu định danh trước xác thực và giảm nguy cơ fixation hoặc reuse token cũ."),
    ("Vì sao logout phải hủy session phía server?", "Nếu chỉ xóa cookie, bản sao token cũ vẫn có thể được server chấp nhận đến khi hết hạn."),
    ("Vì sao signed cookie khó revoke hơn server-side session?", "Token tự chứa state thường hợp lệ đến khi hết hạn nếu server không giữ deny-list/version; session record có thể revoke tức thời."),
    ("Cookie Poisoning khác Parameter Tampering thế nào?", "Cookie Poisoning là tampering trên cookie state; Parameter Tampering rộng hơn và có thể nhắm query, form, path hoặc JSON body."),
    ("XSS ảnh hưởng tới cookie ra sao?", "XSS có thể thực hiện action trong origin và đọc cookie không HttpOnly; HttpOnly giảm đọc token nhưng không loại bỏ XSS."),
    ("Session Fixation khác Cookie Poisoning thế nào?", "Fixation ép dùng Session ID biết trước; Poisoning sửa dữ liệu cookie để thay đổi hành vi server."),
    ("generate_report.py bảo đảm tính trung thực thế nào?", "Script dừng khi thiếu evidence/log, parse kết quả thật, yêu cầu 15 trace, kiểm output trước atomic replace và không thêm ảnh giả hoặc khung ảnh trống."),
)


SEQUENCE_DIAGRAM = """Browser local -> Flask /login: credentials + fixed mode
Flask -> SQLite users: parameterized user lookup
SQLite --> Flask: password hash + current role
Flask -> Browser: Set-Cookie for selected mode
Browser -> Flask protected route: fixed Lab06 cookie
Flask -> Verification/Session service: decode, verify, decrypt, or hash lookup
Verification/Session service -> Authorization: trusted facts only
Authorization -> SQLite: current role lookup for secure modes
Authorization --> Flask: allow/deny + reason
Flask -> Audit/Trace: redacted event + evidence link
Flask --> Browser: response + Inspector + Final Verdict"""


DATA_FLOW_DIAGRAM = """[Untrusted Browser Storage]
        | Cookie header
        v
[Flask Request Boundary] -- Plain/Base64 --> [Vulnerable client-role decision]
        |
        +-- Signed --> [Verify integrity] --> [Database role]
        |
        +-- Encrypted --> [Fernet decrypt] --> [Read-only presentation]
        |
        +-- Session --> [SHA-256 lookup] --> [Active/expiry] --> [Database role]
                                                        |
                                                        v
                                              [Authorization policy]
                                                        |
                                      [Audit + Trace + masked evidence]"""


def _read_required_text(path: Path) -> str:
    if not path.is_file():
        raise ReportInputError(f"Thiếu file bắt buộc: {path}")
    size = path.stat().st_size
    if size <= 0:
        raise ReportInputError(f"File bắt buộc rỗng: {path}")
    if size > MAX_INPUT_FILE_BYTES:
        raise ReportInputError(f"File vượt giới hạn an toàn {MAX_INPUT_FILE_BYTES} bytes: {path}")
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError as exc:
        raise ReportInputError(f"File không phải UTF-8: {path}") from exc


def _count_summary_value(text: str, word: str) -> int:
    matches = re.findall(rf"(?<!\d)(\d+)\s+{word}s?\b", text, flags=re.IGNORECASE)
    return int(matches[-1]) if matches else 0


def parse_pytest_log(path: Path) -> TestSummary:
    text = _read_required_text(path)
    passed = _count_summary_value(text, "passed")
    failed = _count_summary_value(text, "failed")
    errors = _count_summary_value(text, "error")
    skipped = _count_summary_value(text, "skipped")
    if passed < MIN_TESTS:
        raise ReportInputError(f"Pytest chưa đủ {MIN_TESTS} test pass: đọc được {passed} từ {path}")
    if failed or errors:
        raise ReportInputError(f"Pytest chưa sạch: failed={failed}, errors={errors} trong {path}")
    return TestSummary(passed, failed, errors, skipped, path)


def parse_coverage_log(path: Path) -> CoverageSummary:
    text = _read_required_text(path)
    modules: dict[str, int] = {}
    total: int | None = None
    row_pattern = re.compile(
        r"^\s*(?P<name>\S+\.py)\s+\d+\s+\d+\s+(?P<pct>\d{1,3})%",
        re.MULTILINE,
    )
    for match in row_pattern.finditer(text):
        name = Path(match.group("name").replace("\\", "/")).name
        modules[name] = int(match.group("pct"))
    total_matches = re.findall(r"^\s*TOTAL\s+\d+\s+\d+\s+(\d{1,3})%", text, re.MULTILINE)
    if total_matches:
        total = int(total_matches[-1])
    missing = [name for name in CORE_MODULES if name not in modules]
    if missing:
        raise ReportInputError("Coverage log thiếu module lõi: " + ", ".join(missing))
    below = {name: modules[name] for name in CORE_MODULES if modules[name] < MIN_CORE_COVERAGE}
    if below:
        detail = ", ".join(f"{name}={pct}%" for name, pct in below.items())
        raise ReportInputError(f"Coverage module lõi dưới {MIN_CORE_COVERAGE}%: {detail}")
    return CoverageSummary(dict(sorted(modules.items())), total, path)


def _json_has_success(value: Any) -> bool:
    if isinstance(value, Mapping):
        for key, item in value.items():
            key_l = str(key).lower()
            if key_l in {"success", "passed", "ok"} and item is True:
                return True
            if key_l in {"status", "result"} and str(item).lower() in {"ok", "passed", "success", "healthy"}:
                return True
            if _json_has_success(item):
                return True
    elif isinstance(value, list):
        return any(_json_has_success(item) for item in value)
    return False


def _json_has_failure(value: Any) -> bool:
    if isinstance(value, Mapping):
        for key, item in value.items():
            key_l = str(key).lower()
            if key_l in {"success", "passed", "ok"} and item is False:
                return True
            if key_l in {"status", "result"} and str(item).lower() in {
                "failed", "failure", "error", "unhealthy",
            }:
                return True
            if _json_has_failure(item):
                return True
    elif isinstance(value, list):
        return any(_json_has_failure(item) for item in value)
    return False


def parse_smoke_log(path: Path) -> SmokeSummary:
    text = _read_required_text(path)
    lowered = text.lower()
    negative = re.search(r"\b(smoke[_ -]?test[_ -]?failed|status\s*[:=]\s*failed|result\s*[:=]\s*failed)\b", lowered)
    if negative:
        raise ReportInputError(f"Runtime smoke log có marker thất bại: {path}")
    marker = ""
    try:
        parsed = json.loads(text)
    except json.JSONDecodeError:
        parsed = None
    if parsed is not None and _json_has_failure(parsed):
        raise ReportInputError(f"Runtime smoke JSON có trạng thái thất bại: {path}")
    if parsed is not None and _json_has_success(parsed):
        marker = "JSON success/status marker"
    else:
        candidates = (
            r"\bSMOKE_TEST_PASSED\b",
            r"\bruntime[_ -]?smoke\s*[:=]\s*(?:passed|success|ok)\b",
            r"\bstatus\s*[:=]\s*(?:passed|success|ok|healthy)\b",
        )
        for candidate in candidates:
            found = re.search(candidate, text, re.IGNORECASE)
            if found:
                marker = found.group(0)
                break
    if not marker:
        raise ReportInputError(
            f"Runtime smoke log không có success marker tường minh: {path}. "
            "Dùng SMOKE_TEST_PASSED hoặc JSON status=passed/success/ok."
        )
    return SmokeSummary("passed", marker, path)


def _sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _assert_no_sensitive_evidence(path: Path, text: str) -> None:
    known_secrets = ("Student123!", "AdminLab123!")
    for secret in known_secrets:
        if secret in text:
            raise ReportInputError(f"Evidence chứa demo password plaintext: {path}")
    unsafe_field = re.compile(
        r'(?i)["\'](?:password|secret(?:_key)?|signing_key|fernet_key|raw_token|raw_session_id)["\']\s*:\s*["\'](?!\[?(?:redacted|masked))[^"]{4,}',
    )
    if unsafe_field.search(text):
        raise ReportInputError(f"Evidence có trường nhạy cảm chưa mask: {path}")


def _collect_evidence_family(family: str) -> tuple[EvidenceFile, ...]:
    directory = EVIDENCE_ROOT / family
    if not directory.is_dir():
        raise ReportInputError(f"Thiếu thư mục evidence bắt buộc: {directory}")
    paths = sorted(path for path in directory.rglob("*") if path.is_file())
    if not paths:
        raise ReportInputError(f"Thư mục evidence không có file thật: {directory}")
    records: list[EvidenceFile] = []
    for path in paths:
        raw = path.read_bytes()
        if not raw:
            raise ReportInputError(f"Evidence file rỗng: {path}")
        if len(raw) > MAX_INPUT_FILE_BYTES:
            raise ReportInputError(f"Evidence file quá lớn: {path}")
        try:
            text = raw.decode("utf-8")
        except UnicodeDecodeError as exc:
            raise ReportInputError(f"Evidence phải là UTF-8 text/JSON: {path}") from exc
        _assert_no_sensitive_evidence(path, text)
        parsed: Any | None = None
        if path.suffix.lower() == ".json":
            try:
                parsed = json.loads(text)
            except json.JSONDecodeError as exc:
                raise ReportInputError(f"JSON evidence không hợp lệ: {path}: {exc}") from exc
        records.append(EvidenceFile(family, path, len(raw), _sha256_bytes(raw), text, parsed))
    return tuple(records)


def _flatten_trace_records(value: Any) -> Iterable[Mapping[str, Any]]:
    if isinstance(value, Mapping):
        if isinstance(value.get("steps"), list) and (
            value.get("trace_id") or value.get("mode") or value.get("flow_id") or value.get("scenario")
        ):
            yield value
            return
        for key in ("traces", "records", "items", "flows", "data"):
            if key in value:
                yield from _flatten_trace_records(value[key])
    elif isinstance(value, list):
        for item in value:
            yield from _flatten_trace_records(item)


def _normalize_search(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", value.lower()).strip("_")


def _trace_search_blob(source: EvidenceFile, record: Mapping[str, Any]) -> str:
    serialized = json.dumps(record, ensure_ascii=False, sort_keys=True)
    return _normalize_search(source.relative_path + " " + serialized)


def _validate_trace_steps(flow: FlowSpec, record: Mapping[str, Any], source: Path) -> None:
    steps = record.get("steps")
    if not isinstance(steps, list) or not steps:
        raise ReportInputError(f"Trace {flow.flow_id} không có steps: {source}")
    required = ("layer", "technique", "input_data", "output_data", "status")
    for index, step in enumerate(steps, 1):
        if not isinstance(step, Mapping):
            raise ReportInputError(f"Trace {flow.flow_id} step {index} không phải object: {source}")
        missing = [name for name in required if name not in step]
        if "title" not in step and "action" not in step:
            missing.append("title/action")
        if "code_reference" not in step and "code" not in step:
            missing.append("code_reference")
        if "security_meaning" not in step and "meaning" not in step:
            missing.append("security_meaning")
        if missing:
            raise ReportInputError(
                f"Trace {flow.flow_id} step {index} thiếu {', '.join(missing)}: {source}"
            )


def _resolve_flows(trace_files: Sequence[EvidenceFile]) -> tuple[FlowTrace, ...]:
    candidates: list[tuple[EvidenceFile, Mapping[str, Any], str]] = []
    for evidence_file in trace_files:
        if evidence_file.parsed is None:
            continue
        for record in _flatten_trace_records(evidence_file.parsed):
            candidates.append((evidence_file, record, _trace_search_blob(evidence_file, record)))
    if not candidates:
        raise ReportInputError("Không đọc được TraceRecord nào từ evidence/traces")
    resolved: list[FlowTrace] = []
    missing: list[str] = []
    for spec in FLOW_SPECS:
        selected: tuple[EvidenceFile, Mapping[str, Any], str] | None = None
        best_score = -1
        for candidate in candidates:
            blob = candidate[2]
            for alias in spec.aliases:
                normalized = tuple(_normalize_search(token) for token in alias)
                padded = f"_{blob}_"
                if all(f"_{token}_" in padded for token in normalized):
                    score = sum(len(token) for token in normalized)
                    if score > best_score:
                        selected, best_score = candidate, score
        if selected is None:
            missing.append(spec.flow_id)
            continue
        _validate_trace_steps(spec, selected[1], selected[0].path)
        resolved.append(FlowTrace(spec, selected[0], selected[1]))
    if missing:
        raise ReportInputError(
            "Thiếu trace thật cho các flow: " + ", ".join(missing) + ". "
            "Đặt flow_id/scenario hoặc filename theo các tên này trước khi tạo báo cáo."
        )
    return tuple(resolved)


def collect_inputs() -> ReportInputs:
    tests = parse_pytest_log(EVIDENCE_ROOT / "logs" / "pytest.txt")
    coverage = parse_coverage_log(EVIDENCE_ROOT / "logs" / "coverage.txt")
    smoke = parse_smoke_log(EVIDENCE_ROOT / "logs" / "runtime_smoke.txt")
    evidence = {family: _collect_evidence_family(family) for family in REQUIRED_EVIDENCE_FAMILIES}
    flows = _resolve_flows(evidence["traces"])
    sources: dict[str, str] = {}
    source_hashes: dict[str, str] = {}
    for relative in REQUIRED_SOURCE_FILES:
        path = ROOT / relative
        text = _read_required_text(path)
        sources[relative] = text
        source_hashes[relative] = _sha256_bytes(text.encode("utf-8"))
    return ReportInputs(
        generated_at=datetime.now(timezone.utc).isoformat(),
        tests=tests,
        coverage=coverage,
        smoke=smoke,
        evidence={key: tuple(value) for key, value in evidence.items()},
        flows=flows,
        sources=sources,
        source_hashes=source_hashes,
    )


def _compact_json(value: Any, limit: int = 240) -> str:
    def redact(item: Any, key: str = "") -> Any:
        key_l = key.lower()
        allowed_sensitive = any(token in key_l for token in ("fingerprint", "status", "name", "algorithm"))
        if not allowed_sensitive and any(
            token in key_l for token in ("password", "secret", "token", "session_id", "cookie_value")
        ):
            return "[REDACTED]"
        if isinstance(item, Mapping):
            return {str(k): redact(v, str(k)) for k, v in item.items()}
        if isinstance(item, list):
            return [redact(v, key) for v in item[:20]]
        if isinstance(item, str) and len(item) > 96:
            return item[:24] + "...[MASKED]"
        return item

    text = json.dumps(redact(value), ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    return text if len(text) <= limit else text[: limit - 13] + "...[rút gọn]"


def _trace_table(flow: FlowTrace) -> TableSpec:
    rows: list[tuple[str, ...]] = []
    for index, step in enumerate(flow.record.get("steps", []), 1):
        step_no = str(step.get("step_number", index))
        action = str(step.get("title", step.get("action", "")))
        code = str(step.get("code_reference", step.get("code", "")))
        meaning = str(step.get("security_meaning", step.get("meaning", "")))
        rows.append(
            (
                step_no,
                str(step.get("layer", "")),
                action,
                str(step.get("technique", "")),
                _compact_json(step.get("input_data")),
                _compact_json(step.get("output_data")),
                code,
                str(step.get("status", "")),
                meaning,
            )
        )
    return TableSpec(
        title=f"Trace table - {flow.spec.title} ({flow.source.relative_path})",
        headers=("STT", "Layer", "Action", "Technique", "Input", "Output", "Code reference", "Status", "Security meaning"),
        rows=tuple(rows),
        widths_dxa=(420, 700, 1100, 850, 1100, 1100, 1000, 650, 2440),
        kind="trace_table_dense",
    )


def _comparison_matrix() -> TableSpec:
    rows = (
        ("Dữ liệu nằm ở đâu", "Client", "Client", "Client", "Client token", "State server"),
        ("Người dùng đọc được", "Có", "Có sau decode", "Thường có", "Không nếu key an toàn", "Chỉ opaque ID"),
        ("Người dùng sửa được", "Có", "Có", "Có nhưng bị phát hiện", "Có nhưng decrypt fail", "Token không mang role"),
        ("Server phát hiện sửa", "Không", "Không", "Có", "Có", "Lookup/token state"),
        ("Confidentiality", "Không", "Không", "Không", "Có", "State không ở client"),
        ("Integrity", "Không", "Không", "Có", "Có", "Opaque + server state"),
        ("Revocation", "Kém", "Kém", "Khó nếu stateless", "Khó nếu stateless", "Tức thời server-side"),
        ("Role ở client", "Có", "Có", "Payload demo có", "Không dùng", "Không"),
        ("Authorization source", "Cookie role", "Decoded role", "Database sau verify", "Không authorize", "Database"),
        ("Dữ liệu nhạy cảm", "Không phù hợp", "Không phù hợp", "Không che nội dung", "Có thể nếu cần", "Phù hợp hơn"),
        ("Rủi ro còn lại", "Poisoning", "Poisoning", "Replay/stale role", "Key/replay/stale state", "Token theft/CSRF"),
    )
    return TableSpec(
        "Ma trận so sánh năm mô hình",
        ("Tiêu chí", "Plain", "Base64", "Signed", "Encrypted", "Server session"),
        rows,
        (1860, 1500, 1500, 1500, 1500, 1500),
        "matrix_dense",
    )


def _extract_marked_block(sources: Mapping[str, str], marker: str) -> tuple[str, str]:
    start_token = f"LAB06-CODE:{marker}:START"
    end_token = f"LAB06-CODE:{marker}:END"
    for name, source in sources.items():
        start = source.find(start_token)
        end = source.find(end_token)
        if start >= 0 and end > start:
            block_start = source.find("\n", start) + 1
            block = source[block_start:end].strip()
            if block:
                return name, block
    raise ReportInputError(f"Không tìm thấy source marker thật: {marker}")


def _code_comparison_tables(inputs: ReportInputs) -> list[TableSpec]:
    pairs = (
        ("Plain: client role", "plain_authorization", "Signed: verify before use", "signed_verification"),
        ("Base64: decoded role", "base64_authorization", "Server session: database role", "session_resolution"),
        ("Session rotation", "session_rotation", "Logout invalidation", "logout_invalidation"),
        ("Authenticated decryption", "authenticated_decryption", "Cookie flags", "cookie_flags"),
    )
    tables: list[TableSpec] = []
    for left_label, left_marker, right_label, right_marker in pairs:
        left_file, left = _extract_marked_block(inputs.sources, left_marker)
        right_file, right = _extract_marked_block(inputs.sources, right_marker)
        tables.append(
            TableSpec(
                f"{left_label} ↔ {right_label}",
                (f"{left_label} - {left_file}", f"{right_label} - {right_file}"),
                ((left[:1600], right[:1600]),),
                (4680, 4680),
                "code_comparison",
            )
        )
    return tables


def _audit_rows(inputs: ReportInputs) -> tuple[tuple[str, ...], ...]:
    rows: list[tuple[str, ...]] = []
    for evidence_file in inputs.evidence["audit"]:
        parsed = evidence_file.parsed
        items = parsed if isinstance(parsed, list) else [parsed] if isinstance(parsed, Mapping) else []
        for item in items:
            if not isinstance(item, Mapping):
                continue
            rows.append(
                (
                    str(item.get("timestamp", "")),
                    str(item.get("mode", "")),
                    str(item.get("action", "")),
                    str(item.get("route", "")),
                    str(item.get("authorization_decision", item.get("decision", ""))),
                    str(item.get("reason", "")),
                    str(item.get("trace_id", "")),
                )
            )
    if not rows:
        raise ReportInputError("Audit evidence không có object record để lập bảng")
    return tuple(rows[:40])


def _schema_table() -> TableSpec:
    rows = (
        ("users", "id, username, display_name, email, password_hash, role, active, timestamps", "Role có thẩm quyền"),
        ("server_sessions", "session_token_hash, user_id, expiry, active, revoked_at, rotation_reason", "Không lưu raw token"),
        ("audit_logs", "actor, route, mode, cookie status, submitted/database role, decision, reason, trace_id", "Dữ liệu đã che"),
        ("cookie_events", "operation, fingerprint, signature/encryption status, decision", "Không full cookie"),
        ("session_events", "event_type, old/new fingerprint, reason, trace_id", "Rotation/revocation"),
    )
    return TableSpec("Database schema", ("Bảng", "Trường chính", "Ý nghĩa"), rows, (1700, 5260, 2400))


def _account_table() -> TableSpec:
    rows = (
        ("1", "admin_lab", "Quản trị Lab", "admin@lab.local", "admin", "PBKDF2-SHA256:600000"),
        ("10", "student", "Sinh viên Demo", "student@lab.local", "user", "PBKDF2-SHA256:600000"),
    )
    return TableSpec("Tài khoản demo", ("ID", "Username", "Display name", "Email", "Role", "Password storage"), rows, (600, 1300, 1700, 2100, 900, 2760))


def _test_table(inputs: ReportInputs) -> TableSpec:
    t = inputs.tests
    return TableSpec(
        "Pytest summary parse từ log thật",
        ("Passed", "Failed", "Errors", "Skipped", "Nguồn"),
        ((str(t.passed), str(t.failed), str(t.errors), str(t.skipped), t.source.relative_to(ROOT).as_posix()),),
        (1200, 1200, 1200, 1200, 4560),
    )


def _coverage_table(inputs: ReportInputs) -> TableSpec:
    rows = tuple((module, f"{inputs.coverage.modules[module]}%", "Đạt" if inputs.coverage.modules[module] >= MIN_CORE_COVERAGE else "Không đạt") for module in CORE_MODULES)
    if inputs.coverage.total is not None:
        rows += (("TOTAL", f"{inputs.coverage.total}%", "Thông tin tham khảo"),)
    return TableSpec("Coverage module lõi", ("Module", "Coverage", "Gate"), rows, (5200, 1800, 2360))


def _smoke_table(inputs: ReportInputs) -> TableSpec:
    return TableSpec(
        "Runtime smoke summary parse từ log thật",
        ("Status", "Success marker", "Nguồn"),
        ((inputs.smoke.status, inputs.smoke.marker, inputs.smoke.source.relative_to(ROOT).as_posix()),),
        (1400, 3960, 4000),
    )


def _evidence_inventory(inputs: ReportInputs, families: Sequence[str]) -> TableSpec:
    rows: list[tuple[str, ...]] = []
    for family in families:
        for item in inputs.evidence[family]:
            rows.append((family, item.relative_path, str(item.size), item.sha256[:16] + "…"))
    return TableSpec("Evidence inventory", ("Family", "File", "Bytes", "SHA-256"), tuple(rows), (1100, 5260, 1000, 2000), "inventory")


def _safe_text_excerpt(text: str, limit: int = 450) -> str:
    compact = " ".join(text.split())
    compact = re.sub(
        r"\b[A-Za-z0-9_-]{32,}\b",
        lambda match: match.group(0)[:10] + "...[MASKED]",
        compact,
    )
    return compact if len(compact) <= limit else compact[: limit - 13] + "...[rút gọn]"


def _evidence_excerpt_table(inputs: ReportInputs, families: Sequence[str]) -> TableSpec:
    rows: list[tuple[str, ...]] = []
    for family in families:
        for item in inputs.evidence[family]:
            if item.parsed is not None:
                excerpt = _compact_json(item.parsed, limit=450)
            else:
                excerpt = _safe_text_excerpt(item.text, limit=450)
            rows.append((family, item.relative_path, excerpt))
    return TableSpec(
        "Trích đoạn evidence thật đã mask",
        ("Family", "File", "Trích đoạn"),
        tuple(rows),
        (1100, 2500, 5760),
        "inventory",
    )


def build_chapters(inputs: ReportInputs) -> list[Chapter]:
    chapters = [
        Chapter(index, title, list(BASE_CHAPTER_PARAGRAPHS[index]))
        for index, title in enumerate(CHAPTER_TITLES, 1)
    ]
    by_number = {chapter.number: chapter for chapter in chapters}
    by_number[7].code_blocks.extend((SEQUENCE_DIAGRAM, DATA_FLOW_DIAGRAM))
    by_number[8].tables.extend((_schema_table(), _account_table()))

    flow_map = {flow.spec.flow_id: flow for flow in inputs.flows}
    distribution = {
        9: ("plain_login", "plain_admin_denied"),
        10: ("plain_cookie_modified",),
        11: ("base64_original", "base64_modified"),
        13: ("signed_cookie_valid",),
        14: ("signed_cookie_invalid",),
        15: ("encrypted_cookie_valid",),
        16: ("encrypted_cookie_tampered",),
        17: ("server_session_student_login",),
        19: ("session_rotation",),
        20: ("logout_invalidation", "old_session_rejection"),
        21: ("student_admin_denied", "admin_allowed"),
    }
    for chapter_no, flow_ids in distribution.items():
        by_number[chapter_no].tables.extend(_trace_table(flow_map[flow_id]) for flow_id in flow_ids)
    by_number[26].tables.append(_comparison_matrix())
    by_number[27].tables.extend(_code_comparison_tables(inputs))
    by_number[28].tables.append(
        TableSpec(
            "Audit events từ evidence thật",
            ("Timestamp", "Mode", "Action", "Route", "Decision", "Reason", "Trace ID"),
            _audit_rows(inputs),
            (1300, 900, 1500, 1500, 900, 1960, 1300),
            "audit_dense",
        )
    )
    by_number[29].tables.append(_test_table(inputs))
    by_number[30].tables.append(_coverage_table(inputs))
    by_number[31].tables.append(_smoke_table(inputs))
    by_number[34].qa_answers.extend(QA_ANSWERS)
    return chapters


def _set_run_font(run: Any, name: str = "Calibri", size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _configure_docx_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal_rpr = normal._element.get_or_add_rPr()
    normal_rfonts = normal_rpr.get_or_add_rFonts()
    normal_rfonts.set(qn("w:ascii"), "Calibri")
    normal_rfonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(9.0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.05
    for style_name, size, color, before, after in (
        ("Heading 1", 13.5, BLUE, 10, 5),
        ("Heading 2", 11.5, BLUE, 8, 4),
        ("Heading 3", 10.5, DARK_BLUE, 7, 3),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style_rpr = style._element.get_or_add_rPr()
        style_rfonts = style_rpr.get_or_add_rFonts()
        style_rfonts.set(qn("w:ascii"), "Calibri")
        style_rfonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    code_style = doc.styles.add_style("Lab06 Code", WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = "Consolas"
    code_rpr = code_style._element.get_or_add_rPr()
    code_rfonts = code_rpr.get_or_add_rFonts()
    code_rfonts.set(qn("w:ascii"), "Consolas")
    code_rfonts.set(qn("w:hAnsi"), "Consolas")
    code_style.font.size = Pt(7.2)
    code_style.paragraph_format.space_before = Pt(3)
    code_style.paragraph_format.space_after = Pt(4)
    code_style.paragraph_format.line_spacing = 1.0


def _configure_section(section: Any) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.65)
    section.right_margin = Inches(0.6)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.6)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)


def _add_field(paragraph: Any, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, text, end))
    _set_run_font(run, size=9, color=MUTED)


def _configure_header_footer(section: Any) -> None:
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(7.25), WD_TAB_ALIGNMENT.RIGHT
    )
    left = paragraph.add_run("LAB 6 • COOKIE POISONING")
    _set_run_font(left, size=8.5, color=MUTED, bold=True)
    right = paragraph.add_run("\tBÁO CÁO THỰC HÀNH")
    _set_run_font(right, size=8.5, color=MUTED)
    # Editorial-cover override: deliberately no w:pBdr and no bottom rule.
    ppr = paragraph._p.get_or_add_pPr()
    for border in ppr.findall(qn("w:pBdr")):
        ppr.remove(border)
    for existing_run in list(paragraph.runs):
        paragraph._p.remove(existing_run._r)
    simple_header = paragraph.add_run("LAB 6 - COOKIE POISONING")
    _set_run_font(simple_header, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    label = fp.add_run("Trang ")
    _set_run_font(label, size=9, color=MUTED)
    _add_field(fp, "PAGE")
    of_run = fp.add_run(" / ")
    _set_run_font(of_run, size=9, color=MUTED)
    _add_field(fp, "NUMPAGES")
    for existing_run in list(fp.runs):
        fp._p.remove(existing_run._r)
    page_label = fp.add_run("Trang ")
    _set_run_font(page_label, size=9, color=MUTED)
    _add_field(fp, "PAGE")


def _shade_paragraph(paragraph: Any, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell: Any) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    tc_mar = tcpr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    values = {
        "top": CELL_MARGIN_TOP_DXA,
        "bottom": CELL_MARGIN_BOTTOM_DXA,
        "start": CELL_MARGIN_START_DXA,
        "end": CELL_MARGIN_END_DXA,
    }
    for side, value in values.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_cell_width(cell: Any, width_dxa: int) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    tc_w = tcpr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tcpr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def _set_table_geometry(table: Any, widths_dxa: Sequence[int]) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ReportBuildError(f"Table widths phải bằng {CONTENT_WIDTH_DXA} DXA: {widths_dxa}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), "B8C5D1")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            _set_cell_width(cell, widths_dxa[index])
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _fill_cell(cell: Any, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def _add_docx_table(doc: Document, spec: TableSpec) -> None:
    caption = doc.add_paragraph()
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)
    caption.paragraph_format.keep_with_next = True
    run = caption.add_run(spec.title)
    _set_run_font(run, size=8.5, color=DARK_BLUE, bold=True)
    table = doc.add_table(rows=1, cols=len(spec.headers))
    _set_table_geometry(table, spec.widths_dxa)
    header = table.rows[0]
    header_properties = header._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_properties.append(repeat_header)
    for index, text in enumerate(spec.headers):
        cell = header.cells[index]
        _fill_cell(cell, HEADER_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        _set_run_font(run, size=6.0 if "dense" in spec.kind else 6.8, color=INK, bold=True)
    body_size = 6.0 if spec.kind in {"trace_table_dense", "matrix_dense", "audit_dense"} else 7.0
    if spec.kind == "code_comparison":
        body_size = 6.4
    for row_values in spec.rows:
        if len(row_values) != len(spec.headers):
            raise ReportBuildError(f"Sai số cột trong bảng {spec.title}")
        row = table.add_row()
        for index, text in enumerate(row_values):
            cell = row.cells[index]
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0.5)
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.add_run(str(text))
            font = "Consolas" if spec.kind == "code_comparison" else "Calibri"
            _set_run_font(run, name=font, size=body_size, color=INK)
    # Geometry must be applied after rows exist so every tcW is explicit.
    _set_table_geometry(table, spec.widths_dxa)
    for row_index, row in enumerate(table.rows):
        if row_index > 0 and spec.kind == "code_comparison":
            continue
        row_properties = row._tr.get_or_add_trPr()
        row_properties.append(OxmlElement("w:cantSplit"))
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(2)


def _add_docx_cover(doc: Document, inputs: ReportInputs) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(80)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    _set_run_font(kicker.add_run("BÁO CÁO THỰC HÀNH AN TOÀN ỨNG DỤNG WEB"), size=10.5, color=GOLD, bold=True)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    _set_run_font(title.add_run("LAB 6"), size=30, color=INK, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(4)
    _set_run_font(subtitle.add_run("COOKIE POISONING"), size=21, color=BLUE, bold=True)
    description = doc.add_paragraph()
    description.alignment = WD_ALIGN_PARAGRAPH.CENTER
    description.paragraph_format.space_after = Pt(48)
    _set_run_font(description.add_run("Từ dữ liệu client-controlled đến server-side authorization"), size=13.5, color=DARK_BLUE, italic=True)
    for label, value in (
        ("Sinh viên", "Lê Minh"),
        ("MSSV", "21127645"),
        ("Môi trường", "http://127.0.0.1:5006 - dữ liệu giả lập local"),
        ("Generated from evidence", inputs.generated_at),
    ):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(3)
        _set_run_font(paragraph.add_run(f"{label}: "), size=10.5, color=MUTED, bold=True)
        _set_run_font(paragraph.add_run(value), size=10.5, color=MUTED)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(28)
    note.paragraph_format.space_after = Pt(0)
    _set_run_font(note.add_run("Báo cáo chỉ sử dụng source, log và evidence thật đã vượt qua validation gate."), size=9.5, color=MUTED, italic=True)
    # No bottom rule or paragraph border is used on the editorial cover.
    doc.add_page_break()


def build_docx(inputs: ReportInputs, chapters: Sequence[Chapter], destination: Path) -> None:
    doc = Document()
    _configure_docx_styles(doc)
    for section in doc.sections:
        _configure_section(section)
        _configure_header_footer(section)
    doc.core_properties.title = "LAB 6 - COOKIE POISONING"
    doc.core_properties.subject = "Báo cáo thực hành an toàn ứng dụng web dựa trên evidence thật"
    doc.core_properties.author = "Lê Minh - 21127645"
    doc.core_properties.keywords = "Cookie Poisoning, Flask, signed cookie, Fernet, server-side session"
    doc.core_properties.comments = "Generated by Lab06/scripts/generate_report.py after fail-closed validation."
    _add_docx_cover(doc, inputs)

    for chapter in chapters:
        doc.add_heading(f"Chương {chapter.number}. {chapter.title}", level=1)
        for text in chapter.paragraphs:
            doc.add_paragraph(text)
        for block in chapter.code_blocks:
            paragraph = doc.add_paragraph(style="Lab06 Code")
            _shade_paragraph(paragraph, LIGHT_FILL)
            paragraph.add_run(block)
        for table in chapter.tables:
            _add_docx_table(doc, table)
        for index, (question, answer) in enumerate(chapter.qa_answers, 1):
            doc.add_heading(f"Câu {index}. {question}", level=2)
            doc.add_paragraph(answer)

    doc.add_heading("Phụ lục A. Trace JSON", level=1)
    _add_docx_table(doc, _evidence_inventory(inputs, ("traces",)))
    _add_docx_table(doc, _evidence_excerpt_table(inputs, ("traces",)))
    doc.add_heading("Phụ lục B. Request và Response evidence", level=1)
    _add_docx_table(doc, _evidence_inventory(inputs, ("requests", "responses")))
    _add_docx_table(doc, _evidence_excerpt_table(inputs, ("requests", "responses")))
    doc.add_heading("Phụ lục C. Cookie, Session và Audit evidence", level=1)
    _add_docx_table(doc, _evidence_inventory(inputs, ("cookies", "sessions", "audit")))
    _add_docx_table(doc, _evidence_excerpt_table(inputs, ("cookies", "sessions", "audit")))
    doc.add_heading("Phụ lục D. Source manifest", level=1)
    source_rows = tuple((name, digest[:20] + "…") for name, digest in sorted(inputs.source_hashes.items()))
    _add_docx_table(doc, TableSpec("Source files đã đọc", ("File", "SHA-256"), source_rows, (6760, 2600), "inventory"))
    doc.add_heading("Phụ lục E. Lệnh kiểm chứng", level=1)
    commands = "\n".join(
        (
            "pytest",
            "pytest --cov=. --cov-report=term-missing",
            "python scripts/run_runtime_smoke_test.py",
            "python scripts/export_evidence.py",
            "python scripts/generate_report.py",
        )
    )
    paragraph = doc.add_paragraph(style="Lab06 Code")
    _shade_paragraph(paragraph, LIGHT_FILL)
    paragraph.add_run(commands)
    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destination)


def _find_unicode_fonts() -> tuple[Path, Path]:
    env_font = os.environ.get("LAB06_REPORT_FONT")
    env_bold = os.environ.get("LAB06_REPORT_FONT_BOLD")
    regular_candidates = [
        Path(env_font) if env_font else None,
        Path(r"C:\Windows\Fonts\calibri.ttf"),
        Path(r"C:\Windows\Fonts\arial.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf"),
    ]
    bold_candidates = [
        Path(env_bold) if env_bold else None,
        Path(r"C:\Windows\Fonts\calibrib.ttf"),
        Path(r"C:\Windows\Fonts\arialbd.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
        Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf"),
    ]
    regular = next((path for path in regular_candidates if path and path.is_file()), None)
    bold = next((path for path in bold_candidates if path and path.is_file()), None)
    if not regular or not bold:
        raise ReportBuildError(
            "Không tìm thấy font Unicode regular/bold cho tiếng Việt. "
            "Đặt LAB06_REPORT_FONT và LAB06_REPORT_FONT_BOLD tới file TTF hợp lệ."
        )
    return regular, bold


def _register_pdf_fonts() -> None:
    regular, bold = _find_unicode_fonts()
    pdfmetrics.registerFont(TTFont("LabBody", str(regular)))
    pdfmetrics.registerFont(TTFont("LabBold", str(bold)))


def _pdf_styles() -> Mapping[str, ParagraphStyle]:
    styles = getSampleStyleSheet()
    return {
        "body": ParagraphStyle(
            "LabBody", parent=styles["BodyText"], fontName="LabBody", fontSize=9.4,
            leading=11.2, textColor=colors.HexColor("#" + INK), spaceAfter=4,
        ),
        "h1": ParagraphStyle(
            "LabH1", parent=styles["Heading1"], fontName="LabBold", fontSize=13.5,
            leading=16, textColor=colors.HexColor("#" + BLUE), spaceBefore=10,
            spaceAfter=6, keepWithNext=True,
        ),
        "h2": ParagraphStyle(
            "LabH2", parent=styles["Heading2"], fontName="LabBold", fontSize=11.5,
            leading=13.5, textColor=colors.HexColor("#" + BLUE), spaceBefore=8,
            spaceAfter=4, keepWithNext=True,
        ),
        "caption": ParagraphStyle(
            "LabCaption", parent=styles["BodyText"], fontName="LabBold", fontSize=8.5,
            leading=10, textColor=colors.HexColor("#" + DARK_BLUE), spaceBefore=4,
            spaceAfter=4, keepWithNext=True,
        ),
        "table": ParagraphStyle(
            "LabTable", parent=styles["BodyText"], fontName="LabBody", fontSize=6.2,
            leading=7.1, textColor=colors.HexColor("#" + INK), spaceAfter=0,
        ),
        "table_head": ParagraphStyle(
            "LabTableHead", parent=styles["BodyText"], fontName="LabBold", fontSize=6.2,
            leading=7.1, textColor=colors.HexColor("#" + INK), alignment=TA_CENTER,
        ),
        "code": ParagraphStyle(
            "LabCode", parent=styles["Code"], fontName="LabBody", fontSize=6.8,
            leading=8, textColor=colors.HexColor("#" + INK), leftIndent=5,
            rightIndent=5, spaceBefore=3, spaceAfter=4, backColor=colors.HexColor("#" + LIGHT_FILL),
        ),
        "cover_kicker": ParagraphStyle(
            "CoverKicker", parent=styles["BodyText"], fontName="LabBold", fontSize=10.5,
            leading=13, textColor=colors.HexColor("#" + GOLD), alignment=TA_CENTER,
            spaceAfter=18,
        ),
        "cover_title": ParagraphStyle(
            "CoverTitle", parent=styles["Title"], fontName="LabBold", fontSize=30,
            leading=34, textColor=colors.HexColor("#" + INK), alignment=TA_CENTER,
            spaceAfter=8,
        ),
        "cover_subtitle": ParagraphStyle(
            "CoverSubtitle", parent=styles["BodyText"], fontName="LabBold", fontSize=18,
            leading=22, textColor=colors.HexColor("#" + BLUE), alignment=TA_CENTER,
            spaceAfter=6,
        ),
        "cover_meta": ParagraphStyle(
            "CoverMeta", parent=styles["BodyText"], fontName="LabBody", fontSize=10,
            leading=13, textColor=colors.HexColor("#" + MUTED), alignment=TA_CENTER,
            spaceAfter=3,
        ),
    }


def _xml_escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _pdf_table(spec: TableSpec, styles: Mapping[str, ParagraphStyle]) -> list[Any]:
    header = [Paragraph(_xml_escape(value), styles["table_head"]) for value in spec.headers]
    rows = [header]
    for row in spec.rows:
        rows.append([Paragraph(_xml_escape(str(value)).replace("\n", "<br/>"), styles["table"]) for value in row])
    widths = [width / 1440 * inch for width in spec.widths_dxa]
    table = LongTable(rows, colWidths=widths, repeatRows=1, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#" + HEADER_FILL)),
                ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#" + INK)),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#B8C5D1")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 3.5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3.5),
                ("TOPPADDING", (0, 0), (-1, -1), 2.5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2.5),
            ]
        )
    )
    return [Paragraph(_xml_escape(spec.title), styles["caption"]), table, Spacer(1, 6)]


def _draw_pdf_header_footer(canvas: Any, document: Any) -> None:
    canvas.saveState()
    width, height = letter
    canvas.setFont("LabBody", 8)
    canvas.setFillColor(colors.HexColor("#" + MUTED))
    canvas.drawString(72, height - 35, "LAB 6 • COOKIE POISONING")
    right = "BÁO CÁO THỰC HÀNH"
    canvas.drawRightString(width - 72, height - 35, right)
    # Editorial-cover requirement: no line or bottom border under the header.
    canvas.drawRightString(width - 72, 35, f"Trang {document.page}")
    canvas.restoreState()


def build_pdf(inputs: ReportInputs, chapters: Sequence[Chapter], destination: Path) -> None:
    _register_pdf_fonts()
    styles = _pdf_styles()
    document = SimpleDocTemplate(
        str(destination), pagesize=letter,
        leftMargin=0.6 * inch, rightMargin=0.6 * inch,
        topMargin=0.65 * inch, bottomMargin=0.65 * inch,
        title="LAB 6 - COOKIE POISONING", author="Lê Minh - 21127645",
        subject="Báo cáo evidence-backed về Cookie Poisoning",
    )
    story: list[Any] = [
        Spacer(1, 1.25 * inch),
        Paragraph("BÁO CÁO THỰC HÀNH AN TOÀN ỨNG DỤNG WEB", styles["cover_kicker"]),
        Paragraph("LAB 6", styles["cover_title"]),
        Paragraph("COOKIE POISONING", styles["cover_subtitle"]),
        Spacer(1, 0.25 * inch),
        Paragraph("Từ dữ liệu client-controlled đến server-side authorization", styles["cover_meta"]),
        Spacer(1, 0.55 * inch),
        Paragraph("<b>Sinh viên:</b> Lê Minh", styles["cover_meta"]),
        Paragraph("<b>MSSV:</b> 21127645", styles["cover_meta"]),
        Paragraph("<b>Môi trường:</b> http://127.0.0.1:5006 - dữ liệu giả lập local", styles["cover_meta"]),
        Paragraph(f"<b>Generated from evidence:</b> {_xml_escape(inputs.generated_at)}", styles["cover_meta"]),
        PageBreak(),
    ]
    for chapter in chapters:
        story.append(Paragraph(f"Chương {chapter.number}. {_xml_escape(chapter.title)}", styles["h1"]))
        story.extend(Paragraph(_xml_escape(text), styles["body"]) for text in chapter.paragraphs)
        for block in chapter.code_blocks:
            story.append(Preformatted(block, styles["code"]))
        for table in chapter.tables:
            story.extend(_pdf_table(table, styles))
        for index, (question, answer) in enumerate(chapter.qa_answers, 1):
            story.append(Paragraph(f"Câu {index}. {_xml_escape(question)}", styles["h2"]))
            story.append(Paragraph(_xml_escape(answer), styles["body"]))

    appendix_specs = (
        ("Phụ lục A. Trace JSON", (
            _evidence_inventory(inputs, ("traces",)),
            _evidence_excerpt_table(inputs, ("traces",)),
        )),
        ("Phụ lục B. Request và Response evidence", (
            _evidence_inventory(inputs, ("requests", "responses")),
            _evidence_excerpt_table(inputs, ("requests", "responses")),
        )),
        ("Phụ lục C. Cookie, Session và Audit evidence", (
            _evidence_inventory(inputs, ("cookies", "sessions", "audit")),
            _evidence_excerpt_table(inputs, ("cookies", "sessions", "audit")),
        )),
    )
    for title, tables in appendix_specs:
        story.append(Paragraph(title, styles["h1"]))
        for table in tables:
            story.extend(_pdf_table(table, styles))
    story.append(Paragraph("Phụ lục D. Source manifest", styles["h1"]))
    source_rows = tuple((name, digest[:20] + "…") for name, digest in sorted(inputs.source_hashes.items()))
    story.extend(_pdf_table(TableSpec("Source files đã đọc", ("File", "SHA-256"), source_rows, (6760, 2600), "inventory"), styles))
    story.append(Paragraph("Phụ lục E. Lệnh kiểm chứng", styles["h1"]))
    story.append(
        Preformatted(
            "pytest\npytest --cov=. --cov-report=term-missing\n"
            "python scripts/run_runtime_smoke_test.py\n"
            "python scripts/export_evidence.py\npython scripts/generate_report.py",
            styles["code"],
        )
    )
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.build(story, onFirstPage=_draw_pdf_header_footer, onLaterPages=_draw_pdf_header_footer)


def _count_pdf_pages(path: Path) -> int:
    data = path.read_bytes()
    return len(re.findall(rb"/Type\s*/Page(?!s)\b", data))


def verify_docx(path: Path) -> Mapping[str, int]:
    if not path.is_file() or path.stat().st_size < 15000:
        raise ReportBuildError(f"DOCX thiếu hoặc quá nhỏ: {path}")
    try:
        with zipfile.ZipFile(path) as archive:
            bad = archive.testzip()
            if bad:
                raise ReportBuildError(f"DOCX ZIP lỗi tại entry: {bad}")
            names = set(archive.namelist())
            if "word/document.xml" not in names:
                raise ReportBuildError("DOCX thiếu word/document.xml")
            if any(name.startswith("word/media/") for name in names):
                raise ReportBuildError("Báo cáo text-only không được chứa image/media")
            document_xml = archive.read("word/document.xml")
            if b"placeholder" in document_xml.lower():
                raise ReportBuildError("DOCX chứa từ placeholder")
    except zipfile.BadZipFile as exc:
        raise ReportBuildError(f"DOCX không mở được: {path}") from exc
    try:
        reopened = Document(path)
    except Exception as exc:  # python-docx raises several package/XML exceptions
        raise ReportBuildError(f"python-docx không mở lại được output: {path}") from exc
    headings = [p.text for p in reopened.paragraphs if p.style and p.style.name == "Heading 1"]
    chapters = [text for text in headings if text.startswith("Chương ")]
    if len(chapters) != 35:
        raise ReportBuildError(f"DOCX phải có đúng 35 chương, đọc được {len(chapters)}")
    if len(reopened.tables) < 20:
        raise ReportBuildError(f"DOCX thiếu bảng bắt buộc, đọc được {len(reopened.tables)}")
    return {"bytes": path.stat().st_size, "chapters": len(chapters), "tables": len(reopened.tables)}


def verify_pdf(path: Path) -> Mapping[str, int]:
    if not path.is_file() or path.stat().st_size < 20000:
        raise ReportBuildError(f"PDF thiếu hoặc quá nhỏ: {path}")
    data = path.read_bytes()
    if not data.startswith(b"%PDF-") or b"%%EOF" not in data[-2048:]:
        raise ReportBuildError(f"PDF header/EOF không hợp lệ: {path}")
    pages = _count_pdf_pages(path)
    if not MIN_PDF_PAGES <= pages <= MAX_PDF_PAGES:
        raise ReportBuildError(
            f"PDF có {pages} trang; mục tiêu bắt buộc là {MIN_PDF_PAGES}-{MAX_PDF_PAGES}. "
            "Điều chỉnh nội dung/layout thay vì bỏ qua gate."
        )
    return {"bytes": path.stat().st_size, "pages": pages}


def _input_summary(inputs: ReportInputs) -> Mapping[str, Any]:
    return {
        "pytest": {
            "passed": inputs.tests.passed,
            "failed": inputs.tests.failed,
            "errors": inputs.tests.errors,
            "skipped": inputs.tests.skipped,
        },
        "coverage": {module: inputs.coverage.modules[module] for module in CORE_MODULES},
        "runtime_smoke": inputs.smoke.status,
        "evidence_files": {family: len(files) for family, files in inputs.evidence.items()},
        "trace_flows": [flow.spec.flow_id for flow in inputs.flows],
        "source_files": len(inputs.sources),
    }


def generate_reports(inputs: ReportInputs) -> Mapping[str, Any]:
    chapters = build_chapters(inputs)
    if len(chapters) != 35:
        raise ReportBuildError("Internal chapter plan không có đúng 35 chương")
    if sum(len(chapter.tables) for chapter in chapters if chapter.number in {9, 10, 11, 13, 14, 15, 16, 17, 19, 20, 21}) != 15:
        raise ReportBuildError("Internal trace plan không có đúng 15 trace table")
    REPORT_ROOT.mkdir(parents=True, exist_ok=True)
    final_docx = REPORT_ROOT / DOCX_NAME
    final_pdf = REPORT_ROOT / PDF_NAME
    with tempfile.TemporaryDirectory(prefix="lab06-report-", dir=REPORT_ROOT) as temp_name:
        temp = Path(temp_name)
        temp_docx = temp / DOCX_NAME
        temp_pdf = temp / PDF_NAME
        build_docx(inputs, chapters, temp_docx)
        build_pdf(inputs, chapters, temp_pdf)
        docx_check = verify_docx(temp_docx)
        pdf_check = verify_pdf(temp_pdf)
        os.replace(temp_docx, final_docx)
        os.replace(temp_pdf, final_pdf)
    # Re-open final paths after atomic replacement.
    final_docx_check = verify_docx(final_docx)
    final_pdf_check = verify_pdf(final_pdf)
    return {
        "docx": str(final_docx),
        "pdf": str(final_pdf),
        "docx_verification": final_docx_check,
        "pdf_verification": final_pdf_check,
        "temporary_verification": {"docx": docx_check, "pdf": pdf_check},
        "trace_tables": 15,
        "chapters": 35,
        "answers": len(QA_ANSWERS),
    }


def _parse_args(argv: Sequence[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Fail-closed DOCX/PDF generator for Lab06 Cookie Poisoning"
    )
    parser.add_argument(
        "--check-only",
        action="store_true",
        help="Validate required logs/evidence/source without creating report files",
    )
    return parser.parse_args(argv)


def main(argv: Sequence[str] | None = None) -> int:
    args = _parse_args(argv)
    try:
        inputs = collect_inputs()
        if args.check_only:
            print(json.dumps(_input_summary(inputs), ensure_ascii=False, indent=2, sort_keys=True))
            return 0
        result = generate_reports(inputs)
    except (ReportInputError, ReportBuildError) as exc:
        print(f"REPORT_GENERATION_FAILED: {exc}", file=sys.stderr)
        return 2
    except Exception as exc:
        print(f"REPORT_GENERATION_FAILED_UNEXPECTED: {type(exc).__name__}: {exc}", file=sys.stderr)
        return 3
    print(json.dumps(result, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
