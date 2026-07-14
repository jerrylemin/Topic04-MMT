from pathlib import Path
from docx import Document
from scripts import generate_report as report

def test_compact_report_contract():
    assert len(report.REPORT_SECTIONS)==9
    assert len(report.QA_ANSWERS)==5
    assert len(report.SCREENSHOTS)==9

def test_docx_has_all_sections_and_placeholders(tmp_path:Path,monkeypatch):
    monkeypatch.setattr(report,"SHOTS",tmp_path/"screenshots")
    out=tmp_path/"report.docx";missing=report.build_docx(out);doc=Document(out)
    headings=[p.text for p in doc.paragraphs if p.style and p.style.name=="Heading 1"]
    assert headings==list(report.REPORT_SECTIONS)
    text="\n".join(c.text for t in doc.tables for row in t.rows for c in row.cells)
    assert len(missing)==9 and all(name in text for name in report.EXPECTED_FILES)
    assert "ẢNH 01/09" in text and "ẢNH 09/09" in text

def test_missing_pytest_log_is_truthful(tmp_path:Path,monkeypatch):
    monkeypatch.setattr(report,"ROOT",tmp_path)
    assert "không tuyên bố" in report._pytest_summary()
