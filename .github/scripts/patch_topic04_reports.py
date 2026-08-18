from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[2]
MARKER = "PHỤ LỤC BẰNG CHỨNG HÌNH ẢNH TỐI GIẢN"

LABS = {
    "Lab01": {
        "report": "21127645_LeMinh_21127224_NguyenVuBach_Lab01_XSS.docx",
        "topic": "CROSS-SITE SCRIPTING (XSS)",
        "shots": [
            ("01_reflected_vulnerable.png", "Reflected XSS - bản vulnerable", "Hiển thị URL /vulnerable/search, payload đã nhập, alert hoặc DOM img/onerror và Network request q. Một ảnh phải chứng minh input bị phản chiếu thành HTML thực thi."),
            ("02_reflected_secure.png", "Reflected XSS - bản secure", "Gửi đúng payload ở /secure/search. Ảnh phải thấy payload chỉ hiển thị như text/đã encode và không có alert."),
            ("03_stored_persistence.png", "Stored XSS - lưu và tồn tại sau reload", "Ảnh sau khi đăng comment độc hại rồi reload. Phải thấy comment vẫn tồn tại và browser tiếp tục tạo/thi hành nội dung ở bản vulnerable. Nếu giao diện cho phép, mở Network/Response hoặc Database Inspector trong cùng ảnh."),
            ("04_stored_secure.png", "Stored XSS - bản secure", "Gửi cùng nội dung ở /secure/post/1/comments. Ảnh phải thấy event handler/thẻ nguy hiểm bị loại hoặc encode và không có alert."),
            ("05_dom_vulnerable_secure.png", "DOM-based XSS - so sánh vulnerable và secure", "Xếp vulnerable và secure cạnh nhau trong một screenshot. Bên vulnerable dùng fragment và sinh DOM qua innerHTML. Bên secure hiển thị cùng fragment như textContent. Network không có request mới khi chỉ thay hash."),
        ],
    },
    "Lab02": {
        "report": "21127645_LeMinh_21127224_NguyenVuBach_Lab02_BufferOverflow.docx",
        "topic": "BUFFER OVERFLOW",
        "shots": [
            ("01_boundary_and_overflow.png", "Ranh giới buffer và input gây overflow", "Hiển thị UI/Memory Visualizer với buffer 32 byte, input bình thường hoặc 31 byte và input dài 64 byte. Ảnh phải làm rõ dữ liệu dài hơn capacity."),
            ("02_asan_evidence.png", "ASan phát hiện stack-buffer-overflow", "Chạy vulnerable_asan với input 64 byte. Chụp terminal hoặc ASan Inspector thấy stack-buffer-overflow và exit khác 0. Không cần chụp nhiều trang log."),
            ("03_secure_and_hardening.png", "Bản secure và hardening", "Trong một ảnh, chứng minh secure_length hoặc secure_snprintf từ chối/truncate input dài. Nếu có Hardening Inspector, mở kèm để thấy canary/PIE/NX/RELRO là defense in depth."),
        ],
    },
    "Lab03": {
        "report": "21127645_LeMinh_21127224_NguyenVuBach_Lab03_ParameterTampering.docx",
        "topic": "PARAMETER TAMPERING",
        "shots": [
            ("01_price_tampering.png", "Price tampering - vulnerable và secure", "Chụp request sửa price của product 5 từ 100000 thành 1 và kết quả vulnerable chấp nhận. Trong cùng ảnh hoặc bố cục cạnh nhau, thể hiện secure checkout lấy lại giá 100000 từ server/database."),
            ("02_idor.png", "IDOR - Broken Access Control", "Đăng nhập user_a, đổi invoice từ 1001 sang 1002. Ảnh phải chứng minh vulnerable xem được invoice người khác và secure trả 403 hoặc không render dữ liệu."),
            ("03_role_tampering.png", "Role tampering / mass assignment", "Sửa hidden role=user thành admin ở vulnerable và cho thấy role bị thay đổi. Bản secure phải bỏ qua trường role và giữ user."),
            ("04_audit_evidence.png", "Audit log cho hành vi tampering", "Chụp audit/log của secure route có ít nhất một sự kiện price mismatch, invoice_access_denied hoặc sensitive_field_submitted. Một ảnh log đủ cho phần phát hiện và điều tra."),
        ],
    },
    "Lab04": {
        "report": "21127645_LeMinh_21127224_NguyenVuBach_Lab04_CSRF.docx",
        "topic": "CROSS-SITE REQUEST FORGERY (CSRF)",
        "shots": [
            ("01_vulnerable_csrf.png", "CSRF vulnerable - email bị đổi", "Đăng nhập victim, gửi request từ Demo Page. Ảnh phải thấy origin khác, request POST và email đã đổi thành giá trị demo. Có thể đặt Victim UI và Network cạnh nhau."),
            ("02_secure_rejected.png", "CSRF secure - request ngoài luồng bị từ chối", "Gửi request thiếu hoặc sai token vào secure route. Ảnh phải thấy HTTP 403 và email/state không đổi."),
            ("03_secure_legitimate.png", "CSRF secure - form hợp lệ vẫn hoạt động", "Submit trực tiếp form secure có token hợp lệ. Ảnh phải thấy request thành công và state thay đổi đúng, chứng minh kiểm tra bảo mật không phá chức năng hợp lệ."),
        ],
    },
    "Lab05": {
        "report": "21127645_LeMinh_21127224_NguyenVuBach_Lab05_SQLInjection.docx",
        "topic": "SQL INJECTION",
        "shots": [
            ("01_login_injection.png", "Login SQL Injection - vulnerable và secure", "Dùng scenario cố định admin_lab' --. Ảnh phải cho thấy vulnerable tạo session/bypass điều kiện password, còn secure với cùng input không đăng nhập."),
            ("02_search_injection.png", "Search SQL Injection - result set bị mở rộng", "Dùng payload cố định %' OR 1=1 --. Ảnh phải thấy vulnerable trả nhiều/8 sản phẩm thay vì tập con mong đợi."),
            ("03_parameterized_query.png", "Secure query - parameter binding", "Gửi cùng payload vào secure search/login. Chụp Query Construction hoặc log cho thấy SQL template có placeholder ? và payload nằm ở parameter, kết quả không bị mở rộng/bypass."),
        ],
    },
    "Lab06": {
        "report": "21127645_LeMinh_21127224_NguyenVuBach_Lab06_CookiePoisoning.docx",
        "topic": "COOKIE POISONING",
        "shots": [
            ("01_plain_cookie_tamper.png", "Plain cookie - sửa role để vượt quyền", "Đăng nhập student ở Plain Cookie Demo, sửa lab06_role=user thành admin trong DevTools rồi reload admin route. Ảnh phải thấy cookie đã sửa và quyền truy cập vulnerable."),
            ("02_base64_tamper.png", "Base64 không bảo vệ integrity", "Trong Base64 Demo, giải mã role=user, thay bằng giá trị demo role=admin và reload. Ảnh phải chứng minh encoding đảo ngược được và server vulnerable tin dữ liệu."),
            ("03_signed_cookie_rejected.png", "Signed cookie - tamper bị phát hiện", "Sửa một ký tự signed cookie rồi reload. Ảnh phải thấy server từ chối/invalid signature trước khi dùng payload."),
            ("04_server_side_session.png", "Server-side session - role lấy từ server", "Đăng nhập student rồi admin_lab ở Server-side Session Demo. Ảnh phải cho thấy client chỉ giữ opaque session ID, student bị từ chối admin route và admin hợp lệ được phép. Không chụp full token/secret."),
        ],
    },
}


def set_cell_border(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "dashed")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:color"), "808080")
        borders.append(el)
    tc_pr.append(borders)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F2F2")
    tc_pr.append(shd)


def style_run(run, size=12, bold=False, color=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_p(doc, text="", bold=False, size=12, before=0, after=4, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    style_run(r, size=size, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    style_run(r, size=15 if level == 1 else 13, bold=True, color=(47, 85, 151))
    return p


def add_placeholder(doc, idx, fname, caption, desc):
    add_p(doc, f"Hình {idx}. {caption}", bold=True, size=12, before=6, after=3)
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.3)
    cell = table.cell(0, 0)
    cell.width = Inches(6.3)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_border(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(34)
    r = p.add_run(f"[ CHÈN ẢNH {idx:02d} TẠI ĐÂY ]\n{fname}")
    style_run(r, size=14, bold=True)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(2)
    r = p2.add_run("Ảnh phải thể hiện: ")
    style_run(r, bold=True)
    r = p2.add_run(desc)
    style_run(r)
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(5)
    r = p3.add_run(f"Tên file đề nghị: {fname}")
    style_run(r, bold=True)


def remove_old_appendix(doc):
    body = doc._element.body
    marker_el = None
    for p in doc.paragraphs:
        if MARKER in p.text:
            marker_el = p._element
            break
    if marker_el is None:
        return
    elems = list(body)
    start = elems.index(marker_el)
    sect = body.sectPr
    for el in elems[start:]:
        if el is sect:
            continue
        body.remove(el)


def patch_report(lab, cfg):
    path = ROOT / lab / "report" / cfg["report"]
    doc = Document(path)
    remove_old_appendix(doc)
    doc.add_page_break()
    add_heading(doc, MARKER, 1)
    add_p(doc, f"Phần này là bộ ảnh minh chứng tối thiểu cho {lab} - {cfg['topic']}. Danh sách dưới đây thay cho việc chụp riêng từng tab DevTools hoặc nhiều ảnh lặp cùng một kết quả.", after=5)
    add_p(doc, f"Tổng số ảnh cần chèn: {len(cfg['shots'])}.", bold=True, after=4)
    add_p(doc, "Quy tắc: dùng đúng localhost và scenario có sẵn trong repository, giữ URL/route và kết quả chính, che password, session ID, cookie/token/secret dài. Không cần ảnh cài đặt, terminal khởi động, trang chủ hoặc ảnh trung gian.", after=6)
    for idx, (fname, caption, desc) in enumerate(cfg["shots"], 1):
        add_placeholder(doc, idx, fname, caption, desc)
    add_heading(doc, "Checklist ảnh trước khi nộp", 2)
    for text in [
        f"Đủ {len(cfg['shots'])} ảnh theo danh sách, không có ảnh trùng ý.",
        "Ảnh đọc được URL/route và kết quả chính ở mức zoom bình thường.",
        "Bản vulnerable và secure dùng cùng input/scenario khi cần so sánh.",
        "Không lộ password, session ID, cookie đầy đủ, token đầy đủ hoặc secret key.",
        "Caption mô tả đúng điều ảnh chứng minh.",
    ]:
        add_p(doc, "• " + text, after=1)
    doc.save(path)


def guide_text(lab, cfg):
    lines = [
        f"# HƯỚNG DẪN CHỤP ẢNH TỐI GIẢN - {lab.upper()} {cfg['topic']}",
        "",
        "Mục tiêu là chụp ít ảnh nhất nhưng vẫn đủ bằng chứng cho BaiTapTopic04.docx. Không chụp riêng từng tab DevTools nếu một ảnh tổng hợp đã chứng minh được cùng ý.",
        "",
        f"## Tổng số ảnh đề nghị: {len(cfg['shots'])}",
        "",
        "Quy tắc chung: dùng đúng localhost của lab, dùng scenario/payload có sẵn trong repository, giữ URL/route và kết quả chính trong ảnh, che password, session ID, cookie/token/secret dài. Không chụp bước cài đặt, terminal khởi động, trang chủ hoặc ảnh lặp lại.",
        "",
        "## Danh sách ảnh bắt buộc",
        "",
    ]
    for idx, (fname, caption, desc) in enumerate(cfg["shots"], 1):
        lines += [
            f"### Ảnh {idx:02d}. `{fname}`",
            "",
            f"- Caption: {caption}.",
            f"- Phải thấy: {desc}",
            "- Cách chụp: gom UI và vùng DevTools/terminal liên quan vào cùng khung. Khi cần đối chiếu vulnerable/secure, đặt hai cửa sổ cạnh nhau trong một screenshot nếu đọc được rõ.",
            "- Không cần chụp thêm: request trung gian hoặc tab Headers/Payload/Response riêng nếu không bổ sung bằng chứng mới.",
            "",
        ]
    lines += [
        "## Map sang báo cáo",
        "",
        f"File báo cáo `{cfg['report']}` đã có {len(cfg['shots'])} placeholder tương ứng trong phụ lục cuối báo cáo. Thay placeholder theo đúng thứ tự, giữ caption và phần “Ảnh phải thể hiện”.",
        "",
        "## Tiêu chí đủ",
        "",
        "- Có bằng chứng vulnerable hoạt động đúng kịch bản của đề.",
        "- Có bằng chứng secure chặn hoặc xử lý đúng cùng input khi đề yêu cầu so sánh.",
        "- Có đủ thông tin để giải thích root cause và primary fix mà không cần ảnh bổ sung.",
        "- Không có ảnh trang trí hoặc lặp lại trạng thái đã chứng minh.",
        "",
    ]
    return "\n".join(lines)


def main():
    for lab, cfg in LABS.items():
        patch_report(lab, cfg)
        (ROOT / lab / "HUONG_DAN_CHUP_ANH.md").write_text(guide_text(lab, cfg), encoding="utf-8")
        print(f"patched {lab}: {len(cfg['shots'])} screenshots")


if __name__ == "__main__":
    main()
