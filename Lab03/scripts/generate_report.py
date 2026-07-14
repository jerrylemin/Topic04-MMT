from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image as PILImage
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
from screenshot_manifest import SCREENSHOTS  # noqa: E402

OUT = ROOT / "report"
SHOTS = ROOT / "evidence" / "screenshots"
DOCX_PATH = OUT / "21127645_LeMinh_Lab03_ParameterTampering.docx"
PDF_PATH = OUT / "21127645_LeMinh_Lab03_ParameterTampering.pdf"

CHAPTERS = [
    ("Chương 1. Giới thiệu", [
        "Parameter Tampering là việc sửa dữ liệu do client gửi để làm sai lệch quyết định của server. Lab03 minh họa toàn bộ đường đi từ browser, HTTP request, Flask, xác thực, phân quyền, SQLite, audit log đến response.",
        "Mục tiêu là chứng minh ba lỗi logic: sửa giá sản phẩm, IDOR hóa đơn và mass assignment role. Mỗi lỗi có bản vulnerable để quan sát nguyên nhân và bản secure để kiểm chứng bản vá.",
        "Phạm vi an toàn chỉ gồm 127.0.0.1, máy ảo local hoặc Docker local. Dữ liệu đều giả lập; không có thanh toán, email, tài khoản thật, host tùy ý, proxy tổng quát hay request Internet.",
        "Giới hạn: bản vulnerable chỉ phục vụ giáo dục, không phải mẫu triển khai. Ảnh được chụp thủ công; trace và audit phải sinh từ request thật của lab.",
    ]),
    ("Chương 2. Cơ sở lý thuyết", [
        "HTTP request có thể mang tham số trong query string, form body, header và cookie. Hidden field và localStorage vẫn nằm dưới quyền kiểm soát của trình duyệt nên không tạo ranh giới tin cậy.",
        "Session xác định người dùng sau đăng nhập. Authentication trả lời ai đang yêu cầu; authorization trả lời người đó được làm gì với đối tượng cụ thể. Object-level authorization phải kiểm tra quyền trên từng invoice.",
        "IDOR xảy ra khi server dùng mã đối tượng do client gửi nhưng không kiểm tra quyền sở hữu. Mass assignment xảy ra khi server gán trực tiếp nhiều field từ request, kể cả field nhạy cảm như role.",
        "Server-side validation kiểm tra kiểu, phạm vi và trạng thái; business logic phải lấy dữ liệu có thẩm quyền như giá từ database. Audit logging ghi dấu hiệu tampering nhưng không được ghi password, secret hay toàn bộ cookie.",
    ]),
    ("Chương 3. Parameter Tampering", [
        "Lỗ hổng xuất hiện khi ứng dụng coi dữ liệu client là dữ liệu đáng tin và dùng nó để tính tiền, chọn chủ thể hoặc thay đổi quyền. Các bề mặt thường gặp gồm URL, form, hidden field, cookie, localStorage và header tùy chỉnh.",
        "Client-controlled data gồm product_id, quantity, submitted price, invoice id, email và field thừa. Trusted server data gồm user_id từ session, giá từ products, role từ database và quan hệ owner của invoice.",
        "Ảnh hưởng có thể là sai lệch tài chính, rò rỉ dữ liệu, thay đổi tài khoản khác hoặc nâng quyền. Validate kiểu dữ liệu chưa đủ vì một integer hợp lệ vẫn có thể trỏ đến invoice không thuộc quyền.",
    ]),
    ("Chương 4. Kiến trúc Lab03", [
        "Browser render Jinja, gửi request đến Flask tại 127.0.0.1:5003. Flask đọc session, validate input, gọi business logic và authorization trước khi truy vấn SQLite bằng parameterized query.",
        "Trace service lưu các bước Browser UI -> HTTP Request -> Flask Router -> Authentication -> Input Validation -> Business Logic -> Authorization -> SQLite Query/Write -> Audit Logging -> HTTP Response -> Final Result.",
        "Sequence tổng quát: User -> Browser -> Flask route -> Session identity -> Validation -> Business logic -> Authorization -> SQLite transaction -> Audit/Trace -> Response.",
        "Security headers là defense in depth cho browser; chúng không thay thế kiểm tra giá, ownership hoặc field allowlist.",
    ]),
    ("Chương 5. Dữ liệu và tài khoản mẫu", [
        "User A có id 12 và sở hữu invoice 1001, 1003. User B có id 13 và sở hữu invoice 1002. Admin có id 1 và được phép xem invoice theo chính sách owner-or-admin.",
        "Giá có thẩm quyền nằm ở bảng products: USB Security Key 100000 VND, Wireless Mouse 250000 VND, Mechanical Keyboard 1200000 VND, Lab Laptop 15000000 VND.",
        "Password được hash bằng Werkzeug. Session chỉ lưu user_id, username, role và lab_mode; giao diện chỉ hiển thị cookie đã che.",
    ]),
    ("Chương 6. Thay đổi giá sản phẩm", [
        "Luồng hợp lệ gửi product_id=5, quantity=1, price=100000. Trong bản vulnerable, request.form['price'] được chuyển thành số và dùng trực tiếp để tính total rồi ghi invoice.",
        "Khi client sửa price=1, server vẫn tạo invoice có unit_price=1 và total=1. Nguyên nhân là tin hidden field và không truy vấn lại products.price_vnd.",
        "Tác động là tổn thất tài chính giả lập và dữ liệu invoice sai. Database Inspector chứng minh giá client, giá database và giá đã lưu khác nhau.",
    ]),
    ("Chương 7. Vá lỗi thay đổi giá", [
        "Route secure chỉ dùng product_id và quantity sau validation. Server truy vấn sản phẩm, kiểm tra quantity từ 1 đến 10, kiểm tra stock và tính total bằng price_vnd trong database.",
        "Nếu request vẫn chứa price=1, server đánh dấu là untrusted, bỏ qua giá client, ghi checkout_price_mismatch và tạo invoice 100000 VND trong transaction.",
        "Bản vá hiệu quả vì nguồn giá có thẩm quyền được quyết định phía server. Audit mismatch hỗ trợ phát hiện nhưng không phải cơ chế đảm bảo giá.",
    ]),
    ("Chương 8. IDOR hóa đơn", [
        "User A mở invoice 1001 rồi đổi query id thành 1002. Route vulnerable truy vấn WHERE id = ? mà không ràng buộc owner nên trả invoice giả lập của User B.",
        "invoice_id=1002 là số nguyên hợp lệ, cho thấy input validation không đồng nghĩa authorization. Lỗi thuộc nhóm Broken Access Control trong OWASP Top 10.",
        "Tác động là rò rỉ dữ liệu đối tượng. UI chỉ hiển thị dữ liệu lab cần thiết, nhưng verdict vẫn ghi rõ tài nguyên trái phép đã lộ.",
    ]),
    ("Chương 9. Vá IDOR", [
        "Route secure lấy current_user_id từ session. Với user thường, truy vấn invoice theo cả id và user_id; admin được phép theo chính sách công khai owner-or-admin.",
        "Nếu invoice tồn tại nhưng không thuộc user, server trả HTTP 403 trước khi render nội dung và ghi invoice_access_denied. Response không chứa dòng hàng hay tổng tiền invoice 1002.",
        "Không dùng ID khó đoán làm bảo mật. Chính sách phải được áp dụng trên từng object và từng request.",
    ]),
    ("Chương 10. Thay đổi role", [
        "Form vulnerable render hidden user_id=12 và role=user. Client sửa role=admin; route đọc các field và cập nhật trực tiếp, minh họa mass assignment.",
        "Database đổi role từ user thành admin và session được đồng bộ để chứng minh privilege escalation. Client còn có thể sửa user_id để nhắm tài khoản khác trong bản lỗi.",
        "Đây là lỗi phân quyền, không phải SQL Injection: query vẫn parameterized nhưng dữ liệu hợp lệ được dùng sai mục đích.",
    ]),
    ("Chương 11. Vá role tampering", [
        "Form secure chỉ gửi email. Server lấy user_id từ session, validate email và dùng allowlist duy nhất {'email'}; role, user_id, is_admin và field lạ bị từ chối hoặc bỏ qua.",
        "Khi role=admin được thêm thủ công, server ghi sensitive_field_submitted, giữ role trong database và không đồng bộ session từ client.",
        "Chức năng quản trị role, nếu có, phải là route riêng với authorization admin và re-authentication phù hợp; nó không thuộc self-service profile.",
    ]),
    ("Chương 12. Request Inspector và Parameter Diff", [
        "Inspector hiển thị method, path, query, content type, form body, timestamp và handler. Password, secret key và cookie đầy đủ luôn bị loại bỏ.",
        "Mỗi parameter được phân loại trusted server value, untrusted client value, sensitive, modified, ignored, validated hoặc authorized. Diff đặt original cạnh submitted và nguồn tin cậy tương ứng.",
        "Request Tampering Console chỉ có scenario và route allowlist trong Lab03; người dùng không thể nhập host hoặc URL tùy ý.",
    ]),
    ("Chương 13. Authorization Model", [
        "Mô hình quyết định gồm Subject, Action, Object, Owner, Required permission, Policy, Decision và Reason.",
        "IDOR secure: subject user 12, action read invoice, object invoice 1002, owner user 13, policy owner-or-admin, decision deny.",
        "Profile secure: subject user 12, action update role, object user 12, required permission admin, decision deny vì role không thuộc allowlist self-service.",
    ]),
    ("Chương 14. Audit Logging", [
        "Các event chính gồm checkout_price_mismatch, invoice_access_denied, sensitive_field_submitted, target_user_mismatch, invalid_quantity, invalid_product_id, unknown_parameter và authorization_denied.",
        "Log lưu timestamp, user, route, mode, parameter, original/submitted value, decision, reason và trace_id. Không lưu password, secret hoặc cookie đầy đủ.",
        "Audit hỗ trợ phát hiện cùng một user liên tục sửa giá, quét ID hoặc gửi field nhạy cảm. Log cần giám sát và chính sách cảnh báo; bản thân log không chặn tấn công.",
    ]),
    ("Chương 15. So sánh Parameter Tampering và SQL Injection", [
        "Parameter Tampering sửa giá trị hợp lệ để làm sai business logic hoặc access control. SQL Injection chèn cú pháp để thay đổi cấu trúc câu SQL.",
        "Parameterized SQL ngăn dữ liệu bị diễn giải như mã SQL. Nó không tự ngăn server tin price=1, thiếu ownership check hoặc chấp nhận role=admin.",
        "Bản vá Parameter Tampering cần nguồn dữ liệu server, session identity, object-level authorization, field allowlist và validation theo ngữ cảnh.",
    ]),
    ("Chương 16. Phòng chống", [
        "Không tin dữ liệu client; lấy giá từ database; lấy danh tính từ session; authorization trên từng object; allowlist field; không gửi field nhạy cảm; validate kiểu, phạm vi và stock.",
        "Dùng database transaction, parameterized SQL, audit logging, request size limit, least privilege, kiểm thử hồi quy và re-authentication cho thao tác nhạy cảm.",
        "Rate limiting có thể giảm dò ID hoặc spam log nhưng không thay bản vá logic. CSP, SameSite và CSRF là lớp khác; CSRF token không vá IDOR.",
    ]),
    ("Chương 17. Trả lời câu hỏi báo cáo", [
        "1. Parameter Tampering đổi giá trị request để lừa logic; SQL Injection đưa cú pháp vào input để đổi câu truy vấn. Hai lỗi có nguyên nhân và bản vá khác nhau.",
        "2. Hidden field không phải cơ chế bảo mật vì người dùng kiểm soát DOM và request, có thể sửa bằng DevTools hoặc công cụ gửi request.",
        "3. IDOR thuộc Broken Access Control trong OWASP Top 10; cụ thể là thiếu object-level authorization.",
        "4. Trước khi trả invoice, server phải xác thực session, validate id, tìm object, kiểm tra owner hoặc quyền admin, rồi mới render dữ liệu.",
        "5. Không truyền giá như nguồn quyết định vì client sửa được. Server phải lấy giá hiện hành từ database và tính tổng.",
        "6. user_id trong form không xác định danh tính vì do client kiểm soát; session đã xác thực mới là nguồn danh tính.",
        "7. role không nên xuất hiện trong self-service form vì không cần cho cập nhật email và tạo bề mặt mass assignment.",
        "8. Validate kiểu chưa đủ: invoice_id=1002 và role='admin' đều đúng kiểu nhưng có thể trái quyền.",
        "9. Quyền phải kiểm tra trên từng object vì quyền xem invoice 1001 không suy ra quyền xem invoice 1002.",
        "10. Audit log ghi parameter, giá trị gốc/gửi, quyết định và trace, giúp phát hiện mẫu hành vi bất thường mà không lộ secret.",
        "11. Session được Flask ký và dùng làm ngữ cảnh đã xác thực; hidden field là dữ liệu browser có thể sửa. Dù session có role, server vẫn nên đối chiếu database cho thao tác nhạy cảm.",
    ]),
    ("Chương 18. Kiểm thử", [
        "Pytest bao phủ login, hash password, route cần session, cart, vulnerable/secure checkout, IDOR, role, authorization, validation, audit, trace, security headers, report, screenshot guide và local-only scripts.",
        "Demo flows tạo request, response, trace JSON, audit và database snapshot thật. Kết quả pytest được lưu vào evidence/logs/pytest.txt; báo cáo chỉ đọc log đó, không ghi cứng kết quả.",
    ]),
    ("Chương 19. Kết luận", [
        "Lab chứng minh client không phải vùng tin cậy. Ba bản vulnerable cho thấy cùng một nguyên nhân gốc biểu hiện thành sai giá, rò rỉ object và nâng quyền.",
        "Bản secure đưa quyết định về server: giá database, session identity, object authorization và allowlist. Audit/trace làm quyết định có thể quan sát và kiểm chứng.",
        "Giới hạn là dữ liệu cố định, SQLite single-process và không có hệ thống cảnh báo tập trung. Hướng phát triển an toàn là policy tập trung, re-authentication, rate limiting và giám sát audit trong môi trường triển khai thật.",
    ]),
]

ANSI_ESCAPE = re.compile(r"\x1b\[[0-?]*[ -/]*[@-~]")


def clean_text(value: str) -> str:
    value = ANSI_ESCAPE.sub("", value)
    return "".join(character for character in value if character in "\n\r\t" or ord(character) >= 32)


def read_log(path: Path) -> str:
    data = path.read_bytes()
    encoding = "utf-16" if data.startswith((b"\xff\xfe", b"\xfe\xff")) else "utf-8"
    return clean_text(data.decode(encoding, errors="replace"))


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, text, separate, end])


def image_size(path: Path, max_width: float, max_height: float):
    with PILImage.open(path) as image:
        width, height = image.size
    scale = min(max_width / width, max_height / height)
    return width * scale, height * scale


def style_doc(doc: Document):
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.7)
    section.left_margin = section.right_margin = Inches(0.8)
    for name in ("Normal", "Title", "Heading 1", "Heading 2"):
        doc.styles[name].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.styles["Normal"].paragraph_format.space_after = Pt(6)
    doc.styles["Heading 1"].font.color.rgb = RGBColor(16, 70, 110)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("LAB 3 - PARAMETER TAMPERING | 21127645 | Trang ")
    add_field(footer, "PAGE")


def add_docx_table(doc: Document, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = str(value)
    for row in rows:
        for cell, value in zip(table.add_row().cells, row):
            cell.text = str(value)
    return table


def build_docx(missing: list[str]):
    doc = Document()
    style_doc(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LAB 3\nPARAMETER TAMPERING")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(16, 70, 110)
    for line in ("BÁO CÁO THỰC HÀNH AN TOÀN ỨNG DỤNG WEB", "MSSV: 21127645", "Họ tên: Lê Minh"):
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    doc.add_heading("Mục lục", 0)
    toc = doc.add_paragraph()
    add_field(toc, 'TOC \\o "1-3" \\h \\z \\u')
    doc.add_paragraph("Mở file trong Word và chọn Update Field để cập nhật số trang mục lục.")
    doc.add_page_break()
    for heading, paragraphs in CHAPTERS:
        doc.add_heading(heading, 1)
        for paragraph in paragraphs:
            doc.add_paragraph(paragraph)
        if heading.startswith("Chương 5"):
            doc.add_heading("Bảng 1. Tài khoản mẫu", 2)
            add_docx_table(doc, ["ID", "Username", "Email", "Role"], [[12, "user_a", "usera@lab.local", "user"], [13, "user_b", "userb@lab.local", "user"], [1, "admin", "admin@lab.local", "admin"]])
            doc.add_heading("Bảng 2. Sản phẩm mẫu", 2)
            add_docx_table(doc, ["ID", "Tên", "Giá VND", "Stock"], [[5, "USB Security Key", 100000, 20], [6, "Wireless Mouse", 250000, 15], [7, "Mechanical Keyboard", 1200000, 10], [8, "Lab Laptop", 15000000, 5]])
            doc.add_heading("Bảng 3. Hóa đơn mẫu", 2)
            add_docx_table(doc, ["Invoice", "Owner", "Quan hệ"], [[1001, 12, "User A"], [1002, 13, "User B"], [1003, 12, "User A"]])
        if heading.startswith("Chương 15"):
            add_docx_table(doc, ["Tiêu chí", "Parameter Tampering", "SQL Injection"], [["Mục tiêu", "Làm sai logic/quyền", "Thay đổi câu SQL"], ["Kỹ thuật", "Sửa giá trị hợp lệ", "Chèn cú pháp SQL"], ["Bản vá", "Nguồn server + authorization", "Parameterized query"]])
    doc.add_page_break()
    doc.add_heading("Phụ lục A. Ảnh chụp từng bước", 1)
    for index, item in enumerate(SCREENSHOTS, 1):
        doc.add_heading(f"Hình {index}. {item['filename']}", 2)
        path = SHOTS / item["filename"]
        if path.is_file() and path.stat().st_size:
            width, height = image_size(path, 6.3, 7.6)
            doc.add_picture(str(path), width=Inches(width), height=Inches(height))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            missing.append(item["filename"])
            add_docx_table(doc, [f"ẢNH CẦN BỔ SUNG: {item['filename']}"], [[f"Tài khoản: {item['account']}"], [f"URL: {item['url']}"], [f"Dữ liệu cần sửa: {item['modified']}"], [f"Panel: {item['panel']}"], [f"Phải xuất hiện: {item['required']}"]])
        caption = doc.add_paragraph(f"Hình {index}: {item['caption']}")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    doc.add_heading("Phụ lục B. Lệnh và bằng chứng", 1)
    for command in ("python seed.py", "python scripts/run_demo_flows.py", "python scripts/check_screenshots.py", "python scripts/generate_report.py", "pytest"):
        doc.add_paragraph(command, style="List Bullet")
    log = ROOT / "evidence" / "logs" / "pytest.txt"
    doc.add_heading("Kết quả pytest thực tế", 2)
    doc.add_paragraph(read_log(log) if log.exists() else "Chưa có evidence/logs/pytest.txt.")
    doc.save(DOCX_PATH)


def pdf_font():
    candidates = [Path("C:/Windows/Fonts/arial.ttf"), Path("C:/Windows/Fonts/calibri.ttf")]
    for path in candidates:
        if path.exists():
            pdfmetrics.registerFont(TTFont("LabVN", str(path)))
            return "LabVN"
    return "Helvetica"


def build_pdf(missing: list[str]):
    font = pdf_font()
    styles = getSampleStyleSheet()
    body = ParagraphStyle("LabBody", parent=styles["BodyText"], fontName=font, fontSize=9.5, leading=14, spaceAfter=7)
    h1 = ParagraphStyle("LabH1", parent=styles["Heading1"], fontName=font, fontSize=16, leading=20, textColor=colors.HexColor("#10466e"), spaceBefore=10, spaceAfter=8)
    h2 = ParagraphStyle("LabH2", parent=h1, fontSize=12, leading=15)
    title = ParagraphStyle("LabTitle", parent=styles["Title"], fontName=font, fontSize=26, leading=32, alignment=TA_CENTER, textColor=colors.HexColor("#10466e"))
    story = [Spacer(1, 4 * cm), Paragraph("LAB 3<br/>PARAMETER TAMPERING", title), Spacer(1, 1 * cm), Paragraph("BÁO CÁO THỰC HÀNH AN TOÀN ỨNG DỤNG WEB", ParagraphStyle("Center", parent=body, alignment=TA_CENTER)), Paragraph("MSSV: 21127645<br/>Họ tên: Lê Minh", ParagraphStyle("Center2", parent=body, alignment=TA_CENTER)), PageBreak()]
    story += [Paragraph("Mục lục", h1)] + [Paragraph(heading, body) for heading, _ in CHAPTERS] + [PageBreak()]
    grid = [("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#7d8c97")), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dceaf3")), ("FONTNAME", (0, 0), (-1, -1), font), ("FONTSIZE", (0, 0), (-1, -1), 7.5), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 5), ("RIGHTPADDING", (0, 0), (-1, -1), 5)]
    for heading, paragraphs in CHAPTERS:
        story.append(Paragraph(heading, h1))
        story.extend(Paragraph(paragraph, body) for paragraph in paragraphs)
        if heading.startswith("Chương 5"):
            story.append(Paragraph("Bảng 1. Tài khoản mẫu", h2))
            story.append(Table([["ID", "Username", "Email", "Role"], [12, "user_a", "usera@lab.local", "user"], [13, "user_b", "userb@lab.local", "user"], [1, "admin", "admin@lab.local", "admin"]], colWidths=[1.5 * cm, 3 * cm, 6 * cm, 2.5 * cm], style=TableStyle(grid)))
            story.append(Paragraph("Bảng 2. Sản phẩm mẫu", h2))
            story.append(Table([["ID", "Tên", "Giá VND", "Stock"], [5, "USB Security Key", 100000, 20], [6, "Wireless Mouse", 250000, 15], [7, "Mechanical Keyboard", 1200000, 10], [8, "Lab Laptop", 15000000, 5]], colWidths=[1.5 * cm, 7 * cm, 3 * cm, 2 * cm], style=TableStyle(grid)))
        if heading.startswith("Chương 15"):
            story.append(Table([["Tiêu chí", "Parameter Tampering", "SQL Injection"], ["Mục tiêu", "Sai logic/quyền", "Đổi câu SQL"], ["Kỹ thuật", "Sửa giá trị", "Chèn cú pháp"], ["Bản vá", "Nguồn server + authorization", "Parameterized query"]], colWidths=[3 * cm, 7 * cm, 6 * cm], style=TableStyle(grid)))
    story.extend([PageBreak(), Paragraph("Phụ lục A. Ảnh chụp từng bước", h1)])
    for index, item in enumerate(SCREENSHOTS, 1):
        path = SHOTS / item["filename"]
        story.extend([PageBreak(), Paragraph(f"Hình {index}. {item['filename']}", h2)])
        if path.is_file() and path.stat().st_size:
            width, height = image_size(path, 17 * cm, 21 * cm)
            story.append(Image(str(path), width=width, height=height))
        else:
            if item["filename"] not in missing:
                missing.append(item["filename"])
            rows = [[f"ẢNH CẦN BỔ SUNG: {item['filename']}"], [f"Tài khoản: {item['account']}"], [f"URL: {item['url']}"], [f"Dữ liệu cần sửa: {item['modified']}"], [f"Panel: {item['panel']}"], [f"Phải xuất hiện: {item['required']}"]]
            story.append(Table(rows, colWidths=[17 * cm], style=TableStyle(grid + [("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#fde3e3")), ("BOX", (0, 0), (-1, -1), 1, colors.HexColor("#a52a2a"))])))
        story.append(Paragraph(f"Hình {index}: {item['caption']}", body))
    story.extend([PageBreak(), Paragraph("Phụ lục B. Lệnh và bằng chứng", h1)])
    story.extend(Paragraph(command, body) for command in ("python seed.py", "python scripts/run_demo_flows.py", "python scripts/check_screenshots.py", "python scripts/generate_report.py", "pytest"))
    log = ROOT / "evidence" / "logs" / "pytest.txt"
    story.append(Paragraph("Kết quả pytest thực tế", h2))
    story.append(Paragraph((read_log(log) if log.exists() else "Chưa có evidence/logs/pytest.txt.").replace("\n", "<br/>"), body))

    def footer(canvas, document):
        canvas.saveState()
        canvas.setFont(font, 8)
        canvas.drawCentredString(A4[0] / 2, 0.8 * cm, f"LAB 3 - PARAMETER TAMPERING | 21127645 | Trang {document.page}")
        canvas.restoreState()

    SimpleDocTemplate(str(PDF_PATH), pagesize=A4, leftMargin=1.7 * cm, rightMargin=1.7 * cm, topMargin=1.5 * cm, bottomMargin=1.5 * cm, title="LAB 3 - Parameter Tampering", author="Lê Minh - 21127645").build(story, onFirstPage=footer, onLaterPages=footer)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    missing: list[str] = []
    build_docx(missing)
    build_pdf(missing)
    print(f"Created: {DOCX_PATH}")
    print(f"Created: {PDF_PATH}")
    print(f"Missing screenshots ({len(missing)}): {', '.join(missing) if missing else 'none'}")


if __name__ == "__main__":
    main()
