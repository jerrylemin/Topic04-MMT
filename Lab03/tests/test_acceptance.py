import os
import subprocess
import sys
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document

from screenshot_manifest import SCREENSHOTS


ROOT = Path(__file__).parents[1]

REQUIRED_FILES = [
    "README.md", "HUONG_DAN_CHUP_ANH.md", "requirements.txt", ".env.example", ".dockerignore", "Dockerfile", "docker-compose.yml",
    "app.py", "config.py", "database.py", "auth.py", "services.py", "validators.py", "authorization.py",
    "audit_service.py", "trace_models.py", "trace_service.py", "security_utils.py", "schema.sql", "seed.py",
    "scripts/run_lab.sh", "scripts/run_lab.ps1", "scripts/run_lab.bat", "scripts/reset_database.py",
    "scripts/send_request.py", "scripts/run_demo_flows.py", "scripts/export_evidence.py", "scripts/generate_report.py",
    "scripts/check_screenshots.py", "scripts/clean_generated_files.py",
    "templates/base.html", "templates/index.html", "templates/login.html", "templates/products.html", "templates/cart.html",
    "templates/comparison.html", "templates/security_controls.html", "templates/audit_logs.html", "templates/error.html",
    "templates/vulnerable/checkout.html", "templates/vulnerable/checkout_result.html", "templates/vulnerable/invoice.html",
    "templates/vulnerable/profile.html", "templates/vulnerable/profile_result.html", "templates/secure/checkout.html",
    "templates/secure/checkout_result.html", "templates/secure/invoice.html", "templates/secure/profile.html",
    "templates/secure/profile_result.html", "templates/components/trace_panel.html",
    "templates/components/request_inspector.html", "templates/components/session_inspector.html",
    "templates/components/parameter_inspector.html", "templates/components/database_inspector.html",
    "templates/components/authorization_inspector.html", "templates/components/code_comparison.html",
    "templates/components/security_controls.html", "templates/components/final_verdict.html",
    "templates/components/audit_inspector.html", "templates/components/presentation_controls.html",
    "static/css/style.css", "static/css/presentation.css", "static/js/main.js", "static/js/trace-ui.js",
    "static/js/request-editor.js", "static/js/presentation.js", "static/js/parameter-diff.js",
]


@pytest.mark.parametrize("relative_path", REQUIRED_FILES)
def test_required_file_exists(relative_path):
    assert (ROOT / relative_path).is_file(), relative_path


def test_all_templates_parse(app):
    for path in (ROOT / "templates").rglob("*.html"):
        app.jinja_env.get_template(path.relative_to(ROOT / "templates").as_posix())


def test_frontend_has_inspectors_timeline_and_presentation():
    source = "\n".join(path.read_text(encoding="utf-8") for path in (ROOT / "templates").rglob("*.html"))
    for label in ["Request Tampering Console", "Request Inspector", "Session Inspector", "Database Inspector",
                  "Authorization Inspector", "Final Security Verdict", "Action Timeline", "Presentation Mode"]:
        assert label in source


def test_security_control_panel_reports_cookie_and_csrf_state(client):
    text = client.get("/security-controls").text
    assert "Secure cookie" in text and "CSRF protection" in text
    assert "TẮT" in text


def test_all_primary_get_routes_render(client, login):
    login()
    paths = ["/", "/login", "/products", "/products/5", "/cart", "/vulnerable/checkout", "/secure/checkout",
             "/vulnerable/invoice?id=1001", "/secure/invoice?id=1001", "/vulnerable/profile", "/secure/profile",
             "/comparison", "/security-controls", "/audit-logs", "/health"]
    assert all(client.get(path).status_code == 200 for path in paths)
    assert client.get("/secure/invoice?id=1002").status_code == 403


def test_request_editor_uses_fixed_local_routes_only():
    source = (ROOT / "static/js/request-editor.js").read_text(encoding="utf-8")
    assert "FIXED_ROUTES" in source and "fetch(target" in source
    assert "http://" not in source and "https://" not in source
    assert "LOCAL_HOSTS" in source and "location.hostname" in source


def test_client_and_demo_scripts_are_local_only():
    send = (ROOT / "scripts/send_request.py").read_text(encoding="utf-8")
    demo = (ROOT / "scripts/run_demo_flows.py").read_text(encoding="utf-8")
    assert 'BASE_URL = "http://127.0.0.1:5003"' in send
    assert "--host" not in send and "--url" not in send
    assert "import requests" not in demo and ".test_client()" in demo


def test_runtime_has_no_external_url_or_browser_automation():
    paths = [ROOT / "app.py", *list((ROOT / "scripts").glob("*.py")), *list((ROOT / "static/js").glob("*.js"))]
    source = "\n".join(path.read_text(encoding="utf-8") for path in paths).lower()
    urls = [token.split('"', 1)[0].split("'", 1)[0] for token in source.split("http://")[1:]]
    assert all(url.startswith("127.0.0.1:5003") for url in urls)
    assert "https://" not in source and "playwright" not in source and "selenium" not in source


def test_requirements_are_minimal_and_have_no_browser_driver():
    requirements = (ROOT / "requirements.txt").read_text(encoding="utf-8").lower()
    for dependency in ["flask", "pytest", "requests", "python-docx", "reportlab"]:
        assert dependency in requirements
    assert "playwright" not in requirements and "selenium" not in requirements


def test_docker_is_local_non_root_and_unprivileged():
    dockerfile = (ROOT / "Dockerfile").read_text(encoding="utf-8")
    compose = (ROOT / "docker-compose.yml").read_text(encoding="utf-8")
    assert "USER lab" in dockerfile and "HEALTHCHECK" in dockerfile
    assert '127.0.0.1:5003:5003' in compose
    assert "no-new-privileges:true" in compose and "cap_drop" in compose
    assert "network_mode" not in compose and "privileged:" not in compose
    config = (ROOT / "config.py").read_text(encoding="utf-8")
    assert "DATABASE_PATH" in config and "secrets.token_hex" in config
    assert "SECRET_KEY:" not in compose
    assert ".venv" in (ROOT / ".dockerignore").read_text(encoding="utf-8")


def test_no_real_payment_or_email_integration():
    source = "\n".join(path.read_text(encoding="utf-8").lower() for path in ROOT.glob("*.py"))
    for integration in ["stripe", "paypal", "smtplib", "sendgrid", "mailgun"]:
        assert integration not in source


def test_screenshot_guide_lists_every_required_image():
    guide = (ROOT / "HUONG_DAN_CHUP_ANH.md").read_text(encoding="utf-8")
    assert len(SCREENSHOTS) == 8
    assert all(item["filename"] in guide for item in SCREENSHOTS)
    for field in ["Tên file", "Mục đích", "Trạng thái ban đầu", "URL hoặc lệnh", "Dữ liệu cần nhập",
                  "Nút cần bấm", "Tab DevTools hoặc inspector cần mở", "Nội dung bắt buộc phải xuất hiện",
                  "Kết quả đúng", "Caption dùng trong báo cáo"]:
        assert field in guide
    assert all(f"Bước {number}." in guide for number in range(1, 7))


def test_screenshot_checker_runs_without_ocr_or_generation():
    source = (ROOT / "scripts/check_screenshots.py").read_text(encoding="utf-8").lower()
    assert "ocr" not in source and "pytesseract" not in source
    env = {**os.environ, "PYTHONIOENCODING": "cp1252"}
    result = subprocess.run([sys.executable, "scripts/check_screenshots.py"], cwd=ROOT, text=True,
                            encoding="utf-8", errors="replace", capture_output=True, env=env)
    assert result.returncode in (0, 1)
    assert "OK: 8 valid screenshots" in result.stdout or "MISSING:" in result.stdout
    listed = subprocess.run([sys.executable, "scripts/check_screenshots.py", "--list-required"], cwd=ROOT,
                            text=True, encoding="utf-8", capture_output=True, check=True)
    assert [item["filename"] for item in SCREENSHOTS] == [line.split(" - ", 1)[0][4:] for line in listed.stdout.splitlines()]


def test_report_artifacts_exist_and_open():
    docx = ROOT / "report/21127645_LeMinh_Lab03_ParameterTampering.docx"
    pdf = ROOT / "report/21127645_LeMinh_Lab03_ParameterTampering.pdf"
    assert docx.stat().st_size > 20_000 and pdf.stat().st_size > 20_000
    with ZipFile(docx) as archive:
        assert "word/document.xml" in archive.namelist()
    assert pdf.read_bytes().startswith(b"%PDF")


def test_report_has_all_chapters_questions_and_placeholders():
    doc = Document(ROOT / "report/21127645_LeMinh_Lab03_ParameterTampering.docx")
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert all(heading in text for heading in ["1. Mục tiêu và môi trường thực hành", "2. Kịch bản và các bước thực hiện",
                                               "3. Nguyên nhân kỹ thuật", "4. Kết quả và bằng chứng", "5. Mức độ ảnh hưởng",
                                               "6. Bản vá và cách phòng chống", "7. Trả lời các câu hỏi báo cáo",
                                               "8. Kết quả kiểm thử", "9. Kết luận"])
    for answer in ["Parameter Tampering sửa giá trị request", "Hidden field không phải cơ chế bảo mật",
                   "IDOR thuộc Broken Access Control", "Trước khi trả invoice", "Không nên truyền giá sản phẩm từ client"]:
        assert answer in text
    assert len(doc.tables) >= 10
    assert sum("Chèn ảnh tại vị trí này." in cell.text for table in doc.tables for row in table.rows for cell in row.cells) == 8


def test_report_generator_mentions_real_log_and_missing_images():
    source = (ROOT / "scripts/generate_report.py").read_text(encoding="utf-8")
    assert "evidence\" / \"logs\" / \"pytest.txt" in source
    assert "Chèn ảnh tại vị trí này." in source and "image_size" in source
    assert "missing.append" in source


def test_readme_documents_all_operator_commands():
    readme = (ROOT / "README.md").read_text(encoding="utf-8")
    for command in ["python scripts/reset_database.py", "python scripts/run_demo_flows.py",
                    "python scripts/check_screenshots.py", "python scripts/generate_report.py", "pytest"]:
        assert command in readme
    for topic in ["Request Tampering Console", "Presentation Mode", "Authorization Inspector", "Database Inspector"]:
        assert topic in readme
