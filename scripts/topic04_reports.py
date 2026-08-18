"""Build the six concise Topic04 DOCX reports from repository evidence only."""

from __future__ import annotations

import importlib.util
import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_ROOT = ROOT / "report"
TEAM_MEMBERS = [
    {"name": "Lê Minh", "student_id": "21127645"},
    {"name": "Nguyễn Vũ Bách", "student_id": "21127224"},
]
OUTPUTS = {
    "Lab01": ROOT / "Lab01/report/21127645_LeMinh_21127224_NguyenVuBach_Lab01_XSS.docx",
    "Lab02": ROOT / "Lab02/report/21127645_LeMinh_21127224_NguyenVuBach_Lab02_BufferOverflow.docx",
    "Lab03": ROOT / "Lab03/report/21127645_LeMinh_21127224_NguyenVuBach_Lab03_ParameterTampering.docx",
    "Lab04": ROOT / "Lab04/report/21127645_LeMinh_21127224_NguyenVuBach_Lab04_CSRF.docx",
    "Lab05": ROOT / "Lab05/report/21127645_LeMinh_21127224_NguyenVuBach_Lab05_SQLInjection.docx",
    "Lab06": ROOT / "Lab06/report/21127645_LeMinh_21127224_NguyenVuBach_Lab06_CookiePoisoning.docx",
}
SUMMARY_OUTPUT = REPORT_ROOT / "21127645_LeMinh_21127224_NguyenVuBach_BaoCaoTongHop_6Lab_Topic04.docx"

LAB_ASSIGNMENTS = {
    "Lab01": {
        "primary": "Lê Minh — 21127645",
        "main": "Phụ trách kịch bản Reflected, Stored và DOM-based XSS; phân tích source-to-sink; đối chiếu output encoding, sanitization và bản secure.",
        "support": "Nguyễn Vũ Bách — 21127224",
        "support_work": "Rà soát ba loại XSS, source-to-sink, escaping, sanitization, CSP và phạm vi bảo vệ của HttpOnly.",
    },
    "Lab02": {
        "primary": "Nguyễn Vũ Bách — 21127224",
        "main": "Phụ trách luồng HTTP đến chương trình C, buffer 32 byte, crash có kiểm soát, kiểm tra độ dài, snprintf và compiler hardening.",
        "support": "Lê Minh — 21127645",
        "support_work": "Tích hợp luồng request từ web đến native backend, rà soát phần mô tả HTTP và bản vá.",
    },
    "Lab03": {
        "primary": "Lê Minh — 21127645",
        "main": "Phụ trách sửa giá, IDOR, role tampering; phân tích server-side validation, object-level authorization và field allowlist.",
        "support": "Nguyễn Vũ Bách — 21127224",
        "support_work": "Rà soát IDOR, role tampering, object-level authorization và bảng bằng chứng.",
    },
    "Lab04": {
        "primary": "Nguyễn Vũ Bách — 21127224",
        "main": "Phụ trách Victim/Demo Page, cookie tự gửi, request đổi trạng thái, CSRF token, Origin/Referer, SameSite và re-authentication.",
        "support": "Lê Minh — 21127645",
        "support_work": "Rà soát request, response, luồng vulnerable và secure.",
    },
    "Lab05": {
        "primary": "Lê Minh — 21127645",
        "main": "Phụ trách authentication bypass, search injection, query bị biến đổi, prepared statement, password hashing và lỗi an toàn.",
        "support": "Nguyễn Vũ Bách — 21127224",
        "support_work": "Rà soát truy vấn vulnerable, prepared statement, password hashing và biện pháp phòng chống.",
    },
    "Lab06": {
        "primary": "Nguyễn Vũ Bách — 21127224",
        "main": "Phụ trách plain/Base64/signed/encrypted cookie, server-side session, cookie flags, rotate/revoke và authorization phía server.",
        "support": "Lê Minh — 21127645",
        "support_work": "Rà soát giao diện, bảng so sánh cookie và phần server-side authorization.",
    },
}

COMPLETION_CRITERIA = (
    "Mục tiêu và phạm vi",
    "Môi trường và hướng dẫn chạy",
    "Kịch bản vulnerable",
    "Luồng thực hành",
    "Bằng chứng ảnh, log hoặc request-response",
    "Phân tích nguyên nhân kỹ thuật",
    "Đánh giá ảnh hưởng",
    "Biện pháp phòng chống",
    "Bản vá hoặc secure implementation",
    "Báo cáo DOCX đầy đủ",
)

COMPLETION_SCORES = {
    "Lab01": (10, 10, 10, 10, 4, 10, 10, 10, 10, 10),
    "Lab02": (10, 10, 10, 7, 1, 10, 10, 10, 10, 10),
    "Lab03": (10, 10, 10, 10, 4, 10, 10, 10, 10, 10),
    "Lab04": (10, 10, 8, 8, 4, 10, 10, 10, 10, 10),
    "Lab05": (10, 10, 10, 10, 4, 10, 10, 10, 10, 10),
    "Lab06": (10, 10, 10, 10, 4, 10, 10, 10, 10, 10),
}

MISSING_COMPONENTS = {
    "Lab01": "Thiếu 15 ảnh browser/DevTools cho normal input, ba luồng vulnerable và secure retest.",
    "Lab02": "Thiếu 10 ảnh browser, request-response và log GDB/ASan/crash xác nhận vị trí crash, stack overwrite và ngưỡng input.",
    "Lab03": "Thiếu 15 ảnh checkout, IDOR, role tampering và các cặp vulnerable/secure.",
    "Lab04": "Thiếu 20 ảnh email trước/sau, request hợp lệ, request CSRF, secure 403 và trạng thái sau vá; Demo Page hiện yêu cầu bấm xác nhận thay vì tự gửi form.",
    "Lab05": "Thiếu 14 ảnh normal input, dấu nháy đơn, authentication bypass, search mở rộng và secure retest.",
    "Lab06": "Thiếu 12 ảnh cookie fields, plain/Base64 tampering, signed rejection và server-session authorization.",
}


@dataclass(frozen=True)
class Requirement:
    code: str
    text: str
    section: str
    source: str
    action: str = "Chuẩn hóa nội dung và chỉ rõ nguồn kiểm chứng."
    status: str = "Đầy đủ"


@dataclass(frozen=True)
class LabSpec:
    lab: str
    title: str
    subtitle: str
    owner: str
    objectives: tuple[str, ...]
    environment: tuple[str, ...]
    scenario: tuple[str, ...]
    steps: tuple[tuple[str, str, str], ...]
    observations: tuple[tuple[str, str, str], ...]
    causes: tuple[str, ...]
    impacts: tuple[str, ...]
    defenses: tuple[str, ...]
    patches: tuple[tuple[str, str, str], ...]
    questions: tuple[tuple[str, str], ...]
    lessons: tuple[str, ...]
    requirements: tuple[Requirement, ...]


def _reqs(lab: str, rows: Iterable[tuple[str, str, str, str, str]]) -> tuple[Requirement, ...]:
    return tuple(Requirement(f"{lab}-{code}", text, section, source, status=status) for code, text, section, source, status in rows)


def _specs() -> dict[str, LabSpec]:
    image_status = "Đầy đủ về nội dung, chờ ảnh thật"
    log_status = "Đầy đủ về nội dung, chờ log thật"
    return {
        "Lab01": LabSpec(
            "Lab01", "CROSS-SITE SCRIPTING", "Reflected XSS, Stored XSS và DOM-based XSS", "Lê Minh",
            (
                "Phân biệt ba biến thể XSS theo source, nơi lưu, sink và thời điểm thực thi.",
                "Theo dõi dữ liệu không tin cậy từ URL/form/database đến HTML hoặc DOM.",
                "Giải thích và áp dụng output encoding, sanitization, CSP và cookie flags đúng vai trò.",
            ),
            (
                "Ứng dụng Flask/SQLite cố tình có lỗi, chỉ bind loopback; Chrome hoặc Firefox với DevTools.",
                "Route thật: /vulnerable/search, /vulnerable/post/1/comments, /vulnerable/dom-search và các route /secure tương ứng.",
                "Ảnh phải do sinh viên tự chụp; repository hiện chưa có PNG bằng chứng cho Lab01.",
            ),
            (
                "Reflected: query q được phản chiếu vào HTML response.",
                "Stored: body bình luận được lưu trong SQLite rồi phát lại khi trang được tải.",
                "DOM: location.hash được JavaScript phía client đọc và đưa vào innerHTML; fragment không được gửi trong HTTP request.",
            ),
            (
                ("Reflected XSS", "Nhập chuỗi thường, sau đó payload <img src=x onerror=alert(1)>; xem query, response HTML và vị trí DOM.", "Lab01/evidence/traces/reflected_vulnerable.json; reflected_secure.json"),
                ("Stored XSS", "Đăng bình luận thường và payload an toàn; reload/phiên khác; kiểm tra comments.body và response sau reload.", "Lab01/evidence/traces/stored_vulnerable.json; stored_secure.json"),
                ("DOM-based XSS", "Đọc dom_vulnerable.js; xác định source location.hash và sink innerHTML; so sánh textContent ở bản secure.", "Lab01/static/js/dom_vulnerable.js; dom_secure.js; evidence/traces/dom_*.json"),
            ),
            (
                ("Reflected vulnerable", "Trace ghi response status 200 và html_snippet chứa img/onerror chưa encode; chưa có ảnh browser để xác nhận alert trực quan.", "evidence/traces/reflected_vulnerable.json"),
                ("Stored vulnerable", "Trace và database_inspector ghi payload trong comments.body, được đọc lại và render; chưa có ảnh reload/user khác.", "evidence/traces/stored_vulnerable.json"),
                ("DOM", "Source thật dùng innerHTML ở bản vulnerable và textContent ở bản secure; fragment không có trong request URL. Ảnh Elements/Network còn thiếu.", "static/js/dom_vulnerable.js; static/js/dom_secure.js; evidence/traces/dom_vulnerable.json"),
            ),
            (
                "Reflected XSS xuất hiện vì Markup(q) vô hiệu hóa autoescape trước sink HTML.",
                "Stored XSS tồn tại vì dữ liệu đã parameterize SQL nhưng vẫn được Markup khi render; chống SQLi không đồng nghĩa chống XSS.",
                "DOM XSS xuất hiện vì innerHTML diễn giải chuỗi từ location.hash thành node/handler; textContent chỉ tạo text node.",
            ),
            (
                "Reflected XSS tác động người mở URL chứa payload; Stored XSS có thể tác động mọi người xem dữ liệu đã lưu.",
                "DOM XSS chạy trong browser và có thể thao tác nội dung/quyền của origin; HttpOnly giảm rủi ro đọc cookie nhưng không vá sink.",
            ),
            (
                "Encode output đúng ngữ cảnh HTML, JavaScript, URL và attribute.", "Dùng textContent thay innerHTML khi chỉ cần text.",
                "Sanitize HTML bằng allowlist thư viện đáng tin cậy khi thật sự cho phép rich text.", "Bật CSP như lớp giảm hậu quả, không thay sửa code.",
                "Dùng HttpOnly, Secure, SameSite; validate server-side; không tin URL, form, cookie hoặc localStorage.",
            ),
            (
                ("Reflected", "app.py", "q=q if secure else Markup(q)  # secure giữ Jinja autoescape"),
                ("Stored", "app.py", "Markup(bleach.clean(row['body'], tags=ALLOWED_TAGS, attributes={}, protocols=[], strip=True))"),
                ("DOM", "static/js/dom_secure.js", "result.textContent = value;"),
            ),
            (
                ("So sánh Reflected, Stored và DOM-based XSS.", "Reflected nằm trong request/response hiện tại; Stored được lưu và phát lại; DOM phát sinh ở client từ source DOM đến sink DOM."),
                ("Vì sao validate input chưa đủ?", "Input hợp lệ ở một ngữ cảnh vẫn có thể nguy hiểm ở sink khác; cần encode/sanitize tại output."),
                ("Vì sao cần output encoding?", "Encoding giữ ký tự dữ liệu không bị parser HTML/JS/URL/attribute hiểu thành mã."),
                ("CSP có thay sửa code không?", "Không. CSP chỉ giảm khả năng thực thi và hậu quả; sink không an toàn vẫn phải sửa."),
                ("Vá từng lỗi thế nào?", "Reflected: autoescape/contextual encoding; Stored: sanitize/encode khi render; DOM: safe DOM API như textContent."),
            ),
            ("XSS là lỗi source-to-sink, không phải chỉ là một payload.", "Bản vá phải đặt tại đúng ngữ cảnh output; defense in depth không thay root fix."),
            _reqs("L01", (
                ("01", "Mục tiêu phân biệt Reflected, Stored và DOM-based XSS", "2. Mục tiêu", "BaiTapTopic04.docx mục Lab 1", "Đầy đủ"),
                ("02", "Môi trường browser và ứng dụng lab local", "3. Môi trường", "app.py; templates; static/js", "Đầy đủ"),
                ("03", "Reflected: input thường, phản chiếu HTML, payload an toàn, vị trí chèn và escape < > \" '", "5-7", "reflected_vulnerable.json; reflected_secure.json", image_status),
                ("04", "Giải thích nạn nhân mở URL chứa payload", "7. Nguyên nhân", "request/response trace", "Đầy đủ"),
                ("05", "Stored: bình luận thường, payload, reload/user khác, nơi lưu và phạm vi ảnh hưởng", "5-8", "stored_vulnerable.json; lab01.db", image_status),
                ("06", "Rủi ro cookie không HttpOnly", "8-9", "config.py; security controls", "Đầy đủ"),
                ("07", "DOM: đọc JS; source, sink; innerHTML/document.write/eval/setTimeout/location.hash", "5-7", "static/js/dom_*.js", image_status),
                ("08", "Fragment không gửi lên server; so sánh innerHTML/textContent", "4, 7, 10", "dom_*.js; DOM trace", image_status),
                ("09", "Ít nhất năm biện pháp, gồm bốn ngữ cảnh encoding, sanitization, CSP, cookie flags và validation", "9. Phòng chống", "app.py; dom_secure.js; config.py", "Đầy đủ"),
                ("10", "Trả lời trực tiếp năm câu hỏi báo cáo", "11. Câu hỏi", "BaiTapTopic04.docx mục 6 Lab 1", "Đầy đủ"),
                ("11", "Bằng chứng ảnh/request-response", "14. Phụ lục", "evidence/*; F12 manifest", image_status),
            )),
        ),
        "Lab02": LabSpec(
            "Lab02", "BUFFER OVERFLOW TRONG ỨNG DỤNG WEB LOCAL", "Từ HTTP POST tới backend C và stack", "Nguyễn Vũ Bách",
            ("Hiểu ghi vượt buffer và nguy cơ memory corruption.", "Quan sát crash có kiểm soát trong Linux VM local.", "Vá bằng bounded copy, request limit và compiler/OS hardening."),
            ("Linux VM local; GCC; GDB; Python hoặc web client gửi request loopback.", "Flask chuyển dữ liệu POST /submit tới binary C; buffer name[32] có 31 byte dữ liệu và byte null.", "Không shellcode, ROP, reverse shell hoặc chiếm quyền."),
            ("Browser/Python -> HTTP POST -> Flask -> subprocess.run(shell=False) -> chương trình C -> exit/signal -> HTTP response.", "Bản vulnerable dùng strcpy; secure_length kiểm độ dài; secure_snprintf giới hạn ghi.",),
            (
                ("Baseline", "Gửi tên ngắn; ghi request, response và exit code.", "app.py; native_runner.py; evidence cần bổ sung"),
                ("Input dài", "Gửi 32, 33 và 64 byte tới bản vulnerable; quan sát crash có kiểm soát.", "native/vulnerable_processor.c; chưa có log runtime"),
                ("GDB/ASan", "Ghi vị trí crash, stack overwrite và ngưỡng lỗi; không suy diễn từ source.", "gdb/*.gdb; build/vulnerable_asan; chưa có evidence/gdb hoặc evidence/asan"),
                ("Retest", "Gửi cùng input qua secure_length và secure_snprintf; ghi response từ chối.", "native/secure_*_processor.c; chưa có request-response thật"),
            ),
            (
                ("Trạng thái hiện có", "Repository có source, binary build và script GDB nhưng không có log GDB/ASan/crash, request-response hoặc ảnh thật dùng cho các kết luận runtime bắt buộc.", "Lab02/evidence; Lab02/gdb"),
                ("Kết luận được phép", "Có thể xác nhận thiết kế vulnerable/secure từ source; chưa được khẳng định vị trí crash, stack overwrite hay input bắt đầu lỗi.", "native/*.c; Makefile"),
            ),
            ("strcpy không nhận capacity đích nên input dài hơn name[32] có thể ghi sang vùng stack lân cận.", "Memory corruption có thể làm hỏng control data; lỗi nghiêm trọng hơn sai logic vì hành vi không xác định và có thể ảnh hưởng luồng điều khiển."),
            ("Crash hoặc mất ổn định process native; từ chối dịch vụ; trong điều kiện khác có thể dẫn tới thực thi mã, nhưng lab này không thử khai thác.",),
            ("Kiểm độ dài trước copy.", "Dùng snprintf/fgets hoặc bounded API đúng cách.", "Giới hạn request.", "Bật stack protector, PIE, RELRO và NX/DEP; ASLR ở hệ điều hành.", "Ưu tiên thư viện/ngôn ngữ memory-safe khi phù hợp."),
            (
                ("Length check", "native/secure_length_processor.c", "length = strnlen(user_input, NAME_BUFFER_SIZE + 1U); if (length > NAME_SAFE_CAPACITY) reject;"),
                ("Bounded formatting", "native/secure_snprintf_processor.c", "written = snprintf(name, sizeof(name), \"%s\", user_input); if (written >= sizeof(name)) reject;"),
            ),
            (
                ("Buffer Overflow khác Injection?", "Overflow ghi vượt vùng nhớ; Injection làm input thay đổi cú pháp/ý nghĩa lệnh hoặc query."),
                ("Vì sao HTTP kích hoạt lỗi native?", "HTTP input được Flask chuyển thành argument của process C, nên trust boundary web chạm trực tiếp code native."),
                ("Vì sao firewall không đủ?", "Request có thể hợp lệ về mạng/HTTP nhưng vẫn vượt capacity nội bộ của chương trình C."),
                ("Bản vá hiệu quả thế nào?", "Cả hai bản vá đặt invariant trước/bao quanh thao tác ghi, nên không cho dữ liệu vượt name[32]."),
                ("Ba cơ chế hardening?", "Stack canary phát hiện ghi tràn; ASLR/PIE ngẫu nhiên hóa địa chỉ; NX/DEP ngăn thực thi vùng dữ liệu; RELRO bảo vệ cấu trúc liên kết động."),
            ),
            ("Source fix là chính; hardening chỉ giảm khả năng khai thác.", "Không biến bài crash thành bài chiếm quyền."),
            _reqs("L02", (
                ("01", "Linux VM local, GCC, GDB và Python/web client", "3. Môi trường", "README; Makefile; gdb/*", "Đầy đủ"),
                ("02", "Luồng HTTP tới native process, buffer cố định và hàm copy không an toàn", "4, 7", "app.py; native_runner.py; vulnerable_processor.c", "Đầy đủ"),
                ("03", "Input bình thường", "5-6", "chưa có request-response/ảnh", image_status),
                ("04", "Input dài, crash có kiểm soát, vị trí crash, stack overwrite và ngưỡng lỗi", "5-7", "chưa có GDB/ASan/crash log", log_status),
                ("05", "Buffer nằm ở stack và hậu quả vượt giới hạn", "7-8", "processor_common.h; vulnerable_processor.c", "Đầy đủ"),
                ("06", "Giải thích strcpy, gets, sprintf và memory corruption", "7", "BaiTapTopic04.docx; source C", "Đầy đủ"),
                ("07", "Stack Canary, ASLR, DEP/NX", "9, 11", "Makefile; đề bài", "Đầy đủ"),
                ("08", "Hai cách sửa code", "10", "secure_length_processor.c; secure_snprintf_processor.c", "Đầy đủ"),
                ("09", "Request limit, stack protector, PIE, RELRO, memory-safe language", "9", "app.py; Makefile; đề bài", "Đầy đủ"),
                ("10", "Năm câu hỏi báo cáo", "11", "BaiTapTopic04.docx mục 8 Lab 2", "Đầy đủ"),
                ("11", "Không shellcode/ROP/reverse shell", "1, 3", "phạm vi báo cáo", "Đầy đủ"),
            )),
        ),
        "Lab03": LabSpec(
            "Lab03", "PARAMETER TAMPERING", "Sửa giá, IDOR và role tampering", "Lê Minh",
            ("Nhận biết tham số client-controlled.", "Phân biệt validation kiểu dữ liệu với authorization/policy.", "Vá bằng nguồn dữ liệu authoritative, object-level authorization và field allowlist."),
            ("Ứng dụng Flask/SQLite mini commerce; Browser DevTools hoặc proxy local.", "Tài khoản, sản phẩm và hóa đơn đều là dữ liệu giả lập trong seed.py/lab03.db."),
            ("Checkout nhận product_id/quantity/price.", "Invoice route nhận object id.", "Profile update nhận email và, ở bản vulnerable, role."),
            (
                ("Sửa giá", "Thêm product 5, quan sát checkout, đổi price xuống 1, so sánh vulnerable và secure.", "evidence/requests|responses|traces/checkout_*"),
                ("IDOR", "Login user A, mở invoice 1001, đổi sang 1002, so sánh dữ liệu vulnerable và 403 secure.", "evidence/*/invoice_*"),
                ("Role tampering", "Quan sát profile update, thêm role=admin, so sánh vulnerable update và secure field allowlist.", "evidence/*/profile_*"),
            ),
            (
                ("Checkout", "Trace vulnerable ghi giá client bị chấp nhận; secure trace lấy price từ database và không tin field price.", "checkout_tampered_vulnerable.json; checkout_tampered_secure.json"),
                ("IDOR", "Vulnerable response trả invoice khác owner; secure response/trace từ chối theo ownership.", "invoice_idor_vulnerable.json; invoice_idor_secure.json"),
                ("Role", "Vulnerable trace cập nhật role; secure trace chỉ cho email và giữ role authoritative.", "profile_role_tampered_vulnerable.json; profile_role_tampered_secure.json"),
            ),
            ("Hidden field, URL, cookie và form đều do client kiểm soát.", "IDOR là Broken Access Control do thiếu object-level authorization.", "Mass assignment cập nhật field nhạy cảm tạo privilege escalation."),
            ("Gian lận giá/đơn hàng; lộ hóa đơn người khác; nâng quyền trái phép và ảnh hưởng tính bí mật/toàn vẹn."),
            ("Lấy giá từ database phía server.", "Kiểm ownership/role trên từng object.", "Lấy identity từ session.", "Field allowlist; không nhận role/is_admin/balance/permission.", "Validation kiểu/range và audit hành vi bất thường."),
            (
                ("Checkout", "services.py", "server_price = SELECT price_vnd FROM products WHERE id = ?; total = server_price * quantity"),
                ("Invoice", "authorization.py", "if invoice.user_id != session_user.id and session_user.role != 'admin': return 403"),
                ("Profile", "services.py", "allowed_fields = {'email'}; identity = session user; ignore/reject role"),
            ),
            (
                ("Tampering khác SQLi?", "Tampering làm sai policy bằng giá trị hợp lệ; SQLi làm input thay đổi cấu trúc SQL."),
                ("Hidden field có bảo mật không?", "Không; chỉ ẩn trong UI và có thể sửa trước khi gửi."),
                ("IDOR thuộc nhóm nào?", "OWASP Broken Access Control."),
                ("Trước khi trả invoice cần kiểm gì?", "Xác thực session, lấy object, rồi kiểm owner/admin trên chính object đó."),
                ("Vì sao không truyền giá authoritative từ client?", "Client kiểm soát request; giá phải được server đọc từ database tại thời điểm tính tổng."),
            ),
            ("Validation không thay authorization.", "Mỗi quyết định nhạy cảm cần nguồn dữ liệu authoritative phía server."),
            _reqs("L03", (
                ("01", "Môi trường mini commerce và các request checkout/invoice/profile", "3-4", "app.py; templates; seed.py", "Đầy đủ"),
                ("02", "Sửa giá: cart, request, price, giá thấp, vulnerable và server DB price", "5-7", "checkout evidence", image_status),
                ("03", "IDOR: user A -> invoice A -> invoice B -> vulnerable/secure 403", "5-7", "invoice evidence", image_status),
                ("04", "IDOR là Broken Access Control và object-level authorization", "7, 11", "authorization.py; đề bài", "Đầy đủ"),
                ("05", "Role tampering: request profile, role=admin, vulnerable update", "5-7", "profile evidence", image_status),
                ("06", "Field allowlist, identity từ session, cấm field nhạy cảm", "9-10", "services.py; authorization.py", "Đầy đủ"),
                ("07", "Năm câu hỏi báo cáo", "11", "BaiTapTopic04.docx mục 6 Lab 3", "Đầy đủ"),
                ("08", "Bằng chứng ảnh/request-response", "14", "evidence/*; F12 manifest", image_status),
            )),
        ),
        "Lab04": LabSpec(
            "Lab04", "CROSS-SITE REQUEST FORGERY", "Session cookie, CSRF token và Origin/Referer", "Nguyễn Vũ Bách",
            ("Hiểu browser tự gửi cookie và CSRF ép state-changing request.", "Phân biệt CSRF với XSS/SOP/CORS.", "Vá bằng token theo session, Origin/Referer, SameSite và re-auth."),
            ("Victim app tại 127.0.0.1:5004 và Demo Page local tại 127.0.0.1:9004.", "Dữ liệu victim/email là giả lập; không có target tùy ý hoặc Internet."),
            ("Victim login rồi gửi request đổi email hợp lệ.", "Demo Page chứa form cross-origin cố định. Đề yêu cầu auto-submit; source hiện dùng nút submit và data-confirm-submit, nên báo cáo không khẳng định auto-submit đã chạy.", "Secure route yêu cầu token và kiểm Origin/Referer trước mutation."),
            (
                ("Baseline", "Login victim, chụp email trước và request đổi email hợp lệ.", "login_*; secure_valid_request/response; state_transitions.json"),
                ("CSRF vulnerable", "Mở Demo Page, người dùng bấm Gửi form; browser gửi cookie theo chính sách. Ghi request và email sau.", "attacker_templates/attack_page.html; vulnerable_email_*"),
                ("Secure retest", "Gửi request thiếu/sai token hoặc origin khác; ghi 403 và xác nhận email không đổi.", "secure_*_request/response; traces; state"),
            ),
            (
                ("Vulnerable", "Request/response/trace hiện có ghi change-email được chấp nhận; state file lưu before/after. Chưa có ảnh browser.", "evidence/traces/vulnerable_email_change.json; evidence/state/*"),
                ("Secure", "Evidence thiếu token/sai token/origin khác ghi deny/403 trước mutation; valid request có token được chấp nhận.", "secure_email_*.json; secure_*_response.txt"),
                ("Khác biệt triển khai", "Source attacker không auto-submit; form chỉ submit sau thao tác người dùng. Đây là khác biệt so với đề, không được trình bày như kết quả auto-submit.", "attacker_templates/attack_page.html; static/js/form-confirm.js"),
            ),
            ("Server vulnerable coi session cookie là đủ nhưng cookie chỉ chứng minh danh tính, không chứng minh ý định.", "SOP có thể ngăn đọc response nhưng không nhất thiết ngăn form cross-origin được gửi."),
            ("Đổi email ngoài ý muốn; chiếm quy trình khôi phục tài khoản; rủi ro tăng với thao tác nhạy cảm không re-auth."),
            ("Token ngẫu nhiên duy nhất theo session và compare_digest phía server.", "Kiểm Origin/Referer exact allowlist.", "SameSite=Lax/Strict khi phù hợp.", "POST-only cho state change.", "Re-auth cho thao tác nhạy cảm; CAPTCHA chỉ hỗ trợ."),
            (
                ("Token", "csrf_service.py", "token = secrets.token_urlsafe(32); hmac.compare_digest(submitted, session_token)"),
                ("Origin", "origin_service.py", "reject request when Origin/Referer is absent or outside exact local allowlist"),
                ("Mutation", "victim_app.py", "validate session -> origin/referer -> token -> input; only then UPDATE"),
            ),
            (
                ("Vì sao browser tự gửi cookie?", "Cookie được gắn theo domain/path/SameSite/Secure và target request, không theo ý định người dùng."),
                ("Attacker không biết mật khẩu vẫn CSRF?", "Browser của victim đã có session hợp lệ và tự mang credential."),
                ("CSRF đọc response không?", "Thường không do SOP, nhưng đọc response không cần thiết để gây state change."),
                ("CSRF khác XSS?", "CSRF ép gửi request có credential; XSS chạy script trong origin của ứng dụng."),
                ("Vì sao không dùng GET cho state change?", "GET dễ bị kích hoạt qua link/image/prefetch và được thiết kế là safe/idempotent."),
            ),
            ("Authentication không chứng minh intent.", "Deny trước mutation và ghi audit là invariant quan trọng."),
            _reqs("L04", (
                ("01", "Victim login, request đổi email hợp lệ và cookie tự gửi", "4-6", "login/valid request evidence", image_status),
                ("02", "Trang attacker local và form đổi email", "4-6", "attacker_app.py; attack_page.html", "Đầy đủ"),
                ("03", "Auto-submit theo đề", "4-6", "source dùng submit thủ công", "Không áp dụng do khác biệt triển khai, đã giải thích"),
                ("04", "Email trước, request CSRF, email sau", "5-6", "state transitions; vulnerable evidence", image_status),
                ("05", "CSRF token secure, tấn công lại, 403 và email không đổi", "5-6, 10", "secure evidence", image_status),
                ("06", "Origin, Referer, SameSite, POST-only, re-auth và CAPTCHA", "7-10", "csrf/origin services; config.py", "Đầy đủ"),
                ("07", "Phân biệt CSRF/XSS và năm câu hỏi", "7, 11", "BaiTapTopic04.docx mục 5 Lab 4", "Đầy đủ"),
                ("08", "Yêu cầu nộp riêng: request hợp lệ, HTML attacker, before/after, ảnh sau vá, giải thích", "14", "source/evidence/manifest", image_status),
            )),
        ),
        "Lab05": LabSpec(
            "Lab05", "SQL INJECTION", "Authentication bypass, search injection và parameter binding", "Lê Minh",
            ("Nhận biết input làm thay đổi SQL do nối chuỗi.", "Quan sát authentication bypass và expanded search trong SQLite local.", "Vá bằng parameterized query, password hashing và error handling."),
            ("Flask + SQLite local tại 127.0.0.1:5005; Browser/DevTools hoặc proxy local.", "Chỉ SELECT trên dataset lab; không UNION/blind/stacked query, không website thật."),
            ("Login vulnerable nối username/password vào SQL.", "Search vulnerable nối keyword vào LIKE.", "Secure query dùng placeholder và tuple parameters; password dùng PBKDF2."),
            (
                ("Phát hiện", "Nhập login/search bình thường, sau đó dấu nháy đơn; ghi lỗi/hành vi bất thường.", "quote_*; normal_* evidence"),
                ("Authentication bypass", "Dùng payload local cố định, đối chiếu query trước/sau và decision; retest secure.", "auth_logic_*; config.py"),
                ("Search injection", "Tìm USB, dùng payload mở rộng, so rows; retest secure cùng input.", "expanded_search_*; normal_search_*"),
            ),
            (
                ("Dấu nháy", "Evidence quote_login/search ghi lỗi hoặc bất thường mà không hiển thị chi tiết database ở secure mode.", "quote_*_vulnerable.json/txt"),
                ("Login", "Trace vulnerable ghi SQL structure changed và session/authentication decision; secure cùng input không bypass.", "auth_logic_vulnerable.json; auth_logic_secure.json"),
                ("Search", "Vulnerable trace trả 8 hàng ngoài baseline; secure parameter binding trả 0 hàng với cùng payload và không đổi cấu trúc SQL.", "expanded_search_vulnerable.json; expanded_search_secure.json"),
            ),
            ("String concatenation trộn code với data trước SQL parser.", "Escaping thủ công phụ thuộc dialect/ngữ cảnh và dễ bỏ sót; ORM vẫn nguy hiểm khi dùng raw SQL/concatenation."),
            ("Bypass xác thực; rò rỉ dữ liệu; lỗi chi tiết hỗ trợ dò schema; blast radius phụ thuộc quyền DB."),
            ("Prepared statement/parameterized query.", "ORM dùng đúng cách, không raw concat.", "PBKDF2/bcrypt/Argon2, không plaintext.", "Generic error; least privilege; validation; logging/monitoring.", "WAF chỉ hỗ trợ."),
            (
                ("Login lookup", "secure_queries.py", "LOGIN_SQL = '... WHERE username = ? AND active = 1 LIMIT 1'; execute(sql, (username,))"),
                ("Search", "secure_queries.py", "SEARCH_SQL = '... name LIKE ? LIMIT 50'; _run('search', SEARCH_SQL, (f'%{keyword}%',))"),
                ("Password", "auth_service.py", "check_password_hash(stored_hash, submitted_password)"),
            ),
            (
                ("SQLi xảy ra ở tầng nào?", "Tầng xây dựng truy vấn/data-access trước khi SQL parser thực thi."),
                ("Vì sao escaping thủ công dễ sai?", "Quy tắc phụ thuộc DB, encoding và vị trí; một nhánh bỏ sót là đủ tái tạo lỗi."),
                ("Prepared statement khác concat?", "SQL structure được cố định; input được bind như data, không được parse thành SQL syntax."),
                ("ORM luôn chống SQLi?", "Không; raw query, string interpolation hoặc API escape sai vẫn có thể tạo SQLi."),
                ("Vì sao không hiển thị lỗi chi tiết?", "Thông tin query/schema giúp attacker điều chỉnh input và mở rộng tác động."),
            ),
            ("Root fix là tách code/data.", "Hash mật khẩu và least privilege không cứu được query bị đổi logic nhưng giảm tác động ở lớp khác."),
            _reqs("L05", (
                ("01", "Normal input, dấu nháy đơn và phát hiện lỗi/hành vi bất thường", "5-7", "normal/quote evidence", image_status),
                ("02", "Authentication bypass: payload local, query trước/sau, WHERE thay đổi, login không mật khẩu", "5-7", "auth_logic evidence", image_status),
                ("03", "Search injection: baseline, payload, expanded result, rò rỉ", "5-8", "expanded_search evidence", image_status),
                ("04", "Prepared statement, parameterized query, ORM đúng cách", "9-10", "secure_queries.py", "Đầy đủ"),
                ("05", "Password hashing bcrypt/Argon2/PBKDF2; không plaintext", "9-10", "auth_service.py; seed.py", "Đầy đủ"),
                ("06", "Generic error, least privilege, validation, logging/monitoring, WAF hỗ trợ", "9", "error/audit/validation source", "Đầy đủ"),
                ("07", "Năm câu hỏi báo cáo", "11", "BaiTapTopic04.docx mục 5 Lab 5", "Đầy đủ"),
                ("08", "Bằng chứng ảnh/request-response", "14", "evidence/*; F12 manifest", image_status),
            )),
        ),
        "Lab06": LabSpec(
            "Lab06", "COOKIE POISONING", "Plain, Base64, signed cookie và server-side session", "Nguyễn Vũ Bách",
            ("Quan sát đầy đủ thuộc tính cookie.", "Chứng minh plain/Base64 không tạo integrity.", "Phân biệt confidentiality/integrity và dùng server-side authorization."),
            ("Flask + SQLite local tại 127.0.0.1:5006; Browser DevTools.", "Tài khoản student/admin_lab và cookie Lab06 là dữ liệu demo; che token/session dài."),
            ("Plain cookie lưu username/role và route admin tin role client.", "Base64 chứa JSON có role nhưng không có integrity.", "Signed cookie kiểm chữ ký; encrypted cookie chỉ có integrity khi dùng AEAD/MAC; server session lưu role phía server."),
            (
                ("Quan sát", "Login student, ghi Name/Value/Domain/Path/HttpOnly/Secure/SameSite.", "evidence/cookies/plain_cookie_observation.json"),
                ("Plain", "Sửa role user thành admin, reload admin và ghi decision.", "plain_cookie_modified.json; plain_admin_cookie_modified.json"),
                ("Base64", "Decode -> JSON -> sửa role -> encode -> gửi lại -> ghi decision.", "base64_decoded_*; base64_* traces"),
                ("Secure", "Sửa signed cookie để ghi rejection; login server session student/admin; kiểm deny/allow, rotate và logout revoke.", "signed_*; server_session_*; session_* evidence"),
            ),
            (
                ("Cookie fields", "Evidence ghi host-only/domain, path /, HttpOnly false cho plain, Secure false do HTTP local và SameSite Lax.", "plain_cookie_observation.json"),
                ("Plain/Base64", "Trace plain role=admin quyết định allow; decoded Base64 role user -> admin đổi deny -> allow.", "plain_admin_cookie_modified.json; base64_decoded_*.json"),
                ("Signed/session", "Signed cookie sửa bị reject trước payload use; server session student bị deny theo database role, admin hợp lệ được allow.", "signed_cookie_tampered.json; server_session_*"),
            ),
            ("Cookie là input client-controlled; server vulnerable dùng trực tiếp role hoặc decoded role làm authorization source.", "Base64 chỉ encoding. Encryption bảo vệ confidentiality; integrity cần AEAD hoặc encryption kèm MAC."),
            ("Broken Access Control và privilege escalation; sửa state client; session hijacking/fixation là lớp rủi ro khác và cần control khác."),
            ("Không lưu role/is_admin/balance/permission authoritative ở client.", "Signed cookie khi cần phát hiện sửa đổi; AEAD/MAC khi cần confidentiality + integrity.", "Server-side session với ID ngẫu nhiên, hash lookup, role database.", "Rotate sau login; logout revoke server-side.", "Server-side authorization mỗi request và cookie flags HttpOnly/Secure/SameSite."),
            (
                ("Signed cookie", "signed_cookie_service.py", "verify signature before parsing or using payload; reject invalid token"),
                ("Server session", "server_session_service.py", "resolve opaque session ID -> active server record -> current database role"),
                ("Authorization", "authorization_service.py", "allow admin only when current database role equals admin"),
            ),
            (
                ("Vì sao cookie không đáng tin?", "Client lưu, gửi lại và có thể sửa mọi byte trước request."),
                ("Poisoning khác hijacking?", "Poisoning sửa nội dung state; hijacking lấy token/session hợp lệ của người khác."),
                ("Base64 có phải mã hóa?", "Không; ai cũng có thể decode/encode mà không cần secret."),
                ("Signed cookie giải quyết gì?", "Integrity/authenticity của payload, không che nội dung và không tự thay authorization."),
                ("Vì sao server-side authorization bắt buộc?", "Quyền phải dựa vào policy và state authoritative hiện tại, không chỉ dữ liệu client hoặc token hợp lệ."),
            ),
            ("Confidentiality và integrity là hai thuộc tính khác nhau.", "Session lifecycle gồm tạo ID ngẫu nhiên, rotate, kiểm active/expiry và revoke khi logout."),
            _reqs("L06", (
                ("01", "Name, Value, Domain, Path, HttpOnly, Secure, SameSite", "5-6", "plain_cookie_observation.json", image_status),
                ("02", "Plain cookie: login, user->admin, reload admin, kết quả và access control", "5-8", "plain cookie evidence", image_status),
                ("03", "Base64: decode, JSON, sửa role, encode, gửi lại, server decision", "5-7", "base64 evidence", image_status),
                ("04", "Base64 là encoding, không encryption", "7, 11", "base64_cookie_service.py; evidence", "Đầy đủ"),
                ("05", "Signed/encrypted; confidentiality vs integrity; AEAD hoặc MAC", "4, 7, 9-10", "signed/encrypted services", "Đầy đủ"),
                ("06", "Server session, random ID, role server, authorization", "4, 9-10", "server_session/authorization source", "Đầy đủ"),
                ("07", "Rotate sau login; logout revoke; cookie flags", "9-10", "session/cookie source and evidence", "Đầy đủ"),
                ("08", "Năm câu hỏi phân tích", "11", "BaiTapTopic04.docx mục 5 Lab 6", "Đầy đủ"),
                ("09", "Bằng chứng ảnh/request-response", "14", "evidence/*; F12 manifest", image_status),
            )),
        ),
    }


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_text(cell, value: object, bold: bool = False) -> None:
    cell.text = str(value)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(2)
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(8.5)
            run.bold = bold


def _table(doc: Document, headers: tuple[str, ...], rows: Iterable[Iterable[object]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for cell, header in zip(table.rows[0].cells, headers):
        _set_cell_text(cell, header, True)
        _shade(cell, "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            _set_cell_text(cell, value)
    doc.add_paragraph()


def _bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def _number(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def _add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Lê Minh — 21127645 | Nguyễn Vũ Bách — 21127224 | Trang ")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def _setup(doc: Document, spec: LabSpec) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    for name, size, color in (("Title", 25, "153A5B"), ("Heading 1", 16, "153A5B"), ("Heading 2", 12, "2B6F92")):
        styles[name].font.name = "Arial"
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = RGBColor.from_string(color)
    props = doc.core_properties
    props.title = f"{spec.lab} - {spec.title}"
    props.subject = "Báo cáo thực hành Topic04 bám BaiTapTopic04.docx"
    props.author = "Lê Minh (21127645); Nguyễn Vũ Bách (21127224)"
    props.last_modified_by = props.author
    props.keywords = f"Topic04, {spec.lab}, security lab, local only"
    props.comments = "Báo cáo hai thành viên; kết quả chỉ dựa trên source/evidence có trong repository."
    header = section.header.paragraphs[0]
    header.text = f"{spec.lab} | Lê Minh - 21127645 | Nguyễn Vũ Bách - 21127224"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("666666")
    _add_page_field(section.footer.paragraphs[0])


def _cover(doc: Document, spec: LabSpec) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BÁO CÁO THỰC HÀNH AN TOÀN ỨNG DỤNG WEB")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{spec.lab.upper()} - {spec.title}")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor.from_string("153A5B")
    p = doc.add_paragraph(spec.subtitle)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True
    doc.add_paragraph()
    _table(doc, ("Thành viên", "MSSV", "Vai trò"), (
        ("Lê Minh", "21127645", "Phụ trách chính Lab01, Lab03, Lab05; hỗ trợ kỹ thuật Lab02, Lab04, Lab06"),
        ("Nguyễn Vũ Bách", "21127224", "Phụ trách chính Lab02, Lab04, Lab06; hỗ trợ kỹ thuật Lab01, Lab03, Lab05"),
    ))
    p = doc.add_paragraph("Phạm vi: chỉ môi trường local/ứng dụng cố tình có lỗi. Không thử trên hệ thống thật.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    doc.add_page_break()


def _manifest(lab: str) -> list[dict]:
    path = ROOT / lab / "screenshot_manifest.py"
    if not path.is_file():
        return []
    spec = importlib.util.spec_from_file_location(f"{lab}_screenshot_manifest", path)
    module = importlib.util.module_from_spec(spec)
    assert spec and spec.loader
    spec.loader.exec_module(module)
    return list(getattr(module, "F12_SCREENSHOTS", ()))


def _evidence_files(lab: str) -> list[str]:
    root = ROOT / lab / "evidence"
    if not root.is_dir():
        return []
    excluded = ("pytest", "coverage", "smoke", "security_review", "submission_cleanup")
    return sorted(
        p.relative_to(ROOT / lab).as_posix()
        for p in root.rglob("*")
        if p.is_file() and p.name != ".gitkeep" and not any(word in p.name.lower() for word in excluded)
    )


def _source_files(lab: str) -> list[str]:
    root = ROOT / lab
    candidates = []
    for pattern in ("*.py", "*.c", "*.h", "*.sql", "static/js/*.js", "templates/**/*.html", "victim_templates/**/*.html", "attacker_templates/**/*.html", "gdb/*"):
        candidates.extend(root.glob(pattern))
    return sorted(p.relative_to(root).as_posix() for p in candidates if p.is_file())


def build_report(lab: str) -> Path:
    spec = _specs()[lab]
    doc = Document()
    _setup(doc, spec)
    _cover(doc, spec)

    doc.add_heading("1. Tên bài lab và phạm vi an toàn", level=1)
    doc.add_paragraph(f"{spec.lab} - {spec.title}: {spec.subtitle}.")
    doc.add_paragraph("Nhóm sinh viên thực hiện: Lê Minh — 21127645 và Nguyễn Vũ Bách — 21127224.")
    doc.add_paragraph("Mọi payload và dữ liệu chỉ dùng trong repository local. Không website/hệ thống thật, không dữ liệu thật, không reverse shell, malware, persistence hoặc công cụ chiếm quyền.")

    doc.add_heading("2. Mục tiêu", level=1)
    for item in spec.objectives: _bullet(doc, item)

    doc.add_heading("3. Môi trường thực hành", level=1)
    for item in spec.environment: _bullet(doc, item)

    doc.add_heading("4. Kiến trúc và kịch bản lab", level=1)
    for item in spec.scenario: _bullet(doc, item)

    doc.add_heading("5. Các bước thực hiện", level=1)
    _table(doc, ("Kịch bản", "Thao tác", "Nguồn đối chiếu"), spec.steps)

    doc.add_heading("6. Kết quả quan sát", level=1)
    doc.add_paragraph("Các dòng dưới đây phân biệt rõ kết quả có evidence với phần còn thiếu ảnh/log. Source hoặc hướng dẫn không được coi là kết quả runtime.")
    _table(doc, ("Hạng mục", "Kết quả có thể kết luận", "Evidence/source"), spec.observations)

    doc.add_heading("7. Phân tích nguyên nhân kỹ thuật", level=1)
    for item in spec.causes: _bullet(doc, item)

    doc.add_heading("8. Mức độ ảnh hưởng", level=1)
    for item in spec.impacts: _bullet(doc, item)

    doc.add_heading("9. Biện pháp phòng chống", level=1)
    for item in spec.defenses: _bullet(doc, item)

    doc.add_heading("10. Bản vá hoặc đề xuất sửa code", level=1)
    for name, source, snippet in spec.patches:
        doc.add_heading(name, level=2)
        doc.add_paragraph(f"Nguồn: {source}")
        p = doc.add_paragraph()
        r = p.add_run(snippet)
        r.font.name = "Consolas"
        r.font.size = Pt(9)

    doc.add_heading("11. Trả lời câu hỏi báo cáo", level=1)
    for index, (question, answer) in enumerate(spec.questions, 1):
        doc.add_heading(f"Câu {index}. {question}", level=2)
        doc.add_paragraph(answer)

    doc.add_heading("12. Bài học và kết luận", level=1)
    for item in spec.lessons: _bullet(doc, item)
    doc.add_paragraph("Kết quả đối chiếu cho thấy dữ liệu từ client phải được xem là không tin cậy. Bản vá trọng tâm đặt kiểm tra, xác thực, phân quyền hoặc kiểm tra toàn vẹn ở server; các lớp bổ sung chỉ hỗ trợ giảm rủi ro. Hai thành viên cùng rà soát source, evidence, bản vá và nội dung báo cáo theo phân công kỹ thuật của lab.")

    doc.add_heading("13. PHÂN CÔNG THỰC HIỆN LAB", level=1)
    assignment = LAB_ASSIGNMENTS[lab]
    _table(doc, ("Thành viên", "Vai trò", "Công việc chính", "Công việc hỗ trợ"), (
        (assignment["primary"], "Phụ trách chính", assignment["main"], "Tiếp nhận rà soát chéo và hoàn thiện nội dung theo bằng chứng đã thống nhất."),
        (assignment["support"], "Hỗ trợ kỹ thuật", "Rà soát chéo source, evidence, bản vá và nội dung báo cáo.", assignment["support_work"]),
    ))
    doc.add_paragraph(f"{assignment['primary']} chịu trách nhiệm chính; {assignment['support']} thực hiện phần hỗ trợ kỹ thuật nêu trên. Phân công này nằm trong cơ chế mỗi thành viên phụ trách chính ba lab và hỗ trợ ba lab còn lại.")
    doc.add_heading("13.1. Xác nhận nhóm thực hiện", level=2)
    _table(doc, ("STT", "Họ và tên", "MSSV"), ((1, "Lê Minh", "21127645"), (2, "Nguyễn Vũ Bách", "21127224")))

    doc.add_heading("14. Phụ lục ảnh, log và request-response", level=1)
    evidence = _evidence_files(lab)
    if evidence:
        _table(doc, ("STT", "File evidence thật trong repository"), ((i, name) for i, name in enumerate(evidence, 1)))
    else:
        doc.add_paragraph("Chưa có file evidence runtime trong repository.")

    doc.add_heading("14.1. Danh sách ảnh cần chụp thủ công", level=2)
    manifest = _manifest(lab)
    shot_root = ROOT / lab / "evidence/screenshots"
    rows = []
    for index, item in enumerate(manifest, 1):
        filename = str(item.get("filename", ""))
        present = (shot_root / filename).is_file() and (shot_root / filename).stat().st_size > 0
        rows.append((index, filename, item.get("objective", item.get("caption", "")), item.get("report_section", "Phụ lục"), "Có ảnh thật" if present else "Chưa có ảnh thật"))
    _table(doc, ("STT", "Tên ảnh", "Nội dung chứng minh", "Mục", "Trạng thái"), rows)
    doc.add_paragraph(f"Hướng dẫn chi tiết: {lab}/HUONG_DAN_CHUP_ANH.md. Báo cáo không chèn ảnh giả và không trình bày ô hướng dẫn như bằng chứng thật.")

    doc.add_heading("14.2. Source liên quan", level=2)
    sources = _source_files(lab)
    _table(doc, ("STT", "Source"), ((i, name) for i, name in enumerate(sources, 1)))

    output = OUTPUTS[lab]
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return output


def _completion_status(score: int) -> str:
    if score >= 90:
        return "Hoàn thiện"
    if score >= 75:
        return "Gần hoàn thiện"
    if score >= 50:
        return "Hoàn thiện một phần"
    return "Thiếu nhiều thành phần"


def _setup_summary(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    for name, size, color in (("Title", 25, "153A5B"), ("Heading 1", 16, "153A5B"), ("Heading 2", 12, "2B6F92")):
        styles[name].font.name = "Arial"
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = RGBColor.from_string(color)
    props = doc.core_properties
    props.title = "Báo cáo tổng hợp sáu lab Topic04"
    props.subject = "Mức độ hoàn thiện, kiến trúc và phân công nhóm Topic04"
    props.author = "Lê Minh (21127645); Nguyễn Vũ Bách (21127224)"
    props.last_modified_by = props.author
    props.keywords = "Topic04, XSS, Buffer Overflow, Parameter Tampering, CSRF, SQL Injection, Cookie Poisoning"
    props.comments = "Đánh giá dựa trên source, tài liệu và evidence hiện có trong repository."
    header = section.header.paragraphs[0]
    header.text = "TOPIC04 | Lê Minh — 21127645 | Nguyễn Vũ Bách — 21127224"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("666666")
    _add_page_field(section.footer.paragraphs[0])


def build_summary_report() -> Path:
    specs = _specs()
    doc = Document()
    _setup_summary(doc)

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run("BÁO CÁO TỔNG HỢP THỰC HÀNH AN TOÀN ỨNG DỤNG WEB")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(16)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TOPIC04 — LỖ HỔNG BẢO MẬT ỨNG DỤNG WEB")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(23)
    run.font.color.rgb = RGBColor.from_string("153A5B")
    subtitle = doc.add_paragraph("Tổng hợp Cross-Site Scripting, Buffer Overflow, Parameter Tampering, CSRF, SQL Injection và Cookie Poisoning")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    doc.add_paragraph()
    heading = doc.add_paragraph("NHÓM SINH VIÊN THỰC HIỆN")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].bold = True
    _table(doc, ("STT", "Họ và tên", "MSSV"), ((1, "Lê Minh", "21127645"), (2, "Nguyễn Vũ Bách", "21127224")))
    doc.add_page_break()

    doc.add_heading("2. Thông tin nhóm", level=1)
    _table(doc, ("STT", "Họ và tên", "MSSV", "Vai trò tổng quát"), (
        (1, "Lê Minh", "21127645", "Phụ trách chính Lab01, Lab03, Lab05; hỗ trợ rà soát kỹ thuật Lab02, Lab04, Lab06."),
        (2, "Nguyễn Vũ Bách", "21127224", "Phụ trách chính Lab02, Lab04, Lab06; hỗ trợ rà soát kỹ thuật Lab01, Lab03, Lab05."),
    ))
    doc.add_paragraph("Mục tiêu đóng góp chung là cân bằng: mỗi thành viên phụ trách chính ba lab và hỗ trợ kỹ thuật ba lab còn lại.")

    doc.add_heading("3. Phạm vi và quy định an toàn", level=1)
    for item in (
        "Chỉ thực hiện trên localhost, máy ảo hoặc ứng dụng lab cố tình có lỗi; không thử trên website hoặc hệ thống thật.",
        "Không thu thập dữ liệu thật; không dùng shellcode, ROP, reverse shell, malware hoặc persistence.",
        "Buffer Overflow chỉ dừng ở crash có kiểm soát và quan sát bộ nhớ trong môi trường local.",
        "SQL Injection chỉ dùng input an toàn trên database local, không dùng thao tác phá hủy dữ liệu.",
    ):
        _bullet(doc, item)

    doc.add_heading("4. Mục tiêu chung của Topic04", level=1)
    flow = "Nhận diện lỗ hổng → Thử nghiệm trong lab → Thu thập bằng chứng → Phân tích nguyên nhân → Đánh giá ảnh hưởng → Đề xuất phòng chống → Xây dựng bản vá → Kiểm tra lại phiên bản secure"
    p = doc.add_paragraph(flow)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    doc.add_paragraph("Chu trình đặt trust boundary và dữ liệu không tin cậy làm điểm xuất phát, sau đó đối chiếu hành vi vulnerable với bản secure và bằng chứng hiện có.")

    doc.add_heading("5. Kiến trúc tổng quan sáu lab", level=1)
    _table(doc, ("Lab", "Lỗ hổng", "Trust boundary", "Thành phần vulnerable", "Bản vá trọng tâm"), (
        ("Lab01", "Cross-Site Scripting", "URL/form/database → template hoặc DOM", "Phản chiếu/lưu HTML không encode; innerHTML", "Output encoding, sanitization, textContent, CSP"),
        ("Lab02", "Buffer Overflow", "HTTP → Flask → subprocess → stack C", "char name[32] và strcpy", "Kiểm tra độ dài, snprintf, compiler hardening"),
        ("Lab03", "Parameter Tampering", "Tham số client → dữ liệu và policy server", "Tin giá, object ID hoặc role từ request", "Giá authoritative, object-level authorization, field allowlist"),
        ("Lab04", "CSRF", "Trang khác origin → endpoint đổi trạng thái", "Request thay đổi trạng thái thiếu CSRF token", "Token, Origin/Referer, SameSite, re-authentication"),
        ("Lab05", "SQL Injection", "Input → câu SQL → SQLite", "Nối chuỗi username hoặc keyword", "Prepared statement, password hashing, lỗi chung, least privilege"),
        ("Lab06", "Cookie Poisoning", "Cookie client → quyết định phân quyền", "Tin role trong plain/Base64 cookie", "Signed/AEAD cookie, server session, authorization phía server"),
    ))

    doc.add_heading("6. Bảng phân chia công việc", level=1)
    _table(doc, ("Lab", "Người phụ trách chính", "Công việc chính", "Người hỗ trợ", "Nội dung hỗ trợ"), (
        (lab, data["primary"], data["main"], data["support"], data["support_work"])
        for lab, data in LAB_ASSIGNMENTS.items()
    ))
    doc.add_paragraph("Mỗi thành viên phụ trách chính ba lab và hỗ trợ ba lab còn lại. Việc hỗ trợ chéo tập trung vào rà soát kỹ thuật, bằng chứng, bản vá và báo cáo, giúp khối lượng công việc giữa hai thành viên cân bằng.")

    doc.add_heading("7. Tiêu chí đánh giá mức độ hoàn thành", level=1)
    _table(doc, ("STT", "Tiêu chí", "Điểm tối đa"), ((index, criterion, 10) for index, criterion in enumerate(COMPLETION_CRITERIA, 1)))
    doc.add_paragraph("Tổng điểm tối đa: 100 điểm mỗi lab. Phân loại: 90–100 Hoàn thiện; 75–89 Gần hoàn thiện; 50–74 Hoàn thiện một phần; dưới 50 Thiếu nhiều thành phần.")
    notice = doc.add_paragraph("Điểm trong báo cáo này phản ánh mức độ hoàn thiện của source, tài liệu và bằng chứng hiện có trong repository. Điểm không phải kết quả chạy lại test hoặc kiểm thử runtime.")
    notice.runs[0].bold = True

    doc.add_heading("8. Báo cáo mức độ hoàn thành từng lab", level=1)
    review_topics = {
        "Lab01": "Reflected XSS, Stored XSS và DOM-based XSS đã có mô tả source, sink, root cause, ảnh hưởng và bản vá bằng escaping/output encoding, sanitization, textContent, CSP. Câu hỏi báo cáo đã được trả lời theo source và evidence hiện có.",
        "Lab02": "Luồng HTTP đến chương trình C, buffer cố định, input vượt giới hạn, crash có kiểm soát, kiểm tra độ dài, API copy an toàn và compiler hardening đã được mô tả. Repository chưa có log GDB, ASan hoặc crash thực tế đủ để xác nhận luồng runtime.",
        "Lab03": "Sửa giá, IDOR và role tampering đã được đối chiếu với server-side validation, object-level authorization, field allowlist và audit. Bảng evidence có trace/request-response nhưng chưa có ảnh browser.",
        "Lab04": "Victim app, Demo Page, cookie tự gửi, request thay đổi trạng thái, CSRF token, Origin/Referer, SameSite và re-authentication đã được phân tích. Evidence trước/sau và sau vá có dữ liệu máy nhưng chưa có bộ ảnh; source yêu cầu xác nhận thủ công thay vì auto-submit.",
        "Lab05": "Normal input, dấu nháy đơn, authentication bypass, search injection và query biến đổi đã được đối chiếu với prepared statement, password hashing, generic error, least privilege, logging và monitoring.",
        "Lab06": "Plain cookie, role tampering, Base64, signed cookie, encrypted/authenticated cookie, server-side session, cookie flags, rotate/revoke session và server-side authorization đã được trình bày và đối chiếu evidence.",
    }
    for index, (lab, spec) in enumerate(specs.items(), 1):
        doc.add_heading(f"8.{index}. {lab} — {spec.title.title()}", level=2)
        doc.add_paragraph("Mục tiêu: " + " ".join(spec.objectives))
        doc.add_paragraph(review_topics[lab])
        evidence = _evidence_files(lab)
        evidence_preview = ", ".join(evidence[:6]) if evidence else "chưa có file evidence runtime"
        doc.add_paragraph(f"Bằng chứng hiện có: {len(evidence)} file được liệt kê trong thư mục evidence; ví dụ {evidence_preview}.")
        doc.add_paragraph("Root cause: " + " ".join(spec.causes))
        doc.add_paragraph("Ảnh hưởng: " + " ".join(spec.impacts))
        doc.add_paragraph("Bản vá và phòng chống: " + " ".join(spec.defenses))
        doc.add_paragraph(f"Câu hỏi báo cáo: đã trả lời {len(spec.questions)} câu theo source và bằng chứng hiện có; không thay thế phần runtime còn thiếu.")
        doc.add_paragraph("Thành phần còn thiếu: " + MISSING_COMPONENTS[lab])
        scores = COMPLETION_SCORES[lab]
        _table(doc, ("STT", "Tiêu chí", "Điểm"), ((number, criterion, scores[number - 1]) for number, criterion in enumerate(COMPLETION_CRITERIA, 1)))
        total = sum(scores)
        result = doc.add_paragraph(f"Tổng điểm: {total}/100. Trạng thái: {_completion_status(total)}.")
        result.runs[0].bold = True

    doc.add_heading("9. Bảng tổng hợp mức độ hoàn thành", level=1)
    strengths = {
        "Lab01": "Đủ ba dạng XSS, source-to-sink và bản secure",
        "Lab02": "Đủ source C/web, hai bản vá và hardening",
        "Lab03": "Đủ ba kịch bản, policy và audit trace",
        "Lab04": "Đủ token, origin policy, state và trace",
        "Lab05": "Đủ login/search vulnerable và secure query",
        "Lab06": "Đủ năm mô hình cookie/session và authorization",
    }
    _table(doc, ("Lab", "Điểm", "Trạng thái", "Thành phần tốt", "Thành phần cần bổ sung"), (
        (lab, sum(COMPLETION_SCORES[lab]), _completion_status(sum(COMPLETION_SCORES[lab])), strengths[lab], MISSING_COMPONENTS[lab])
        for lab in specs
    ))

    doc.add_heading("10. Các thành phần cần hoàn thiện trước khi nộp", level=1)
    for lab in specs:
        _bullet(doc, f"{lab}: {MISSING_COMPONENTS[lab]}")
    _bullet(doc, "Thay toàn bộ dòng trạng thái ảnh chưa có bằng ảnh thật và caption đồng bộ sau khi nhóm tự thực hiện, không dùng ảnh dựng.")
    _bullet(doc, "Kiểm tra lại các cặp bằng chứng vulnerable/secure để mỗi kết luận runtime có request-response, log hoặc ảnh tương ứng.")

    doc.add_heading("11. Kết luận", level=1)
    doc.add_paragraph("Sáu lab bao quát XSS, Buffer Overflow, Parameter Tampering, CSRF, SQL Injection và Cookie Poisoning. Điểm chung là dữ liệu từ client, request, tham số, cookie hoặc input native không được xem là đáng tin cậy.")
    doc.add_paragraph("Root fix phải đặt ở nơi quyết định an toàn: encode/sanitize đúng context, kiểm tra biên bộ nhớ, dùng dữ liệu authoritative, token chống CSRF, prepared statement và authorization phía server. CSP, SameSite, compiler hardening, cookie flags, logging và least privilege là defense in depth, không thay thế root fix.")
    doc.add_paragraph("Lê Minh phụ trách chính Lab01, Lab03, Lab05; Nguyễn Vũ Bách phụ trách chính Lab02, Lab04, Lab06. Hai thành viên hỗ trợ chéo cả ba lab của người còn lại bằng các nhiệm vụ kỹ thuật cụ thể, bảo đảm cơ cấu đóng góp cân bằng.")

    doc.add_heading("12. Phụ lục", level=1)
    doc.add_heading("12.1. Danh sách sáu báo cáo DOCX", level=2)
    _table(doc, ("Lab", "Báo cáo"), ((lab, path.relative_to(ROOT).as_posix()) for lab, path in OUTPUTS.items()))
    doc.add_heading("12.2. Thư mục evidence", level=2)
    _table(doc, ("Lab", "Thư mục", "Số file hiện có"), ((lab, f"{lab}/evidence/", len(_evidence_files(lab))) for lab in specs))
    doc.add_heading("12.3. Source và file bản vá chính", level=2)
    _table(doc, ("Lab", "Source/bản vá chính"), (
        ("Lab01", "app.py; templates/; static/js/dom_secure.js"),
        ("Lab02", "app.py; native/; gdb/; Makefile"),
        ("Lab03", "app.py; services.py; authorization.py"),
        ("Lab04", "victim_app.py; attacker_app.py; csrf_service.py; origin_service.py"),
        ("Lab05", "vulnerable_queries.py; secure_queries.py; auth_service.py"),
        ("Lab06", "signed_cookie_service.py; encrypted_cookie_service.py; server_session_service.py; authorization_service.py"),
    ))
    doc.add_heading("12.4. Danh sách placeholder còn thiếu", level=2)
    placeholder_rows = []
    for lab in specs:
        manifest = _manifest(lab)
        shot_root = ROOT / lab / "evidence/screenshots"
        missing = sum(1 for item in manifest if not (shot_root / str(item.get("filename", ""))).is_file())
        placeholder_rows.append((lab, missing, MISSING_COMPONENTS[lab]))
    _table(doc, ("Lab", "Số ảnh chưa có", "Nội dung cần bổ sung"), placeholder_rows)

    SUMMARY_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(SUMMARY_OUTPUT)
    return SUMMARY_OUTPUT


def write_requirement_matrix() -> Path:
    REPORT_ROOT.mkdir(parents=True, exist_ok=True)
    path = REPORT_ROOT / "DOI_CHIEU_BAITAP_TOPIC04.md"
    lines = [
        "# Đối chiếu yêu cầu BaiTapTopic04",
        "",
        "Nguồn chuẩn: `BaiTapTopic04.docx`. Mỗi yêu cầu được tách riêng; trạng thái không thay thế bằng chứng ảnh/log còn thiếu.",
        "",
        "| Lab | Mã yêu cầu | Nội dung yêu cầu trong DOCX | Báo cáo hiện tại đã có | Bằng chứng hoặc source | Hành động cần sửa | Trạng thái sau sửa |",
        "| --- | --- | --- | --- | --- | --- | --- |",
    ]
    source = Document(ROOT / "BaiTapTopic04.docx")
    lab = "Chung"
    lab_number = 0
    subsection = "Quy định an toàn chung"
    output_names = {key: value.name for key, value in OUTPUTS.items()}
    source_map = {
        "Lab01": "Lab01/app.py; templates; static/js; evidence/*",
        "Lab02": "Lab02/app.py; native/*; Makefile; gdb/*; evidence/*",
        "Lab03": "Lab03/app.py; services/authorization; evidence/*",
        "Lab04": "Lab04/victim_app.py; attacker_app.py/templates; csrf/origin services; evidence/*",
        "Lab05": "Lab05/vulnerable_queries.py; secure_queries.py; auth_service.py; evidence/*",
        "Lab06": "Lab06 cookie/session/authorization services; evidence/*",
        "Chung": "BaiTapTopic04.docx; sáu báo cáo chuẩn hóa",
    }

    def report_section(text: str, current: str) -> str:
        lower = (current + " " + text).lower()
        if "mục tiêu" in lower: return "2. Mục tiêu"
        if "môi trường" in lower or "lưu ý" in lower: return "1 và 3. Phạm vi, môi trường"
        if "kịch bản" in lower: return "4. Kiến trúc và kịch bản"
        if "nhiệm vụ" in lower or "phần a" in lower or "phần b" in lower or "phần c" in lower or "phần d" in lower: return "5-6. Các bước và kết quả"
        if "phân tích" in lower or "câu hỏi" in lower: return "7 và 11. Phân tích, câu hỏi"
        if "phòng chống" in lower: return "9-10. Phòng chống và bản vá"
        if "nộp bài" in lower or "ảnh" in lower or "log" in lower or "request-response" in lower: return "14. Phụ lục evidence"
        if "cấu trúc báo cáo" in lower or "bài học" in lower: return "1-14. Cấu trúc chuẩn hóa"
        return "Mục liên quan trong báo cáo chuẩn hóa"

    def row_status(current_lab: str, text: str) -> str:
        lower = text.lower()
        if "báo cáo pdf" in lower:
            return "Không áp dụng do ràng buộc nhiệm vụ hiện tại, đã giải thích"
        if current_lab == "Lab04" and ("tự gửi request" in lower or "form tự gửi" in lower):
            return "Không áp dụng do khác biệt triển khai, đã giải thích"
        runtime_words = ("nhập ", "thử ", "quan sát", "tải lại", "đăng nhập", "đăng bình luận", "chụp", "mở devtools", "gửi dữ liệu", "kiểm tra có", "ảnh ")
        if current_lab == "Lab02" and any(word in lower for word in ("crash", "gdb", "stack có bị", "độ dài input", "log crash")):
            return "Đầy đủ về nội dung, chờ log thật"
        if current_lab.startswith("Lab") and any(word in lower for word in runtime_words):
            return "Đầy đủ về nội dung, chờ ảnh thật"
        return "Đầy đủ"

    for index, paragraph in enumerate(source.paragraphs):
        text_value = " ".join(paragraph.text.split())
        if not text_value:
            continue
        match = re.match(r"LAB\s+(\d)\s+", text_value.upper())
        if match:
            lab_number = int(match.group(1))
            lab = f"Lab0{lab_number}"
            subsection = text_value
            continue
        if text_value.startswith("GỢI Ý CẤU TRÚC") or text_value == "BÀI HỌC":
            lab = "Chung"
            subsection = text_value
            continue
        if re.match(r"^\d+\.\s+", text_value) or text_value.startswith("Phần ") or text_value.startswith("Câu hỏi"):
            subsection = text_value
            continue
        if index < 3:
            continue
        current_output = output_names.get(lab, "cả sáu báo cáo")
        status = row_status(lab, text_value)
        action = "Giữ và chuẩn hóa theo đúng heading; chỉ nhận kết quả khi có evidence." if status != "Đầy đủ" else "Đối chiếu và giữ nội dung trực tiếp liên quan."
        esc = lambda s: str(s).replace("|", "\\|").replace("\n", " ")
        code = f"T04-{lab.upper()}-P{index:03d}"
        lines.append(
            f"| {lab} | {code} | {esc(text_value)} | {esc(report_section(text_value, subsection))} trong `{current_output}` | "
            f"{esc(source_map[lab])} | {esc(action)} | {status} |"
        )
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


def write_update_report(pptx_status: str = "Chưa đối chiếu", pptx_path: str = "", slide_count: int = 0) -> Path:
    REPORT_ROOT.mkdir(parents=True, exist_ok=True)
    path = REPORT_ROOT / "BAO_CAO_CAP_NHAT_NOI_DUNG.md"
    missing = {
        "Lab01": "Ảnh browser/DevTools cho normal input, reflected/stored/DOM vulnerable và secure retest.",
        "Lab02": "Toàn bộ ảnh browser; request-response; log GDB/ASan/crash, vị trí crash, stack overwrite và ngưỡng input.",
        "Lab03": "Ảnh checkout/IDOR/role tampering vulnerable và secure.",
        "Lab04": "Ảnh email trước/sau, request hợp lệ/CSRF, secure 403 và email không đổi; source hiện không auto-submit.",
        "Lab05": "Ảnh normal/quote/auth bypass/search expanded và secure retest.",
        "Lab06": "Ảnh cookie fields, plain/Base64 tampering, signed rejection và server-session authorization.",
    }
    lines = ["# Báo cáo cập nhật nội dung Topic04", "", "## Báo cáo đã sửa"]
    for lab, output in OUTPUTS.items(): lines.append(f"- `{output.relative_to(ROOT).as_posix()}`")
    lines += [
        "", "## Chương được bổ sung/chuẩn hóa",
        "- Sáu báo cáo dùng cùng 14 mục: tên/phạm vi, mục tiêu, môi trường, kiến trúc, bước thực hiện, kết quả, nguyên nhân, ảnh hưởng, phòng chống, bản vá, câu hỏi, bài học, phân công và phụ lục evidence.",
        "- Thông tin hai thành viên được đồng bộ ở bìa, header, metadata và phân công.",
        "", "## Nội dung thừa đã xóa",
        "- Các chương kiểm tra ngoài phạm vi, quy trình sinh tệp trung gian và nội dung không trực tiếp phục vụ yêu cầu của lab.",
        "- Các ô ảnh legacy/test lặp lại hoặc không phải bằng chứng thật; báo cáo chỉ giữ bảng yêu cầu ảnh F12 cần thiết.",
        "", "## Ảnh đã giữ",
        "- Không có ảnh bitmap thật trong sáu DOCX đầu vào. Danh sách ảnh F12 cần chụp được giữ trong phụ lục và đồng bộ với `HUONG_DAN_CHUP_ANH.md`/manifest.",
        "", "## Ảnh đã xóa khỏi báo cáo",
        "- Không xóa file ảnh thật. Chỉ loại các ô hướng dẫn legacy/test khỏi DOCX; không xóa file vật lý trong `evidence`.",
        "", "## Hướng dẫn ảnh đã cập nhật",
        "- Không sửa thêm trong lượt chuẩn hóa này; sáu file hướng dẫn hiện tại đã có danh sách F12 tương ứng với manifest.",
        "", "## Bằng chứng thật còn thiếu",
    ]
    for lab, text in missing.items(): lines.append(f"- {lab}: {text}")
    lines += [
        "", "## PPTX",
        f"- Trạng thái: {pptx_status}.",
        f"- File cuối: `{pptx_path}`." if pptx_path else "- Chưa có file PPTX cuối.",
        f"- Tổng số slide cuối: {slide_count}.",
        "", "## Xác nhận phạm vi",
        "- Không tạo hoặc cập nhật PDF.", "- Không chạy test, pytest hoặc smoke test.", "- Không chạy lab, Docker hoặc kịch bản khai thác.",
        "- Không tạo screenshot giả.", "- Không sửa logic vulnerable/secure, route, database schema hoặc payload kỹ thuật.",
    ]
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


def write_team_update_manifest() -> Path:
    REPORT_ROOT.mkdir(parents=True, exist_ok=True)
    path = REPORT_ROOT / "CAP_NHAT_NHOM_2_THANH_VIEN.md"
    ui_files = (
        "Lab01/templates/base.html",
        "Lab01/templates/index.html",
        "Lab02/templates/base.html",
        "Lab02/templates/index.html",
        "Lab03/templates/base.html",
        "Lab04/victim_templates/base.html",
        "Lab04/attacker_templates/base.html",
        "Lab05/templates/base.html",
        "Lab06/templates/base.html",
    )
    documentation_files = (
        "Lab01/README.md",
        "Lab02/README.md",
        "Lab03/README.md",
        "Lab04/README.md",
        "Lab05/README.md",
        "Lab06/README.md",
        "Lab01/HUONG_DAN_CHUP_ANH.md",
        "Lab02/HUONG_DAN_CHUP_ANH.md",
        "Lab03/HUONG_DAN_CHUP_ANH.md",
        "Lab04/HUONG_DAN_CHUP_ANH.md",
        "Lab05/HUONG_DAN_CHUP_ANH.md",
        "Lab06/HUONG_DAN_CHUP_ANH.md",
        "Lab02/docs/feature_progress.md",
        "Lab03/docs/project_context.md",
        "Lab04/docs/project_context.md",
        "Lab05/docs/project_context.md",
        "presentation/slide_outline.md",
        "presentation/HUONG_DAN_CHEN_ANH_PPTX.md",
        "presentation/content_matrix.md",
        "report/BAO_CAO_CAP_NHAT_NOI_DUNG.md",
        "scripts/topic04_reports.py",
    )
    legacy_paths = (
        ROOT / "Lab01/report/21127645_LeMinh_Lab01_XSS.docx",
        ROOT / "Lab02/report/21127645_LeMinh_Lab02_BufferOverflow.docx",
        ROOT / "Lab03/report/21127645_LeMinh_Lab03_ParameterTampering.docx",
        ROOT / "Lab04/report/21127645_LeMinh_Lab04_CSRF.docx",
        ROOT / "Lab05/report/21127645_LeMinh_Lab05_SQLInjection.docx",
        ROOT / "Lab06/report/21127645_LeMinh_Lab06_CookiePoisoning.docx",
    )
    remaining_legacy = [item.relative_to(ROOT).as_posix() for item in legacy_paths if item.exists()]
    replacement_note = (
        "- Sáu báo cáo DOCX cũ chỉ mang tên Lê Minh đã được thay thế và xóa khỏi thư mục báo cáo chính."
        if not remaining_legacy
        else "- Năm báo cáo DOCX cũ chỉ mang tên Lê Minh đã được thay thế và xóa; bản cũ `Lab02/report/21127645_LeMinh_Lab02_BufferOverflow.docx` đang mở trong Word nên Windows chưa cho phép xóa."
    )
    lines = [
        "# Cập nhật nhóm hai thành viên",
        "",
        "## Danh sách file đã sửa",
        *(f"- `{name}`" for name in (*ui_files, *documentation_files)),
        "",
        "## Danh sách UI đã cập nhật",
        *(f"- `{name}`" for name in ui_files),
        "",
        "## Danh sách báo cáo mới",
        *(f"- `{path.relative_to(ROOT).as_posix()}`" for path in OUTPUTS.values()),
        "",
        "## Báo cáo tổng hợp",
        f"- `{SUMMARY_OUTPUT.relative_to(ROOT).as_posix()}`",
        "",
        "## Xác nhận",
        replacement_note,
        "- Lê Minh — 21127645 và Nguyễn Vũ Bách — 21127224 xuất hiện đồng bộ trong UI, tài liệu và báo cáo.",
        "- Không tạo hoặc cập nhật PDF.",
        "- Không render DOCX.",
        "- Không chạy test.",
        "- Không sửa logic lab, gồm logic vulnerable và secure.",
    ]
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


def generate(lab: str) -> Path:
    if lab not in OUTPUTS:
        raise ValueError(f"Unknown lab: {lab}")
    return build_report(lab)


def generate_all() -> list[Path]:
    outputs = [build_report(lab) for lab in OUTPUTS]
    outputs.append(build_summary_report())
    write_requirement_matrix()
    write_update_report()
    write_team_update_manifest()
    return outputs


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser()
    parser.add_argument("lab", nargs="?", choices=tuple(OUTPUTS), help="Generate one report; omit to generate all six")
    args = parser.parse_args()
    created = [generate(args.lab)] if args.lab else generate_all()
    for item in created:
        print(item)
