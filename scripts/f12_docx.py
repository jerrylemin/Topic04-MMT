"""Shared DOCX-only helpers for manual DevTools evidence sections."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable, Mapping

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


FIELDS = (
    ("Tên file", "filename"),
    ("Mục tiêu chứng minh", "objective"),
    ("Trạng thái ban đầu", "setup"),
    ("URL hoặc lệnh", "url"),
    ("Tài khoản", "account"),
    ("Dữ liệu nhập", "data"),
    ("Thao tác", "action"),
    ("Tab UI", "ui_tab"),
    ("Tab F12", "f12_tab"),
    ("Request cần chọn", "request"),
    ("Trường cần mở", "field"),
    ("Nội dung bắt buộc", "required"),
    ("Kết quả mong đợi", "expected"),
    ("Nếu không thấy kết quả", "troubleshooting"),
    ("Phạm vi ảnh", "scope"),
    ("Caption dự kiến", "caption"),
    ("Mục báo cáo", "report_section"),
)


def _guide_section(guide: Path, filename: str) -> str:
    if not guide.is_file():
        return ""
    text = guide.read_text(encoding="utf-8", errors="replace")
    marker = text.find(filename)
    if marker < 0:
        return ""
    start = text.rfind("\n#", 0, marker)
    start = 0 if start < 0 else start + 1
    following = re.search(r"\n#{2,4}\s+", text[marker + len(filename):])
    end = len(text) if following is None else marker + len(filename) + following.start()
    return re.sub(r"\s+", " ", text[start:end]).strip()[:1800]


def _add_placeholder(doc, spec: Mapping[str, object], guide_text: str) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, key in FIELDS:
        value = str(spec.get(key, "") or "").strip()
        if not value and key not in {"filename", "caption"}:
            value = "Xem hướng dẫn chi tiết đúng mục ảnh này trong HUONG_DAN_CHUP_ANH.md."
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        cells[0].paragraphs[0].runs[0].bold = True
        for cell in cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
    if guide_text:
        cells = table.add_row().cells
        cells[0].text = "Hướng dẫn gốc"
        cells[1].text = guide_text
        cells[0].paragraphs[0].runs[0].bold = True
        for run in cells[1].paragraphs[0].runs:
            run.font.size = Pt(8)


def append_devtools_evidence(
    doc,
    lab_root: Path,
    specs: Iterable[Mapping[str, object]],
    *,
    start_figure: int = 1,
) -> list[str]:
    """Append a manual-evidence section and return missing screenshot filenames."""
    specs = list(specs)
    if not specs:
        return []
    doc.add_page_break()
    doc.add_heading("CHỨNG CỨ TRÌNH DUYỆT VÀ HTTP QUA DEVTOOLS F12", level=1)
    doc.add_paragraph(
        "Mục này liên kết thao tác trên giao diện với request/response HTTP thật trong "
        "Chrome hoặc Microsoft Edge DevTools. Ảnh vulnerable chứng minh dữ liệu client đi "
        "vào request và tác động tới xử lý; ảnh secure chứng minh server kiểm tra lại dữ liệu, "
        "từ chối request hoặc dùng nguồn tin cậy. Cookie, session ID và chữ ký dài phải được che."
    )
    doc.add_paragraph(
        "Trình tự mỗi kịch bản: trạng thái trước → request → payload/query parameters → response "
        "→ trạng thái sau → bản secure → phân tích. Mọi ảnh đều do sinh viên tự chụp từ môi trường local; "
        "generator chỉ chèn ảnh đã có hoặc tạo placeholder chi tiết."
    )
    shots_dir = lab_root / "evidence" / "screenshots"
    guide = lab_root / "HUONG_DAN_CHUP_ANH.md"
    missing: list[str] = []
    for offset, spec in enumerate(specs):
        number = start_figure + offset
        filename = str(spec["filename"])
        caption = str(spec.get("caption") or spec.get("objective") or filename)
        doc.add_heading(f"Hình {number}. {caption}", level=2)
        image_path = shots_dir / filename
        if image_path.is_file() and image_path.stat().st_size > 0:
            try:
                doc.add_picture(str(image_path), width=Inches(6.1))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                missing.append(filename)
                doc.add_paragraph("ẢNH CẦN BỔ SUNG", style=None).runs[0].bold = True
                _add_placeholder(doc, spec, _guide_section(guide, filename))
        else:
            missing.append(filename)
            marker = doc.add_paragraph("ẢNH CẦN BỔ SUNG")
            marker.runs[0].bold = True
            _add_placeholder(doc, spec, _guide_section(guide, filename))
        source = doc.add_paragraph(f"Tên file: {filename}. Nguồn: sinh viên chụp từ môi trường local.")
        source.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in source.runs:
            run.italic = True
            run.font.size = Pt(9)
        analysis = str(spec.get("analysis") or "Ảnh này cần đối chiếu dữ liệu hiển thị trong DevTools với trạng thái UI. Khi ảnh chưa tồn tại, báo cáo chỉ mô tả bằng chứng cần thu thập và không khẳng định kết quả đã được quan sát.")
        doc.add_paragraph(analysis)
    return missing
