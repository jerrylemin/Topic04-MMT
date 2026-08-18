import os
import subprocess
import sys
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document

from screenshot_manifest import SCREENSHOTS


ROOT = Path(__file__).parents[1]

REQUIRED_FILES = [
    "README.md", "HUONG_DAN_CHUP_ANH.md", "requirements.txt", ".env.example", ".dockerignore", "Dockerfile", "docker-compose.yml",
    "app.py", "config.py", "database.py", "auth.py", "services.py", "validators.py", "authorization.py",
    "audit_service.py", "trace_models.py", "trace_service.py", "security_utils.py", "schema.sql", "seed.py",
    "scripts/run_lab.sh", "scripts/run_lab.ps1", "scripts/run_lab.bat", "scripts/reset_database.py",
    "scripts/send_request.py", "scripts/run_demo_flows.py", "scripts/export_evidence.py", "scripts/generate_report.py",
    "scripts/check_screenshots.py", "scripts/clean_generated_files.py",
    "templates/base.html", "templates/index.html", "templates/login.html", "templates/products.html", "templates/cart.html",
    "templates/comparison.html", "templates/security_controls.html", "templates/audit_logs.html", "templates/error.html",
    "templates/vulnerable/checkout.html", "templates/vulnerable/checkout_result.html", "templates/vulnerable/invoice.html",
    "templates/vulnerable/profile.html", "templates/vulnerable/profile_result.html", "templates/secure/checkout.html",
    "templates/secure/checkout_result.html", "templates/secure/invoice.html", "templates/secure/profile.html",
    "templates/secure/profile_result.html", "templates/components/trace_panel.html",
    "templates/components/request_inspector.html", "templates/components/session_inspector.html",
    "templates/components/parameter_inspector.html", "templates/components/database_inspector.html",
    "templates/components/authorization_inspector.html", "templates/components/code_comparison.html",
    "templates/components/security_controls.html", "templates/components/final_verdict.html",
    "templates/components/audit_inspector.html", "templates/components/presentation_controls.html",
    "static/css/style.css", "static/css/presentation.css", "static/js/main.js", "static/js/trace-ui.js",
    "static/js/request-editor.js", "static/js/presentation.js", "static/js/parameter-diff.js",
]


@pytest.mark.parametrize("relative_path", REQUIRED_FILES)
def test_required_file_exists(relative_path):
    assert (ROOT / relative_path).is_file(), relative_path


def test_all_templates_parse(app):
    for path in (ROOT / "templates").rglob("*.html"):
        app.jinja_env.get_template(path.relative_to(ROOT / "templates").as_posix())


def test_frontend_has_inspectors_timeline_and_presentation():
    source = "\n".join(path.read_text(encoding="utf-8") for path in (ROOT / "templates").rglob("*.html"))
    for label in ["Request Tampering Console", "Request Inspector", "Session Inspector", "Database Inspector",
                  "Authorization Inspector", "Final Security Verdict", "Action Timeline", "Presentation Mode"]:
        assert label in source


def test_security_control_panel_reports_cookie_and_csrf_state(client):
    text = client.get("/security-controls").text
    assert "Secure cookie" in text and "CSRF protection" in text
    assert "TẮT" in text


def test_all_primary_get_routes_render(client, login):
    login()
    paths = ["/", "/login", "/products", "/products/5", "/cart", "/vulnerable/checkout", "/secure/checkout",
             "/vulnerable/invoice?id=1001", "/secure/invoice?id=1001", "/vulnerable/profile", "/secure/profile",
             "/comparison", "/security-controls", "/audit-logs", "/health"]
    assert all(client.get(path).status_code == 200 for path in paths)
    assert client.get("/secure/invoice?id=1002").status_code == 403


def test_request_editor_uses_fixed_local_routes_only():
    source = (ROOT / "static/js/request-editor.js").read_text(encoding="utf-8")
    assert "FIXED_ROUTES" in source and "fetch(target" in source
    assert "http://" not in source and "https://" not in source
    assert "LOCAL_HOSTS" in source and "location.hostname" in source


def test_client_and_demo_scripts_are_local_only():
    send = (ROOT / "scripts/send_request.py").read_text(encoding="utf-8")
    demo = (ROOT / "scripts/run_demo_flows.py").read_text(encoding="utf-8")
    assert 'BASE_URL = "http://127.0.0.1:5003"' in send
    assert "--host" not in send and "--url" not in send
    assert "import requests" not in demo and ".test_client()" in demo


def test_runtime_has_no_external_url_or_browser_automation():
    paths = [ROOT / "app.py", *list((ROOT / "scripts").glob("*.py")), *list((ROOT / "static/js").glob("*.js"))]
    source = "\n".join(path.read_text(encoding="utf-8") for path in paths).lower()
    urls = [token.split('"', 1)[0].split("'", 1)[0] for token in source.split("http://")[1:]]
    assert all(url.startswith("127.0.0.1:5003") for url in urls)
    assert "https://" not in source and "playwright" not in source and "selenium" not in source


def test_requirements_are_minimal_and_have_no_browser_driver():
    requirements = (ROOT / "requirements.txt").read_text(encoding="utf-8").lower()
    for dependency in ["flask", "pytest", "requests", "python-docx"]:
        assert dependency in requirements
    assert "playwright" not in requirements and "selenium" not in requirements


def test_docker_is_local_non_root_and_unprivileged():
    dockerfile = (ROOT / "Dockerfile").read_text(encoding="utf-8")
    compose = (ROOT / "docker-compose.yml").read_text(encoding="utf-8")
    assert "USER lab" in dockerfile and "HEALTHCHECK" in dockerfile
    assert '127.0.0.1:5003:5003' in compose
    assert "no-new-privileges:true" in compose and "cap_drop" in compose
    assert "network_mode" not in compose and "privileged:" not in compose
    config = (ROOT / "config.py").read_text(encoding="utf-8")
    assert "DATABASE_PATH" in config and "LAB03_SECRET_KEY" in config
    assert "LAB03_BIND_HOST" in compose and '127.0.0.1:5003:5003' in compose
    assert ".venv" in (ROOT / ".dockerignore").read_text(encoding="utf-8")


def test_no_real_payment_or_email_integration():
    source = "\n".join(path.read_text(encoding="utf-8").lower() for path in ROOT.glob("*.py"))
    for integration in ["stripe", "paypal", "smtplib", "sendgrid", "mailgun"]:
        assert integration not in source


def test_screenshot_guide_lists_every_required_image():
    guide = (ROOT / "HUONG_DAN_CHUP_ANH.md").read_text(encoding="utf-8")
    assert len(SCREENSHOTS) == 41
    assert all(item["filename"] in guide for item in SCREENSHOTS)
    for field in ["Tài khoản", "URL", "Dữ liệu cần sửa", "Panel cần mở", "Bước timeline", "Caption báo cáo"]:
        assert field in guide


def test_screenshot_checker_runs_without_ocr_or_generation():
    source = (ROOT / "scripts/check_screenshots.py").read_text(encoding="utf-8").lower()
    assert "ocr" not in source and "pytesseract" not in source
    env = {**os.environ, "PYTHONIOENCODING": "utf-8"}
    result = subprocess.run([sys.executable, "scripts/check_screenshots.py"], cwd=ROOT, text=True,
                            encoding="utf-8", errors="replace", capture_output=True, env=env)
    assert result.returncode in (0, 1)
    assert "Đúng tên trong manifest:" in result.stdout


def test_report_artifacts_exist_and_open():
    result = subprocess.run([sys.executable, "scripts/generate_report.py"], cwd=ROOT, check=False)
    docx = ROOT / "report/21127645_LeMinh_21127224_NguyenVuBach_Lab03_ParameterTampering.docx"
    assert result.returncode == 0 and docx.stat().st_size > 20_000
    with ZipFile(docx) as archive:
        assert "word/document.xml" in archive.namelist()


def test_report_has_all_chapters_questions_and_placeholders():
    doc = Document(ROOT / "report/21127645_LeMinh_21127224_NguyenVuBach_Lab03_ParameterTampering.docx")
    text = "\n".join(
        [paragraph.text for paragraph in doc.paragraphs]
        + [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    )
    assert all(f"{number}." in text for number in range(1, 15))
    for answer in ["Tampering khác SQLi", "Hidden field có bảo mật không",
                   "IDOR thuộc nhóm nào", "Trước khi trả invoice", "giá authoritative"]:
        assert answer in text
    assert len(doc.tables) >= 8
    assert "42_login_user_a_network.png" in text


def test_report_generator_mentions_real_log_and_missing_images():
    source = (ROOT / "scripts/generate_report.py").read_text(encoding="utf-8")
    central = (ROOT.parent / "scripts/topic04_reports.py").read_text(encoding="utf-8")
    assert "topic04_reports" in source and 'generate("Lab03")' in source
    assert "evidence" in central and "screenshot_manifest.py" in central


def test_readme_documents_all_operator_commands():
    readme = (ROOT / "README.md").read_text(encoding="utf-8")
    for command in ["python scripts/reset_database.py", "python scripts/run_demo_flows.py",
                    "python scripts/check_screenshots.py", "python scripts/generate_report.py", "pytest"]:
        assert command in readme
    for topic in ["Request Tampering Console", "Presentation Mode", "Authorization Inspector", "Database Inspector"]:
        assert topic in readme
