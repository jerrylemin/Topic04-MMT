from __future__ import annotations

import subprocess
import sys
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "report" / "21127645_LeMinh_21127224_NguyenVuBach_Lab05_SQLInjection.docx"


def test_report_wrapper_generates_current_two_member_docx():
    completed = subprocess.run(
        [sys.executable, "scripts/generate_report.py"],
        cwd=ROOT,
        capture_output=True,
        text=True,
        check=False,
    )
    assert completed.returncode == 0, completed.stderr
    assert REPORT.is_file() and REPORT.stat().st_size > 20_000

    document = Document(REPORT)
    text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )
    assert "21127645" in text and "Lê Minh" in text
    assert "21127224" in text and "Nguyễn Vũ Bách" in text
    assert "SQL INJECTION" in text
    assert "37_login_normal_request.png" in text


def test_report_wrapper_uses_shared_docx_only_generator():
    source = (ROOT / "scripts" / "generate_report.py").read_text(encoding="utf-8").lower()
    assert "topic04_reports" in source
    assert "reportlab" not in source
    assert "soffice" not in source
    assert ".pdf" not in source
