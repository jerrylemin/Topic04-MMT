"""Generate the evidence-driven Lab05 DOCX and PDF report.

The generator intentionally fails closed: it never invents traces, test
results, coverage, smoke results, or a PDF fallback.  Run the fixed demo and
evidence export first, then run this script.
"""

from __future__ import annotations

import ast
import json
import re
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "report"
DOCX_PATH = REPORT_DIR / "21127645_LeMinh_Lab05_SQLInjection.docx"
PDF_PATH = REPORT_DIR / "21127645_LeMinh_Lab05_SQLInjection.pdf"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
NAVY = RGBColor(0x20, 0x37, 0x48)
TEAL = RGBColor(0x2B, 0x51, 0x63)
GOLD = RGBColor(0x9A, 0x7B, 0x27)
MUTED = RGBColor(0x55, 0x55, 0x55)

TRACE_FILES = {
    "normal_login_vulnerable": "normal_login_vulnerable.json",
    "quote_login_vulnerable": "quote_login_vulnerable.json",
    "auth_logic_vulnerable": "auth_logic_vulnerable.json",
    "auth_logic_secure": "auth_logic_secure.json",
    "normal_login_secure": "normal_login_secure.json",
    "normal_search_vulnerable": "normal_search_vulnerable.json",
    "quote_search_vulnerable": "quote_search_vulnerable.json",
    "expanded_search_vulnerable": "expanded_search_vulnerable.json",
    "expanded_search_secure": "expanded_search_secure.json",
    "normal_search_secure": "normal_search_secure.json",
    "user_detail_vulnerable": "user_detail_vulnerable.json",
    "user_detail_secure": "user_detail_secure.json",
}

TEXT_EVIDENCE_KEYS = (
    "normal_login_vulnerable",
    "quote_login_vulnerable",
    "auth_logic_vulnerable",
    "auth_logic_secure",
    "normal_login_secure",
    "normal_search_vulnerable",
    "quote_search_vulnerable",
    "expanded_search_vulnerable",
    "expanded_search_secure",
    "normal_search_secure",
)

REPORT_FLOW_KEYS = (
    "normal_login_vulnerable",
    "quote_login_vulnerable",
    "auth_logic_vulnerable",
    "auth_logic_secure",
    "normal_login_secure",
    "normal_search_vulnerable",
    "quote_search_vulnerable",
    "expanded_search_vulnerable",
    "expanded_search_secure",
)

QUERY_FILES = {
    "vulnerable_login": "vulnerable_login_queries.json",
    "secure_login": "secure_login_queries.json",
    "vulnerable_search": "vulnerable_search_queries.json",
    "secure_search": "secure_search_queries.json",
}

SOURCE_SYMBOLS = (
    ("vulnerable_queries.py", "vulnerable_login"),
    ("secure_queries.py", "secure_login_lookup"),
    ("auth_service.py", "authenticate_secure"),
    ("vulnerable_queries.py", "vulnerable_search"),
    ("secure_queries.py", "secure_search"),
    ("validation.py", "validate_username"),
    ("error_service.py", "error_inspector"),
)

CHAPTER_TITLES = (
    "Giới thiệu",
    "Mục tiêu và phạm vi an toàn",
    "Cơ sở lý thuyết SQL và SQLite",
    "SQL Injection là gì",
    "Source, transform và SQL sink",
    "Kiến trúc Lab05",
    "Database schema và dữ liệu mẫu",
    "Phát hiện lỗi bằng input đơn giản",
    "Authentication flow bình thường",
    "Authentication logic bị thay đổi trong vulnerable mode",
    "Secure authentication với parameterized query và PBKDF2",
    "Search flow bình thường",
    "Search condition bị thay đổi trong vulnerable mode",
    "Secure search với LIKE parameter",
    "Error handling vulnerable và secure",
    "Prepared statement và parameter binding",
    "Password hashing",
    "ORM và raw SQL",
    "Input validation",
    "Least privilege",
    "Logging và monitoring",
    "WAF và giới hạn",
    "So sánh trước và sau vá",
    "Kết quả kiểm thử",
    "Coverage",
    "Runtime smoke test",
    "Mức độ ảnh hưởng",
    "Phòng chống",
    "Bài học rút ra",
    "Kết luận",
)

FLOW_LABELS = {
    "normal_login_vulnerable": "Normal vulnerable login",
    "quote_login_vulnerable": "Quote login",
    "auth_logic_vulnerable": "Local authentication logic demonstration",
    "auth_logic_secure": "Secure rejection",
    "normal_login_secure": "Secure normal login",
    "normal_search_vulnerable": "Normal vulnerable search",
    "quote_search_vulnerable": "Quote search",
    "expanded_search_vulnerable": "Expanded vulnerable search",
    "expanded_search_secure": "Secure search",
}

FLOW_OBJECTIVES = {
    "normal_login_vulnerable": "Xác nhận credential demo hợp lệ đi qua luồng legacy cố ý yếu.",
    "quote_login_vulnerable": "Quan sát dấu nháy đơn làm hỏng SQL nối chuỗi mà không lộ traceback.",
    "auth_logic_vulnerable": "Chứng minh input local cố định làm thay đổi điều kiện WHERE.",
    "auth_logic_secure": "Chứng minh cùng input chỉ là dữ liệu khi bind parameter.",
    "normal_login_secure": "Xác nhận lookup bằng placeholder và PBKDF2 cho credential hợp lệ.",
    "normal_search_vulnerable": "Xác nhận tìm kiếm bình thường trong bảng products local.",
    "quote_search_vulnerable": "Quan sát lỗi cú pháp do dấu nháy trong chuỗi LIKE nối trực tiếp.",
    "expanded_search_vulnerable": "Chứng minh điều kiện tìm kiếm bị mở rộng nhưng chỉ đọc products.",
    "expanded_search_secure": "Chứng minh LIKE parameter giữ nguyên cấu trúc và giới hạn kết quả.",
}

QUESTIONS_BY_CHAPTER = {
    4: ((1, "SQL Injection xảy ra ở tầng nào của ứng dụng?", "Lỗi xuất hiện tại ranh giới xây dựng truy vấn ở tầng ứng dụng: dữ liệu không tin cậy bị ghép thành cú pháp trước khi đi vào SQL parser. SQLite chỉ thực thi văn bản mà ứng dụng đã tạo."),),
    16: (
        (2, "Vì sao escaping thủ công dễ sai?", "Escaping thủ công phụ thuộc ngữ cảnh, encoding và dialect; chỉ một trường hợp bị bỏ sót cũng làm dữ liệu trở lại thành cú pháp. Parameter binding giao việc phân tách mã và dữ liệu cho driver."),
        (3, "Prepared statement khác nối chuỗi SQL như thế nào?", "Nối chuỗi tạo một văn bản SQL mới từ input. Prepared statement giữ template SQL cố định và truyền giá trị qua kênh parameter riêng."),
        (9, "Parameterized query xử lý dấu nháy đơn ra sao?", "Driver mã hóa dấu nháy như một phần của giá trị bound; dấu này không đóng literal và không đổi cây cú pháp của câu lệnh."),
    ),
    18: ((4, "ORM có tự động chống SQL Injection trong mọi trường hợp không?", "Không. ORM an toàn khi dùng API parameterized đúng cách; raw fragments, dynamic identifiers và escape tự chế vẫn có thể tái tạo lỗ hổng."),),
    15: ((5, "Vì sao không hiển thị lỗi SQL chi tiết cho người dùng?", "Thông báo chi tiết có thể lộ dialect, tên bảng, cột, đường dẫn hoặc cấu trúc truy vấn. UI chỉ trả thông báo chung; chi tiết đã rút gọn dành cho audit local."),),
    19: ((6, "Vì sao input validation không thay thế prepared statement?", "Validation kiểm soát định dạng nghiệp vụ, còn parameter binding bảo vệ cấu trúc SQL. Một giá trị hợp lệ về nghiệp vụ vẫn có thể chứa ký tự đặc biệt, nên cần cả hai lớp."),),
    17: (
        (7, "Vì sao password phải dùng PBKDF2, bcrypt hoặc Argon2?", "Các hàm này có salt và chi phí tính toán điều chỉnh được, làm chậm thử đoán hàng loạt. Hàm băm nhanh không có thuộc tính này."),
        (8, "Vì sao SHA-256 không salt không phù hợp lưu password?", "Cùng password sinh cùng digest, hỗ trợ bảng tra cứu và thử đoán tốc độ cao. Lab chỉ giữ digest legacy để minh họa mô hình cũ cố ý yếu."),
        (12, "Vì sao không lưu plaintext password?", "Lộ database sẽ lộ ngay credential và tạo nguy cơ tái sử dụng trên hệ thống khác. Chỉ verifier có salt và chi phí phù hợp được lưu."),
    ),
    22: ((10, "Vì sao WAF không thay thế sửa code?", "WAF chỉ là lớp lọc hỗ trợ và có thể bị bypass hoặc gây false positive. Bản vá gốc phải loại bỏ khả năng input trở thành cú pháp bằng parameter binding."),),
    20: ((11, "Least privilege giảm tác động ra sao?", "Giới hạn quyền, câu lệnh và dữ liệu truy cập làm giảm phạm vi hậu quả nếu một truy vấn bị điều khiển. SQLite lab mô phỏng bằng process non-root, file local và SELECT định nghĩa sẵn."),),
    5: ((13, "SQL Injection khác Parameter Tampering thế nào?", "SQL Injection thay đổi cách SQL parser hiểu câu lệnh; Parameter Tampering thay đổi giá trị tham số nghiệp vụ nhưng không nhất thiết thay đổi cú pháp SQL."),),
    10: ((14, "Authentication bypass xảy ra vì điều kiện WHERE thay đổi ra sao?", "Comment marker trong input local cố định làm phần kiểm tra digest phía sau không còn tham gia điều kiện cuối; bản ghi đầu tiên phù hợp có thể được trả về."),),
    21: ((15, "Vì sao log không được chứa password hoặc full query nhạy cảm?", "Log thường được sao chép và giữ lâu. Password, hash đầy đủ hoặc truy vấn chứa secret sẽ biến hệ thống quan sát thành nguồn rò rỉ thứ hai."),),
}

SEQUENCE_SPECS = (
    ("Sơ đồ tuần tự 1. Vulnerable login", "auth_logic_vulnerable"),
    ("Sơ đồ tuần tự 2. Secure login", "auth_logic_secure"),
    ("Sơ đồ tuần tự 3. Vulnerable search", "expanded_search_vulnerable"),
    ("Sơ đồ tuần tự 4. Secure search", "expanded_search_secure"),
)

DATA_FLOW_SPECS = (
    (
        "Sơ đồ luồng dữ liệu 1. Vulnerable data flow",
        ("Untrusted Input", "Flask Request", "String Concatenation", "SQL Text", "SQLite Parser", "Unexpected Result"),
    ),
    (
        "Sơ đồ luồng dữ liệu 2. Secure data flow",
        ("Untrusted Input", "Flask Request", "Validation", "SQL Template", "Parameter Binding", "SQLite Parser", "Expected Result"),
    ),
)


class EvidenceError(RuntimeError):
    """Raised when required observed evidence is missing or inconsistent."""


def _required_paths(root: Path = ROOT) -> dict[str, Path]:
    paths: dict[str, Path] = {
        "audit": root / "evidence/audit/audit_logs.json",
        "pytest": root / "evidence/logs/pytest.txt",
        "coverage": root / "evidence/logs/coverage.txt",
        "smoke": root / "evidence/logs/runtime_smoke_test.txt",
    }
    paths.update({f"trace:{key}": root / "evidence/traces" / filename for key, filename in TRACE_FILES.items()})
    paths.update({f"request:{key}": root / "evidence/requests" / f"{key}.txt" for key in TEXT_EVIDENCE_KEYS})
    paths.update({f"response:{key}": root / "evidence/responses" / f"{key}.txt" for key in TEXT_EVIDENCE_KEYS})
    paths.update({f"query:{key}": root / "evidence/queries" / filename for key, filename in QUERY_FILES.items()})
    for filename, _symbol in SOURCE_SYMBOLS:
        paths[f"source:{filename}"] = root / filename
    paths["source:schema.sql"] = root / "schema.sql"
    paths["source:seed.py"] = root / "seed.py"
    missing = [str(path.relative_to(root)) for path in paths.values() if not path.is_file() or path.stat().st_size == 0]
    if missing:
        raise FileNotFoundError("Missing real Lab05 evidence/source: " + ", ".join(sorted(set(missing))))
    return paths


def _read_json(path: Path):
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (UnicodeDecodeError, json.JSONDecodeError) as exc:
        raise EvidenceError(f"Invalid JSON evidence: {path}") from exc


def _records(value, label: str) -> list[dict]:
    if isinstance(value, list):
        records = value
    elif isinstance(value, dict):
        records = next((value[key] for key in ("records", "queries", "events", "items") if isinstance(value.get(key), list)), [value])
    else:
        raise EvidenceError(f"{label} must contain a JSON object or list")
    if not records or not all(isinstance(record, dict) for record in records):
        raise EvidenceError(f"{label} contains no usable records")
    return records


def _scan_evidence_for_secrets(label: str, text: str) -> None:
    forbidden = ("AdminLab123!", "StudentA123!", "StudentB123!")
    found = next((secret for secret in forbidden if secret in text), None)
    if found or re.search(r"(?i)\b[0-9a-f]{64}\b", text) or re.search(r"pbkdf2:[^\s\"]+\$[^\s\"]+\$", text):
        raise EvidenceError(f"Sensitive value is not redacted in {label}")


def _validate_trace(key: str, trace: dict) -> None:
    if not isinstance(trace, dict) or not trace.get("trace_id") or not isinstance(trace.get("steps"), list):
        raise EvidenceError(f"Trace {key} is incomplete")
    required = {"step_number", "layer", "technique", "input_data", "output_data", "code_reference", "status", "security_meaning"}
    for index, step in enumerate(trace["steps"], 1):
        if not isinstance(step, dict) or not (step.get("title") or step.get("action")):
            raise EvidenceError(f"Trace {key} step {index} has no action/title")
        missing = sorted(name for name in required if name not in step)
        if missing:
            raise EvidenceError(f"Trace {key} step {index} misses: {', '.join(missing)}")


def _parse_pytest(text: str) -> dict:
    passed = sum(int(value) for value in re.findall(r"(?<!\d)(\d+) passed\b", text, flags=re.IGNORECASE))
    failed = sum(int(value) for value in re.findall(r"(?<!\d)(\d+) failed\b", text, flags=re.IGNORECASE))
    errors = sum(int(value) for value in re.findall(r"(?<!\d)(\d+) errors?\b", text, flags=re.IGNORECASE))
    if passed == failed == errors == 0:
        raise EvidenceError("pytest.txt has no terminal pytest summary")
    return {"passed": passed, "failed": failed, "errors": errors, "status": "passed" if passed and not failed and not errors else "failed"}


def _parse_coverage(text: str) -> dict:
    rows = {}
    for line in text.splitlines():
        match = re.match(r"\s*(\S+\.py|TOTAL)\s+\d+\s+(?:\d+\s+)?(\d{1,3})%", line)
        if match:
            rows[match.group(1)] = int(match.group(2))
    if not rows:
        raise EvidenceError("coverage.txt has no parseable coverage rows")
    return {"rows": rows, "total": rows.get("TOTAL"), "status": "observed"}


def _parse_smoke(text: str) -> dict:
    failed = len(re.findall(r"(?im)^FAIL\s*\|", text))
    passed = len(re.findall(r"(?im)^PASS\s*\|", text))
    if not passed and not failed:
        raise EvidenceError("runtime_smoke_test.txt has no PASS/FAIL result")
    return {"passed_checks": passed, "failed_checks": failed, "status": "passed" if passed and not failed else "failed"}


def _normalise_query_record(record: dict, label: str) -> dict:
    normal = dict(record)
    normal["construction_type"] = record.get("construction_type", record.get("construction_method"))
    normal["rows_returned"] = record.get("rows_returned", record.get("result_count"))
    required = ("query_template", "construction_type", "final_query_masked", "parameters_masked", "rows_returned", "error_category", "trace_id")
    missing = [name for name in required if name not in normal or (normal[name] is None and name not in {"error_category"})]
    if missing:
        raise EvidenceError(f"{label} query record misses: {', '.join(missing)}")
    return normal


def _load_evidence(root: Path = ROOT) -> dict:
    paths = _required_paths(root)
    traces = {}
    for key in TRACE_FILES:
        trace = _read_json(paths[f"trace:{key}"])
        _validate_trace(key, trace)
        _scan_evidence_for_secrets(f"trace:{key}", json.dumps(trace, ensure_ascii=False))
        traces[key] = trace

    requests = {key: paths[f"request:{key}"].read_text(encoding="utf-8", errors="strict") for key in TEXT_EVIDENCE_KEYS}
    responses = {key: paths[f"response:{key}"].read_text(encoding="utf-8", errors="strict") for key in TEXT_EVIDENCE_KEYS}
    for key in TEXT_EVIDENCE_KEYS:
        trace_id = str(traces[key]["trace_id"])
        request_text, response_text = requests[key], responses[key]
        _scan_evidence_for_secrets(f"request:{key}", request_text)
        _scan_evidence_for_secrets(f"response:{key}", response_text)
        request_panel = traces[key].get("request_inspector", {})
        request_markers = (
            f"Method: {request_panel.get('method', '')}",
            f"Path: {request_panel.get('path', '')}",
            "Timestamp:",
        )
        if any(marker not in request_text for marker in request_markers):
            raise EvidenceError(f"request:{key} does not match its observed trace request")
        if trace_id not in response_text:
            raise EvidenceError(f"response:{key} is not linked to trace ID {trace_id}")

    query_records = []
    for key in QUERY_FILES:
        raw_records = _records(_read_json(paths[f"query:{key}"]), f"query:{key}")
        for record in raw_records:
            normal = _normalise_query_record(record, f"query:{key}")
            _scan_evidence_for_secrets(f"query:{key}", json.dumps(normal, ensure_ascii=False))
            query_records.append(normal)

    audits = _records(_read_json(paths["audit"]), "audit")
    _scan_evidence_for_secrets("audit", json.dumps(audits, ensure_ascii=False))
    query_by_trace = {str(record["trace_id"]): record for record in query_records}
    audits_by_trace: dict[str, list[dict]] = {}
    for audit in audits:
        trace_id = str(audit.get("trace_id", ""))
        if trace_id:
            audits_by_trace.setdefault(trace_id, []).append(audit)

    flows = {}
    for key in REPORT_FLOW_KEYS:
        trace_id = str(traces[key]["trace_id"])
        if trace_id not in query_by_trace or trace_id not in audits_by_trace:
            raise EvidenceError(f"Flow {key} is not linked across trace, query and audit evidence")
        flows[key] = {
            "key": key,
            "trace": traces[key],
            "request": requests[key],
            "response": responses[key],
            "query": query_by_trace[trace_id],
            "audits": audits_by_trace[trace_id],
        }

    pytest_text = paths["pytest"].read_text(encoding="utf-8", errors="replace")
    coverage_text = paths["coverage"].read_text(encoding="utf-8", errors="replace")
    smoke_text = paths["smoke"].read_text(encoding="utf-8", errors="replace")
    return {
        "paths": paths,
        "traces": traces,
        "requests": requests,
        "responses": responses,
        "query_records": query_records,
        "audits": audits,
        "flows": flows,
        "pytest_text": pytest_text,
        "coverage_text": coverage_text,
        "smoke_text": smoke_text,
        "pytest": _parse_pytest(pytest_text),
        "coverage": _parse_coverage(coverage_text),
        "smoke": _parse_smoke(smoke_text),
    }


def _compact(value, limit: int = 120) -> str:
    text = json.dumps(value, ensure_ascii=False, separators=(",", ":")) if isinstance(value, (dict, list)) else str(value)
    text = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text if len(text) <= limit else text[: limit - 3] + "..."


def _source_function(root: Path, filename: str, function: str) -> tuple[int, int, str]:
    path = root / filename
    source = path.read_text(encoding="utf-8")
    node = next((item for item in ast.walk(ast.parse(source)) if isinstance(item, (ast.FunctionDef, ast.AsyncFunctionDef)) and item.name == function), None)
    if node is None or node.end_lineno is None:
        raise EvidenceError(f"Source symbol not found: {filename}:{function}")
    snippet = "\n".join(source.splitlines()[node.lineno - 1 : node.end_lineno])
    return node.lineno, node.end_lineno, snippet


def _set_run_font(run, size: float = 11, bold=None, color=None, italic=None, name: str = "Calibri") -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name)
    rpr.rFonts.set(qn("w:hAnsi"), name)
    rpr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = tr_pr.find(qn("w:tblHeader"))
    if marker is None:
        marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    if marker.getparent() is None:
        tr_pr.append(marker)


def _set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
    if tc_mar.getparent() is None:
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        if node.getparent() is None:
            tc_mar.append(node)


def _set_table_geometry(
    table,
    widths_dxa: list[int] | tuple[int, ...],
    indent_dxa: int = 120,
    cell_margin_tb: int = 80,
) -> None:
    if len(widths_dxa) != len(table.columns) or any(width <= 0 for width in widths_dxa):
        raise ValueError("Table width count must match positive table columns")
    total = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr
    for tag, attrs in (
        ("w:tblW", {"w:w": str(total), "w:type": "dxa"}),
        ("w:tblInd", {"w:w": str(indent_dxa), "w:type": "dxa"}),
        ("w:tblLayout", {"w:type": "fixed"}),
    ):
        node = tbl_pr.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
        for key, value in attrs.items():
            node.set(qn(key), value)
        if node.getparent() is None:
            tbl_pr.append(node)

    old_grid = table._tbl.tblGrid
    grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    table._tbl.replace(old_grid, grid)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            if tc_w.getparent() is None:
                tc_pr.append(tc_w)
            _set_cell_margins(cell, top=cell_margin_tb, bottom=cell_margin_tb)


def _shade(cell, fill: str = "F2F4F7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    if node.getparent() is None:
        tc_pr.append(node)


def _add_table(
    doc,
    headers,
    rows,
    widths_dxa,
    font_size: float = 8,
    header_fill: str = "F2F4F7",
    indent_dxa: int = 120,
    cell_margin_tb: int = 80,
):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.style = "Table Grid"
    _set_repeat_header(table.rows[0])
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _shade(cell, header_fill)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(paragraph.add_run(str(text)), font_size, bold=True, color=DARK_BLUE)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cell = cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index == 0 or len(str(value)) < 16 else WD_ALIGN_PARAGRAPH.LEFT
            _set_run_font(paragraph.add_run(str(value)), font_size)
    _set_table_geometry(table, widths_dxa, indent_dxa, cell_margin_tb)
    return table


def _add_page_field(paragraph) -> None:
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def _set_header_footer(section, label: str) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    _set_run_font(header.add_run(label), 8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_after = Pt(0)
    _set_run_font(footer.add_run("Trang "), 8.5, color=MUTED)
    _add_page_field(footer)


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True
    _set_header_footer(section, "LAB05 · SQL INJECTION · 21127645")
    section.first_page_header.paragraphs[0].text = ""
    section.first_page_footer.paragraphs[0].text = ""

    normal = doc.styles["Normal"]
    normal.font.name, normal.font.size = "Calibri", Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name, style.font.size, style.font.color.rgb = "Calibri", Pt(size), color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    doc.core_properties.title = "LAB 5 - SQL INJECTION"
    doc.core_properties.subject = "Báo cáo thực hành bảo mật ứng dụng web từ evidence local"
    doc.core_properties.author = "21127645 - Lê Minh"
    doc.core_properties.keywords = "SQL Injection, SQLite, parameterized query, PBKDF2, audit, trace"


def _add_cover(doc: Document) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(118)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    _set_run_font(p.add_run("BÁO CÁO THỰC HÀNH BẢO MẬT ỨNG DỤNG WEB"), 11, bold=True, color=GOLD)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    _set_run_font(p.add_run("LAB 5 · SQL INJECTION"), 30, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    _set_run_font(p.add_run("String Concatenation · Parameter Binding · PBKDF2"), 15, color=TEAL)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(70)
    _set_run_font(p.add_run("Evidence-driven local security lab"), 10.5, italic=True, color=GOLD)
    for label, value in (
        ("Sinh viên", "Lê Minh"),
        ("MSSV", "21127645"),
        ("Môi trường", "Flask + SQLite · http://127.0.0.1:5005"),
        ("Phạm vi", "Dữ liệu giả lập · SELECT-only · Không kết nối Internet"),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        _set_run_font(p.add_run(f"{label}: "), 10.5, bold=True, color=NAVY)
        _set_run_font(p.add_run(value), 10.5, color=MUTED)
    doc.add_page_break()


def _add_toc(doc: Document) -> None:
    doc.add_heading("Mục lục báo cáo", level=1)
    rows = []
    for index in range(15):
        left = (index + 1, CHAPTER_TITLES[index])
        right = (index + 16, CHAPTER_TITLES[index + 15])
        rows.append([left[0], left[1], right[0], right[1]])
    _add_table(doc, ["Chương", "Nội dung", "Chương", "Nội dung"], rows, [650, 4030, 650, 4030], 8)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    _set_run_font(p.add_run("Danh mục sơ đồ: 4 sơ đồ tuần tự và 2 sơ đồ luồng dữ liệu Word-native. Danh mục bảng: 9 bảng trace, code comparison, test, coverage, smoke và evidence index."), 9, color=MUTED)
    doc.add_page_break()


def _trace_value(trace: dict, panel: str, key: str, default="-"):
    value = trace.get(panel, {})
    return value.get(key, default) if isinstance(value, dict) else default


def _decision(trace: dict) -> str:
    return str(_trace_value(trace, "decision_inspector", "final_decision", trace.get("final_result", "-")))


def _rows_returned(trace: dict) -> int:
    value = _trace_value(trace, "execution_inspector", "rows_returned", 0)
    try:
        return int(value)
    except (TypeError, ValueError):
        return 0


def _chapter_content(evidence: dict) -> list[str]:
    traces = evidence["traces"]
    pytest = evidence["pytest"]
    coverage = evidence["coverage"]
    smoke = evidence["smoke"]
    total = coverage.get("total")
    return [
        "Lab05 là ứng dụng học tập độc lập minh họa nguyên nhân SQL Injection và bản vá tại tầng xây dựng truy vấn. Báo cáo được tạo từ request, response, query event, trace, audit, log kiểm thử và source code đã tồn tại; ảnh chụp thủ công không phải dependency của báo cáo.",
        "Runtime chỉ bind 127.0.0.1:5005, dùng SQLite nằm trong Lab05 và các scenario local cố định. Ứng dụng không nhận host, URL, port, connection string, đường dẫn database hay câu SQL tùy ý; không có scanner, UNION, blind, time-based, stacked query hoặc schema enumeration.",
        "SQL mô tả phép đọc dữ liệu theo cấu trúc do parser hiểu. SQLite là database nhúng một file, phù hợp lab local nhưng vẫn phân biệt rõ SQL text và bound value. SELECT-only không loại bỏ SQL Injection; nó chỉ giới hạn tác động ghi trong mô hình này.",
        "SQL Injection xảy ra khi input không tin cậy tham gia cấu tạo cú pháp SQL. Trong vulnerable mode, string concatenation đặt input vào SQL text trước khi SQLite parse; trong secure mode, template và parameter đi theo hai kênh riêng.",
        "Source là request.form hoặc request.args; transform gồm normalize, legacy digest hoặc wildcard LIKE; sink là execute_read_only. Ranh giới nguy hiểm nằm ở chỗ transform bằng nối chuỗi có thể đổi cấu trúc, còn parameter binding giữ input ở vai trò dữ liệu.",
        "Flask routes gọi validation, auth/query service, SQLite, trace và audit. Trace ID liên kết Request Inspector, Query Construction Inspector, SQL Execution Inspector, Authentication/Result Set Inspector, Database Inspector, Error Inspector, response và audit event.",
        "Schema thật định nghĩa users, products, audit_logs, login_attempts, query_events và trace_records. Seed tạo 3 user và 8 product giả lập; secure password dùng PBKDF2, còn legacy digest chỉ phục vụ luồng yếu cố ý và không được secure authentication sử dụng.",
        f"Quote login có quyết định {_decision(traces['quote_login_vulnerable'])}; quote search có quyết định {_decision(traces['quote_search_vulnerable'])}. Hai trace ghi Error Inspector thay vì trả traceback, full path hay thông báo SQLite chi tiết ra UI.",
        f"Normal vulnerable login trả {_rows_returned(traces['normal_login_vulnerable'])} row và quyết định {_decision(traces['normal_login_vulnerable'])}. Server băm password demo bằng legacy SHA-256 rồi ghép username và digest vào câu SQL cố ý yếu.",
        f"Input local cố định trong authentication trace tạo quyết định {_decision(traces['auth_logic_vulnerable'])}. Query Inspector và Final Verdict ghi nhận cấu trúc WHERE bị thay đổi; đây chỉ là hành vi của database giả lập local, không phải công cụ nhắm mục tiêu bên ngoài.",
        f"Cùng input trong secure trace có quyết định {_decision(traces['auth_logic_secure'])}. Lookup dùng WHERE username = ?, sau đó check_password_hash xác minh PBKDF2 và chỉ tạo session khi verifier thành công; lỗi trả thông báo credential chung.",
        f"Normal vulnerable search trả {_rows_returned(traces['normal_search_vulnerable'])} row từ products. Luồng bình thường cho thấy chức năng hợp lệ trước khi so sánh với quote và expanded scenario.",
        f"Expanded vulnerable search trả {_rows_returned(traces['expanded_search_vulnerable'])} row và quyết định {_decision(traces['expanded_search_vulnerable'])}. Result Set Inspector xác nhận chỉ products được đọc, không truy cập users và không thay đổi database.",
        f"Secure search với cùng input trả {_rows_returned(traces['expanded_search_secure'])} row và quyết định {_decision(traces['expanded_search_secure'])}; normal secure search trả {_rows_returned(traces['normal_search_secure'])} row với quyết định {_decision(traces['normal_search_secure'])}. SQL giữ LIKE ? và LIMIT 50; wildcard là một phần của bound parameter thay vì SQL text.",
        "Vulnerable error evidence giữ exception category và nguyên nhân kỹ thuật đã rút gọn cho Inspector. Secure route xem dấu nháy là dữ liệu và không lộ SQLite detail; cả hai đường đều không lộ traceback hoặc absolute database path.",
        "Prepared statement tách bước compile/parse template khỏi giá trị. sqlite3 placeholder ? và tuple parameters là cơ chế thật được dùng; không có replace, manual escaping hoặc parameterized query giả.",
        "Werkzeug generate_password_hash dùng pbkdf2:sha256:600000 với salt; check_password_hash thực hiện verifier. Database Inspector chỉ được phép trình bày thuật toán, độ dài và fingerprint rút gọn, không đưa full hash vào báo cáo.",
        "Lab không thêm ORM vì sqlite3 parameter binding đã đáp ứng phạm vi. ORM chỉ an toàn khi API tạo bound parameters; raw SQL, dynamic identifier hoặc literal fragment vẫn cần review như SQL viết tay.",
        f"Validation giới hạn độ dài, kiểu integer và chỉ cho phép scenario cố định ở route vulnerable. User detail vulnerable có quyết định {_decision(traces['user_detail_vulnerable'])}; secure user detail validate integer, bind placeholder và có quyết định {_decision(traces['user_detail_secure'])}. Validation giảm bề mặt nhưng bản vá cấu trúc vẫn là parameter binding.",
        "SQLite không có user/role permission model như server database. Lab mô phỏng least privilege trung thực bằng file local, process non-root trong Docker, query định nghĩa sẵn, SELECT-only cho input và không cung cấp database path hay connection string.",
        f"Audit evidence có {len(evidence['audits'])} record liên kết trace_id; query evidence có {len(evidence['query_records'])} record. Log giữ template, parameter count, decision, reason, row count và error category, đồng thời redaction password và digest đầy đủ.",
        "WAF có thể phát hiện pattern phổ biến nhưng không hiểu đầy đủ ngữ cảnh code và có thể bị né tránh. Nó là lớp hỗ trợ cho monitoring; không thay thế parameterized query, password hashing, validation, least privilege và error handling.",
        "Code Comparison dưới đây được trích bằng ast.parse từ function đang chạy. Bản vulnerable nối input vào SQL; bản secure giữ SQL constant, truyền tuple parameter và tách password verification khỏi truy vấn.",
        f"pytest.txt ghi {pytest['passed']} passed, {pytest['failed']} failed và {pytest['errors']} errors; trạng thái dẫn xuất là {pytest['status']}. Báo cáo không tự đổi kết quả này thành thành công nếu log có failure.",
        f"coverage.txt có {len(coverage['rows'])} dòng module parse được; TOTAL là {total if total is not None else 'không có dòng TOTAL'}%. Các giá trị trong bảng coverage được đọc trực tiếp từ log, không ghi cứng.",
        f"Runtime smoke evidence có {smoke['passed_checks']} PASS và {smoke['failed_checks']} FAIL; trạng thái dẫn xuất là {smoke['status']}. Smoke chỉ gọi loopback và không thay thế quyết định SameSite hoặc hành vi trình duyệt thực.",
        "Trong hệ thống thật, SQL Injection có thể dẫn đến đọc trái phép, bypass authentication, sửa/xóa dữ liệu hoặc leo thang tùy quyền database. Lab thu hẹp thành SELECT local để quan sát nguyên nhân mà không mở rộng năng lực tấn công.",
        "Phòng chống ưu tiên parameter binding ở mọi giá trị, allowlist cho identifier không bind được, PBKDF2/bcrypt/Argon2 cho password, error chung, least privilege, audit có redaction, review raw SQL và kiểm thử regression qua public route.",
        "Bài học chính là validation, WAF và ORM không thể bù cho truy vấn nối chuỗi. Evidence tốt phải liên kết request -> trace -> query -> audit -> response và chỉ kết luận từ flow đã chạy.",
        "Lab05 chứng minh cùng input có ý nghĩa khác nhau tùy construction method: vulnerable SQL có thể đổi logic hoặc lỗi parse, còn secure SQL giữ cấu trúc. Bản vá là parameter binding kết hợp PBKDF2, safe errors, least privilege và monitoring.",
    ]


def _add_question(doc: Document, number: int, question: str, answer: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    _set_run_font(p.add_run(f"Câu hỏi {number}. {question} "), 9.3, bold=True, color=DARK_BLUE)
    _set_run_font(p.add_run(answer), 9.3)


def _add_chapter(doc: Document, number: int, title: str, paragraph: str) -> None:
    doc.add_heading(f"Chương {number}. {title}", level=1)
    p = doc.add_paragraph(paragraph)
    p.paragraph_format.keep_together = True
    for question in QUESTIONS_BY_CHAPTER.get(number, ()):
        _add_question(doc, *question)


def _add_code_comparison(doc: Document, root: Path) -> None:
    rows = []
    for filename, function in SOURCE_SYMBOLS:
        start, end, snippet = _source_function(root, filename, function)
        rows.append([filename, function, f"{start}-{end}", _compact(snippet, 330)])
    doc.add_heading("Code Comparison trích từ source thật", level=2)
    _add_table(doc, ["File", "Function", "Lines", "Source excerpt"], rows, [1300, 1700, 900, 5460], 6.4)


def _add_data_flow_diagram(doc: Document, title: str, nodes: tuple[str, ...]) -> None:
    doc.add_heading(title, level=3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    _set_run_font(p.add_run(" -> ".join(nodes)), 9, bold=True, color=NAVY)
    rows = [[index, node, "Dữ liệu" if index == 1 else "Chuyển tiếp có kiểm soát" if "Binding" in node or "Validation" in node else "Biến đổi/diễn giải"] for index, node in enumerate(nodes, 1)]
    _add_table(doc, ["Bước", "Node", "Security meaning"], rows, [700, 2860, 5800], 7.5, "E8EEF5")


def _add_sequence_diagram(doc: Document, title: str, trace: dict) -> None:
    doc.add_heading(title, level=3)
    actors = ("User", "Browser", "Flask", "Query Builder", "SQLite", "Password Verifier", "Session")
    _add_table(doc, actors, [["actor"] * len(actors)], [1337, 1337, 1337, 1337, 1337, 1337, 1338], 6.4, "E8EEF5")
    request = trace.get("request_inspector", {})
    query = trace.get("query_inspector", {})
    decision = trace.get("decision_inspector", {})
    raw_input = _trace_value(trace, "input_inspector", "raw_input", "-")
    messages = [
        [1, "User -> Browser", "Submit fixed local input", _compact(raw_input, 60)],
        [2, "Browser -> Flask", f"{request.get('method', '-')} {request.get('path', '-')}", trace.get("trace_id", "-")],
        [3, "Flask -> Query Builder", str(query.get("construction_method", "-")), _compact(query.get("query_template", "-"), 85)],
        [4, "Query Builder -> SQLite", "Execute SELECT", _compact(query.get("final_query_masked", "-"), 85)],
        [5, "SQLite -> Flask", "Return result/error", f"rows={_rows_returned(trace)}"],
    ]
    if trace.get("feature") == "login":
        messages.append([6, "Flask -> Password Verifier", "Verify only in secure mode", str(decision.get("password_verification_executed", False))])
        messages.append([7, "Flask -> Session", "Create only after allow", str(decision.get("session_created", False))])
    messages.append([len(messages) + 1, "Flask -> Browser", "Safe response", _decision(trace)])
    _add_table(doc, ["STT", "From -> To", "Message", "Observed evidence"], messages, [600, 2300, 2460, 4000], 7.0)


def _flow_summary_rows(flow: dict) -> list[list[str]]:
    trace, query = flow["trace"], flow["query"]
    input_panel = trace.get("input_inspector", {})
    verdict = trace.get("final_verdict", {})
    audit = flow["audits"][0]
    code_refs = sorted({str(step.get("code_reference", "-")) for step in trace.get("steps", [])})
    impact = "authentication bypass" if verdict.get("authentication_bypassed") else "unexpected product rows" if verdict.get("unexpected_rows_returned") else "query error observed" if verdict.get("database_error_occurred") else "expected local result"
    patch = "Dùng SQL template + parameter binding; login secure tách PBKDF2 verification." if trace.get("mode") == "secure" else "Đối chiếu secure flow dùng parameter binding và safe error."
    return [
        ["Mục tiêu", FLOW_OBJECTIVES[flow["key"]], "Điều kiện đầu", "Database seed local; fixed scenario; SELECT-only."],
        ["Input", _compact(input_panel.get("raw_input", "-"), 110), "Request", _compact(flow["request"], 130)],
        ["Source/normalize", _compact(f"{input_panel.get('source', '-')} -> {input_panel.get('normalized_input', '-')}", 120), "Query template", _compact(query.get("query_template", "-"), 130)],
        ["Construction", str(query.get("construction_type", "-")), "Final query masked", _compact(query.get("final_query_masked", "-"), 130)],
        ["SQLite", _compact(_trace_value(trace, "query_inspector", "input_interpreted_as", "-"), 100), "Rows", str(query.get("rows_returned", _rows_returned(trace)))],
        ["Decision", _decision(trace), "Response", _compact(flow["response"], 130)],
        ["State change", f"database_modified={verdict.get('database_modified', False)}", "Root cause", str(verdict.get("root_cause", "-"))],
        ["Impact", impact, "Code vulnerable/secure", _compact(" | ".join(code_refs), 130)],
        ["Bản vá", patch, "Trace JSON", str(trace.get("trace_id", "-"))],
        ["Audit event", _compact({key: audit.get(key) for key in ("action", "decision", "reason", "trace_id")}, 125), "Final verdict", _compact(verdict, 130)],
    ]


def _trace_rows(trace: dict) -> list[list[str]]:
    rows = []
    for step in trace["steps"]:
        rows.append([
            step["step_number"],
            _compact(step["layer"], 24),
            _compact(step.get("action", step.get("title", "-")), 32),
            _compact(step["technique"], 36),
            _compact(step["input_data"], 44),
            _compact(step["output_data"], 44),
            _compact(step["code_reference"], 38),
            _compact(step["status"], 16),
            _compact(step["security_meaning"], 58),
        ])
    return rows


def _add_flow_page(doc: Document, flow: dict, index: int) -> None:
    trace = flow["trace"]
    doc.add_heading(f"Phụ lục Trace {index}. {FLOW_LABELS[flow['key']]}", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    _set_run_font(p.add_run(f"Trace ID: {trace['trace_id']} · Mode: {trace.get('mode', '-')} · Feature: {trace.get('feature', '-')} · Decision: {_decision(trace)}"), 7.2, bold=True, color=NAVY)
    _add_table(doc, ["Field", "Observed value", "Field", "Observed value"], _flow_summary_rows(flow), [1150, 6050, 1150, 6050], 5.8, "E8EEF5")
    caption = doc.add_paragraph()
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(2)
    caption.paragraph_format.keep_with_next = True
    _set_run_font(caption.add_run("Trace table - các bước quan sát được"), 7, bold=True, color=DARK_BLUE)
    _add_table(
        doc,
        ["STT", "Layer", "Action", "Technique", "Input", "Output", "Code reference", "Status", "Security meaning"],
        _trace_rows(trace),
        [420, 1200, 1700, 1600, 1900, 1900, 1700, 800, 3180],
        5.6,
        cell_margin_tb=30,
    )


def _add_evidence_appendix(doc: Document, evidence: dict) -> None:
    doc.add_heading("Phụ lục A. Evidence và lệnh tái lập", level=1)
    rows = [
        ["Trace JSON", len(evidence["traces"]), "evidence/traces/*.json"],
        ["Request", len(evidence["requests"]), "evidence/requests/*.txt"],
        ["Response", len(evidence["responses"]), "evidence/responses/*.txt"],
        ["Query evidence", len(evidence["query_records"]), "evidence/queries/*.json"],
        ["Audit event", len(evidence["audits"]), "evidence/audit/audit_logs.json"],
        ["Code snippets", len(SOURCE_SYMBOLS), "AST extraction from current source"],
        ["Lệnh chạy", 5, "seed.py; pytest; coverage; runtime smoke; generate_report.py"],
    ]
    _add_table(doc, ["Loại", "Số lượng", "Nguồn thật"], rows, [1800, 1100, 6460], 8)
    p = doc.add_paragraph()
    _set_run_font(p.add_run("Lệnh tái lập: python seed.py; python -m pytest; python -m pytest --cov=. --cov-report=term-missing; python scripts/run_runtime_smoke_test.py; python scripts/generate_report.py."), 8.5)


def _new_landscape_section(doc: Document):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = Inches(11), Inches(8.5)
    section.top_margin = section.bottom_margin = Inches(0.55)
    section.left_margin = section.right_margin = Inches(0.5)
    section.header_distance = section.footer_distance = Inches(0.3)
    _set_header_footer(section, "LAB05 · SQL INJECTION · TRACE EVIDENCE · 21127645")
    return section


def _validate_docx_structure(path: Path) -> None:
    if not path.is_file() or path.stat().st_size < 20_000:
        raise RuntimeError("DOCX generation failed or output is unexpectedly small")
    with zipfile.ZipFile(path) as archive:
        if archive.testzip() is not None:
            raise RuntimeError("DOCX ZIP package is corrupt")
        xml = archive.read("word/document.xml").decode("utf-8")
    required_text = [f"Chương {index}. {title}" for index, title in enumerate(CHAPTER_TITLES, 1)]
    required_text += [title for title, _key in SEQUENCE_SPECS]
    required_text += [title for title, _nodes in DATA_FLOW_SPECS]
    required_text += [FLOW_LABELS[key] for key in REPORT_FLOW_KEYS]
    missing = [text for text in required_text if text not in xml]
    if missing:
        raise RuntimeError("DOCX is missing required report content: " + ", ".join(missing))


def build_docx(evidence: dict | None = None, root: Path = ROOT) -> Path:
    evidence = evidence or _load_evidence(root)
    doc = Document()
    _configure_document(doc)
    _add_cover(doc)
    _add_toc(doc)
    paragraphs = _chapter_content(evidence)

    for index in range(1, 8):
        _add_chapter(doc, index, CHAPTER_TITLES[index - 1], paragraphs[index - 1])
    doc.add_page_break()
    doc.add_heading("Sơ đồ luồng dữ liệu", level=2)
    for title, nodes in DATA_FLOW_SPECS:
        _add_data_flow_diagram(doc, title, nodes)
    doc.add_page_break()

    for index in range(8, 16):
        _add_chapter(doc, index, CHAPTER_TITLES[index - 1], paragraphs[index - 1])
    doc.add_page_break()
    doc.add_heading("Sơ đồ tuần tự từ trace thật", level=2)
    for diagram_index, (title, trace_key) in enumerate(SEQUENCE_SPECS, 1):
        _add_sequence_diagram(doc, title, evidence["traces"][trace_key])
        if diagram_index == 2:
            doc.add_page_break()
    doc.add_page_break()

    for index in range(16, 31):
        _add_chapter(doc, index, CHAPTER_TITLES[index - 1], paragraphs[index - 1])
        if index == 23:
            _add_code_comparison(doc, root)
        elif index == 25:
            coverage_rows = [[name, value, "Observed"] for name, value in evidence["coverage"]["rows"].items()]
            _add_table(doc, ["Module", "Coverage %", "Nguồn"], coverage_rows, [5000, 1500, 2860], 7.5)
        elif index == 26:
            smoke = evidence["smoke"]
            _add_table(doc, ["PASS", "FAIL", "Status"], [[smoke["passed_checks"], smoke["failed_checks"], smoke["status"]]], [2500, 2500, 4360], 8)

    _add_evidence_appendix(doc, evidence)
    _new_landscape_section(doc)
    for index, key in enumerate(REPORT_FLOW_KEYS, 1):
        _add_flow_page(doc, evidence["flows"][key], index)
        if index < len(REPORT_FLOW_KEYS):
            doc.add_page_break()

    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)
    _validate_docx_structure(DOCX_PATH)
    return DOCX_PATH


def _find_soffice() -> Path:
    candidates = [
        shutil.which("soffice"),
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    path = next((Path(candidate) for candidate in candidates if candidate and Path(candidate).is_file()), None)
    if path is None:
        raise RuntimeError("PDF generation requires LibreOffice/soffice; no flattening fallback is allowed")
    return path


def build_pdf(docx_path: Path = DOCX_PATH) -> Path:
    if not docx_path.is_file():
        raise RuntimeError(f"DOCX does not exist: {docx_path}")
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    if PDF_PATH.exists():
        PDF_PATH.unlink()
    soffice = _find_soffice()
    with tempfile.TemporaryDirectory(prefix="lab05-lo-") as profile:
        profile_uri = Path(profile).resolve().as_uri()
        result = subprocess.run(
            [str(soffice), "--headless", f"-env:UserInstallation={profile_uri}", "--convert-to", "pdf", "--outdir", str(REPORT_DIR), str(docx_path)],
            capture_output=True,
            text=True,
            timeout=180,
            check=False,
        )
    if result.returncode or not PDF_PATH.is_file() or PDF_PATH.stat().st_size < 10_000:
        detail = (result.stderr or result.stdout or "no converter output").strip()
        raise RuntimeError(f"PDF conversion failed; report is not complete: {detail}")
    return PDF_PATH


def _count_pdf_pages(path: Path) -> int:
    try:
        from pypdf import PdfReader

        return len(PdfReader(path).pages)
    except ImportError:
        data = path.read_bytes()
        count = len(re.findall(rb"/Type\s*/Page\b", data))
        if count == 0:
            raise RuntimeError("Cannot validate PDF page count; install pypdf or provide pdfinfo")
        return count


def _validate_outputs(docx_path: Path = DOCX_PATH, pdf_path: Path = PDF_PATH) -> int:
    _validate_docx_structure(docx_path)
    if not pdf_path.is_file() or pdf_path.stat().st_size < 10_000 or pdf_path.read_bytes()[:5] != b"%PDF-":
        raise RuntimeError("PDF output is missing, invalid or unexpectedly small")
    pages = _count_pdf_pages(pdf_path)
    if not 18 <= pages <= 25:
        raise RuntimeError(f"Report has {pages} pages; required range is 18-25")
    return pages


def main() -> int:
    evidence = _load_evidence(ROOT)
    docx = build_docx(evidence, ROOT)
    pdf = build_pdf(docx)
    pages = _validate_outputs(docx, pdf)
    print(f"DOCX: {docx}")
    print(f"PDF: {pdf}")
    print(f"Pages: {pages}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
